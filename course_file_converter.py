"""
Word → Brightspace HTML Converter  +  Padlet → Word Converter
──────────────────────────────────────────────────────────────
Three-tab GUI:
  • Convert      — file picker, output options, live preview (HTML source / rendered)
  • Settings     — element transform dropdowns, blockquote HR toggle, CSS upload
  • Padlet → Word — curriculum-map markdown + Word template → populated .docx

Accordion detection: a single-cell Word table whose cell contains one or more
H4 paragraphs is treated as an accordion group.  Each H4 → card title; the
normal paragraphs that follow it (until the next H4 or end of cell) → card body.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import json, re, tempfile, webbrowser, shutil, subprocess, sys, zipfile, threading
from pathlib import Path


# ═══════════════════════════════════════════════════════════════
#  DEPENDENCY BOOTSTRAP
# ═══════════════════════════════════════════════════════════════

_DEPS = [
    ("docx",        "python-docx",  True),
    ("tkinterdnd2", "tkinterdnd2",  False),
]

def _check_and_install_deps():
    import importlib
    missing_required = []
    missing_optional = []
    for import_name, pip_name, required in _DEPS:
        try:
            importlib.import_module(import_name)
        except ImportError:
            (missing_required if required else missing_optional).append(
                (import_name, pip_name))

    if not missing_required and not missing_optional:
        return True

    to_install = missing_required + missing_optional

    root = tk.Tk()
    root.title("Install Required Packages")
    root.resizable(False, False)
    root.configure(bg="#1a2632")

    PAD = 16
    W   = 520

    tk.Label(root, text="📦  Missing Python Packages",
             font=("Segoe UI", 13, "bold"),
             bg="#1a2632", fg="#ffffff",
             pady=14).pack(fill="x", padx=PAD)

    if missing_required:
        req_names = ", ".join(p for _, p in missing_required)
        desc = (f"The following package{'s are' if len(missing_required) > 1 else ' is'} "
                f"required to run the converter:\n\n  {req_names}")
    else:
        desc = ""
    if missing_optional:
        opt_names = ", ".join(p for _, p in missing_optional)
        opt_note  = (f"\nThe following optional package adds drag-and-drop support:\n\n"
                     f"  {opt_names}")
        desc += opt_note

    tk.Label(root, text=desc.strip(),
             font=("Segoe UI", 10), justify="left",
             bg="#1a2632", fg="#b0bec5",
             wraplength=W - PAD * 2).pack(anchor="w", padx=PAD)

    tk.Frame(root, bg="#3a4a58", height=1).pack(fill="x", padx=PAD, pady=(12, 0))

    log_frame = tk.Frame(root, bg="#0f1117")
    log_frame.pack(fill="both", padx=PAD, pady=(0, 0))

    log_text = tk.Text(
        log_frame, height=10, width=62,
        font=("Cascadia Code", 9),
        bg="#0f1117", fg="#e2e8f0",
        relief="flat", bd=0, padx=10, pady=8,
        state="disabled", wrap="word")
    log_vsb = ttk.Scrollbar(log_frame, orient="vertical", command=log_text.yview)
    log_text.configure(yscrollcommand=log_vsb.set)
    log_vsb.pack(side="right", fill="y")
    log_text.pack(fill="both", expand=True)

    btn_frame  = tk.Frame(root, bg="#1a2632")
    btn_frame.pack(fill="x", padx=PAD, pady=12)

    status_var = tk.StringVar(
        value=f"Ready to install {len(to_install)} package"
              f"{'s' if len(to_install) > 1 else ''}.")
    tk.Label(btn_frame, textvariable=status_var,
             font=("Segoe UI", 9), bg="#1a2632", fg="#78909c",
             anchor="w").pack(side="left", expand=True, fill="x")

    install_btn = tk.Button(btn_frame)
    skip_btn    = tk.Button(btn_frame)
    _result     = {"ok": False}

    def _log(line, colour="#e2e8f0"):
        log_text.config(state="normal")
        log_text.insert("end", line + "\n", (colour,))
        log_text.tag_configure(colour, foreground=colour)
        log_text.see("end")
        log_text.config(state="disabled")
        root.update_idletasks()

    def _run_install():
        install_btn.config(state="disabled", text="Installing…")
        skip_btn.config(state="disabled")
        status_var.set("Installing — please wait…")
        root.update_idletasks()

        all_ok = True
        for import_name, pip_name in to_install:
            _log(f"▶  pip install {pip_name}", "#78909c")
            try:
                proc = subprocess.Popen(
                    [sys.executable, "-m", "pip", "install", pip_name],
                    stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                    text=True, bufsize=1)
                for line in proc.stdout:
                    _log("   " + line.rstrip(), "#94a3b8")
                proc.wait()
                if proc.returncode == 0:
                    _log(f"✅  {pip_name} installed successfully.", "#4ade80")
                else:
                    _log(f"❌  {pip_name} install failed (exit {proc.returncode}).", "#f87171")
                    if any(n == import_name for n, _ in missing_required):
                        all_ok = False
            except Exception as exc:
                _log(f"❌  Error: {exc}", "#f87171")
                all_ok = False

        if all_ok:
            status_var.set("All packages installed. Starting the app…")
            _log("\n✅  Done! The converter will now open.", "#4ade80")
            _result["ok"] = True
            root.after(1200, root.destroy)
        else:
            status_var.set("Some required packages failed to install.")
            install_btn.config(state="normal", text="Retry")
            skip_btn.config(state="normal")

    def _skip():
        if missing_required:
            if not messagebox.askyesno(
                "Skip installation?",
                "Required packages are missing. The app will not work correctly "
                "without them.\n\nSkip anyway?",
                parent=root):
                return
        root.destroy()

    install_btn.config(
        text=f"Install {len(to_install)} Package{'s' if len(to_install) > 1 else ''}",
        font=("Segoe UI", 10, "bold"),
        bg="#006fbf", fg="#ffffff", relief="flat", bd=0,
        padx=14, pady=6, cursor="hand2",
        activebackground="#004a8f", activeforeground="#ffffff",
        command=_run_install)
    install_btn.pack(side="right", padx=(8, 0))

    skip_btn.config(
        text="Skip",
        font=("Segoe UI", 10),
        bg="#2a2f45", fg="#94a3b8", relief="flat", bd=0,
        padx=10, pady=6, cursor="hand2",
        activebackground="#383d56", activeforeground="#e2e8f0",
        command=_skip)
    skip_btn.pack(side="right")

    root.update_idletasks()
    sw, sh = root.winfo_screenwidth(), root.winfo_screenheight()
    rw, rh = root.winfo_width(), root.winfo_height()
    root.geometry(f"+{(sw - rw) // 2}+{(sh - rh) // 2}")

    root.mainloop()
    return _result["ok"]


if __name__ == "__main__":
    if not _check_and_install_deps():
        sys.exit(1)

from docx import Document
from docx.oxml.ns import qn


# ═══════════════════════════════════════════════════════════════
#  PADLET → WORD  CONVERSION LOGIC  (non-GUI)
#  Ported from padlet_to_docx.py — all GUI removed, pure logic.
# ═══════════════════════════════════════════════════════════════

import os
import urllib.parse
from copy import deepcopy
from dataclasses import dataclass, field
from html.parser import HTMLParser
from docx.oxml import OxmlElement


@dataclass
class _Span:
    text:   str
    bold:   bool = False
    italic: bool = False
    href:   str  = ""


@dataclass
class _Para:
    spans:       list = field(default_factory=list)
    list_type:   object = None
    list_level:  int = 0
    list_number: object = None

    def plain_text(self):
        return "".join(s.text for s in self.spans)


class _HTMLToParas(HTMLParser):
    def __init__(self):
        super().__init__()
        self._paras      = []
        self._current    = None
        self._list_stack = []
        self._ol_ctrs    = []
        self._bold  = 0
        self._italic = 0
        self._href   = ""
        self._li_type   = None
        self._li_level  = 0
        self._li_number = None
        self._li_ctx_stack = []

    def _ensure_para(self):
        if self._current is None:
            self._current = _Para(
                list_type=self._li_type,
                list_level=self._li_level,
                list_number=self._li_number)

    def _push_para(self):
        if self._current is not None and self._current.spans:
            self._paras.append(self._current)
        self._current = None

    def _add_span(self, text):
        if not text:
            return
        self._ensure_para()
        self._current.spans.append(_Span(
            text=text, bold=bool(self._bold),
            italic=bool(self._italic), href=self._href))

    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs)
        if tag in ("p", "h1", "h2", "h3", "h4", "h5", "h6"):
            self._push_para(); self._ensure_para()
        elif tag == "li":
            self._push_para()
            level = len(self._list_stack) - 1
            ltype = self._list_stack[-1] if self._list_stack else "ul"
            self._li_type  = ltype; self._li_level = level
            if ltype == "ol":
                self._ol_ctrs[-1] += 1; self._li_number = self._ol_ctrs[-1]
            else:
                self._li_number = None
            self._ensure_para()
        elif tag == "ul":
            self._li_ctx_stack.append((self._li_type, self._li_level, self._li_number))
            self._list_stack.append("ul"); self._ol_ctrs.append(0)
        elif tag == "ol":
            self._li_ctx_stack.append((self._li_type, self._li_level, self._li_number))
            self._list_stack.append("ol"); self._ol_ctrs.append(0)
        elif tag == "br":
            self._push_para()
        elif tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1
        elif tag == "a":
            href = attrs.get("href", "").strip()
            self._href = _p_resolve_padlet_url(href) if href else ""

    def handle_endtag(self, tag):
        if tag in ("h1","h2","h3","h4","h5","h6","p"):
            self._push_para()
        elif tag == "li":
            self._push_para(); self._li_type = self._li_level = self._li_number = None
        elif tag in ("ul", "ol"):
            self._push_para()
            if self._list_stack: self._list_stack.pop()
            if self._ol_ctrs:    self._ol_ctrs.pop()
            if self._li_ctx_stack:
                self._li_type, self._li_level, self._li_number = self._li_ctx_stack.pop()
            else:
                self._li_type = None
        elif tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)
        elif tag == "a":
            self._href = ""

    def handle_data(self, data):
        self._add_span(data.replace("\xa0", " "))

    def handle_entityref(self, name):
        ch = {"amp":"&","lt":"<","gt":">","nbsp":" ","quot":'"',"apos":"'"}.get(name,"")
        if ch: self._add_span(ch)

    def handle_charref(self, name):
        try:
            ch = chr(int(name[1:], 16) if name.startswith("x") else int(name))
        except ValueError:
            ch = ""
        if ch: self._add_span(ch)

    def result(self):
        self._push_para()
        return [p for p in self._paras if p.plain_text().strip()]


def _p_html_to_paras(html):
    if not html or not html.strip():
        return []
    p = _HTMLToParas(); p.feed(html)
    return p.result()


_P_MOD_RE    = re.compile(r"^## Module \d+:\s*(.+)", re.MULTILINE)
_P_H3_RE     = re.compile(r"^### \d+\.\s*(.+)")
_P_ATTACH_RE = re.compile(r"^\[Attachment \d+\]\((.+?)\)")
_P_SKIP_COLORS = {"Purple", "Red", "Green", "Blue"}


def _p_resolve_padlet_url(url):
    is_signed_gcs = (
        "storage.googleapis.com" in url
        and "Signature=" in url
        and "original-url=" in url)
    if is_signed_gcs:
        return re.sub(r"&original-url=[^&)]*", "", url).strip(")")
    orig = re.search(r"original-url=([^&)]+)", url)
    if orig:
        return urllib.parse.unquote(orig.group(1))
    return url


def _p_display_text_for_url(url):
    if "storage.googleapis.com" in url or "padlet-uploads" in url:
        path = url.split("?")[0]
        filename = path.rstrip("/").split("/")[-1]
        if filename:
            return urllib.parse.unquote(filename)
    return re.sub(r"^https?://", "", url)[:80]


def _p_get_color(lines):
    for ln in lines:
        m = re.search(r"\*\*Post color:\*\*\s*(\w+)", ln)
        if m: return m.group(1)
    return None


def _p_get_html_body(lines):
    parts, in_add = [], False
    for ln in lines:
        if ln.strip().startswith("#### Additional"): in_add = True
        if in_add: continue
        s = ln.strip()
        if s.startswith("<") or s.startswith("</"): parts.append(ln)
    return "\n".join(parts).strip()


def _p_get_attachments(lines):
    urls, in_add = [], False
    for ln in lines:
        if ln.strip().startswith("#### Additional"): in_add = True
        if in_add: continue
        m = _P_ATTACH_RE.match(ln.strip())
        if m:
            url = _p_resolve_padlet_url(m.group(1))
            urls.append(url)
    return urls


def _p_clean_title(raw):
    s = raw.strip()
    s = re.sub(r"^\d+\.\s*",                              "", s)
    s = re.sub(r"^module culminating task[:\s]*",         "", s, flags=re.IGNORECASE)
    s = re.sub(r"^course culminating task[:\s]+",         "", s, flags=re.IGNORECASE)
    s = re.sub(r"^\[optional\]\s*task \d+[:\s]*",         "", s, flags=re.IGNORECASE)
    s = re.sub(r"^task \d+[:\s]*",                        "", s, flags=re.IGNORECASE)
    s = re.sub(r"  +",                                    " ", s)
    return s.strip()


def _p_classify(title, color):
    t = title.lower()
    for pat in ("progress update","final report","additional key resource",
                "additional planning resource","guiding resource",
                "other key resources","course tasks","[see earlier note]"):
        if pat in t: return "skip"
    if "course learning objectives" in t:  return "course_obj"
    if "course culminating task"    in t:  return "course_culm"
    if "learning objectives"        in t:  return "mod_obj"
    if "module culminating task"    in t:  return "mod_culm"
    if color in _P_SKIP_COLORS:            return "skip"
    if re.search(r"\btask \d+", t):
        return "task_intro" if "introduce yourself" in t else "task"
    if color == "Yellow":                  return "mod_culm"
    return "skip"


def _p_collect_links(paras, attachment_urls):
    seen = set(); links = []
    for p in paras:
        for span in p.spans:
            if span.href and span.href not in seen:
                seen.add(span.href); links.append(span.href)
    for url in attachment_urls:
        if url and url not in seen:
            seen.add(url); links.append(url)
    return links


def _p_hoist_links(paras, links):
    if not links: return paras
    label_para = _Para()
    label_para.spans.append(_Span(text="Resources:", bold=True))
    url_paras = []
    for url in links:
        p = _Para()
        display = _p_display_text_for_url(url)
        p.spans.append(_Span(text=display, href=url))
        url_paras.append(p)
    return [label_para] + url_paras + paras


def _p_extract_list_paras(paras):
    return [p for p in paras if p.list_type in ("ul","ol") and p.plain_text().strip()]


def p_parse_markdown(md_path):
    with open(md_path, encoding="utf-8") as f:
        raw = f.read()

    warnings = []
    result = {"course_objectives":[], "course_culminating":None, "modules":[]}
    current_module = None

    for section in re.split(r"\n---\n", raw):
        lines = section.strip().splitlines()
        if not lines: continue

        for ln in lines:
            m = _P_MOD_RE.match(ln.strip())
            if m:
                mod_title = re.sub(r"  +", " ", m.group(1).strip())
                current_module = {"title":mod_title,"objectives":[],"tasks":[],"module_culminating":None}
                result["modules"].append(current_module)
                break

        h3_raw = None
        for ln in lines:
            h = _P_H3_RE.match(ln.strip())
            if h: h3_raw = h.group(1).strip(); break
        if not h3_raw: continue

        color = _p_get_color(lines)
        kind  = _p_classify(h3_raw, color)
        if kind == "skip": continue

        html_body   = _p_get_html_body(lines)
        attachments = _p_get_attachments(lines)
        paras       = _p_html_to_paras(html_body)
        links       = _p_collect_links(paras, attachments)
        clean       = _p_clean_title(h3_raw)

        if kind == "course_obj":
            result["course_objectives"] = _p_extract_list_paras(paras)
        elif kind == "course_culm":
            paras = _p_hoist_links(paras, links)
            result["course_culminating"] = {"title":clean,"paras":paras}
        elif kind == "mod_obj" and current_module is not None:
            current_module["objectives"] = _p_extract_list_paras(paras)
        elif kind == "mod_culm" and current_module is not None:
            paras = _p_hoist_links(paras, links)
            current_module["module_culminating"] = {"title":clean,"paras":paras}
        elif kind in ("task","task_intro") and current_module is not None:
            paras = _p_hoist_links(paras, links)
            current_module["tasks"].append({"title":clean,"paras":paras,"is_intro":kind=="task_intro"})
        elif current_module is None and kind not in ("course_obj","course_culm"):
            warnings.append(f"Card '{h3_raw[:60]}' appears before any module heading — skipped.")

    return result, warnings


# ── Word helpers for Padlet converter ─────────────────────────

def _p_get_numbering_ids(doc):
    try:
        num_xml = doc.part.numbering_part._element
        bullet_id = None
        for p in doc.paragraphs:
            if p.style.name == "List Paragraph":
                numPr = p._element.find(qn("w:pPr") + "/" + qn("w:numPr"))
                if numPr is None:
                    pPr = p._element.find(qn("w:pPr"))
                    if pPr is not None: numPr = pPr.find(qn("w:numPr"))
                if numPr is None:
                    numPr = p._element.find(
                        ".//{%s}numPr" % "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
                if numPr is not None:
                    numId_el = numPr.find(qn("w:numId"))
                    if numId_el is not None:
                        nid = numId_el.get(qn("w:val"),"")
                        if nid and nid != "0":
                            bullet_id = int(nid); break

        numbered_id = None
        for an in num_xml.findall(qn("w:abstractNum")):
            lvl0 = an.find(qn("w:lvl"))
            if lvl0 is None: continue
            fmt = lvl0.find(qn("w:numFmt"))
            if fmt is None or fmt.get(qn("w:val")) != "decimal": continue
            ind = lvl0.find(qn("w:pPr") + "/" + qn("w:ind"))
            if ind is None:
                pPr = lvl0.find(qn("w:pPr"))
                if pPr is not None: ind = pPr.find(qn("w:ind"))
            if ind is None: continue
            left = ind.get(qn("w:left"),""); hanging = ind.get(qn("w:hanging"),"")
            if left == "720" and hanging == "360":
                abs_id = an.get(qn("w:abstractNumId"))
                for num in num_xml.findall(qn("w:num")):
                    ref = num.find(qn("w:abstractNumId"))
                    if ref is not None and ref.get(qn("w:val")) == abs_id:
                        numbered_id = int(num.get(qn("w:numId"))); break
                if numbered_id is not None: break

        return (bullet_id or 1, numbered_id or 2)
    except Exception:
        return (1, 2)


def _p_add_hyperlink_rel(doc, url):
    part = doc.part
    HL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    for rId, rel in part.rels.items():
        try:
            if rel.reltype == HL_TYPE and rel.target_ref == url:
                return rId
        except Exception:
            pass
    return part.relate_to(url, HL_TYPE, is_external=True)


def _p_make_run_xml(span):
    r = OxmlElement("w:r")
    if span.bold or span.italic:
        rPr = OxmlElement("w:rPr")
        if span.bold:   rPr.append(OxmlElement("w:b"))
        if span.italic: rPr.append(OxmlElement("w:i"))
        r.append(rPr)
    t = OxmlElement("w:t"); t.text = span.text
    if span.text and (span.text[0]==" " or span.text[-1]==" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
    r.append(t); return r


def _p_make_hyperlink_xml(span, rId):
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), rId); hl.set(qn("w:history"),"1")
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle"); style.set(qn("w:val"),"Hyperlink")
    rPr.append(style)
    if span.bold:   rPr.append(OxmlElement("w:b"))
    if span.italic: rPr.append(OxmlElement("w:i"))
    r.append(rPr); t = OxmlElement("w:t"); t.text = span.text
    if span.text and (span.text[0]==" " or span.text[-1]==" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
    r.append(t); hl.append(r); return hl


def _p_get_abstract_num_id(doc, concrete_num_id):
    num_xml = doc.part.numbering_part._element
    for num in num_xml.findall(qn("w:num")):
        if num.get(qn("w:numId")) == str(concrete_num_id):
            ref = num.find(qn("w:abstractNumId"))
            if ref is not None: return int(ref.get(qn("w:val"),0))
    return None


def _p_clone_numbering_id(doc, abstract_num_id):
    num_xml = doc.part.numbering_part._element
    existing = [int(n.get(qn("w:numId"),0)) for n in num_xml.findall(qn("w:num"))]
    new_id = max(existing, default=0) + 1
    new_num = OxmlElement("w:num"); new_num.set(qn("w:numId"),str(new_id))
    abs_ref = OxmlElement("w:abstractNumId"); abs_ref.set(qn("w:val"),str(abstract_num_id))
    new_num.append(abs_ref)
    lvl_override = OxmlElement("w:lvlOverride"); lvl_override.set(qn("w:ilvl"),"0")
    start_override = OxmlElement("w:startOverride"); start_override.set(qn("w:val"),"1")
    lvl_override.append(start_override); new_num.append(lvl_override)
    num_xml.append(new_num); return new_id


def _p_build_para_xml(para, bullet_id, numbered_id, doc_part):
    new_p = OxmlElement("w:p"); pPr = OxmlElement("w:pPr")
    if para.list_type is not None:
        ps = OxmlElement("w:pStyle"); ps.set(qn("w:val"),"ListParagraph"); pPr.append(ps)
        numPr = OxmlElement("w:numPr")
        ilvl  = OxmlElement("w:ilvl");  ilvl.set(qn("w:val"),str(para.list_level)); numPr.append(ilvl)
        numId = OxmlElement("w:numId"); numId.set(qn("w:val"),str(numbered_id if para.list_type=="ol" else bullet_id)); numPr.append(numId)
        pPr.append(numPr)
    else:
        ps = OxmlElement("w:pStyle"); ps.set(qn("w:val"),"Normal"); pPr.append(ps)
    new_p.append(pPr)
    for span in para.spans:
        if not span.text: continue
        if span.href:
            try:
                rId = _p_add_hyperlink_rel(doc_part, span.href)
                new_p.append(_p_make_hyperlink_xml(span, rId))
            except Exception:
                new_p.append(_p_make_run_xml(span))
        else:
            new_p.append(_p_make_run_xml(span))
    return new_p


def _p_build_rich_paras(rich_paras, bullet_id, numbered_id, doc):
    result = []; in_ol = False; current_ol_id = numbered_id
    abstract_id = _p_get_abstract_num_id(doc, numbered_id)
    for para in rich_paras:
        if not para.plain_text().strip(): continue
        if para.list_type == "ol":
            if not in_ol:
                if abstract_id is not None:
                    current_ol_id = _p_clone_numbering_id(doc, abstract_id)
                else:
                    current_ol_id = numbered_id
                in_ol = True
        else:
            in_ol = False
        result.append(_p_build_para_xml(para, bullet_id, current_ol_id, doc))
    return result


def _p_has_page_break(el):
    from lxml import etree
    try:
        return b'w:type="page"' in etree.tostring(el)
    except Exception:
        return False


def _p_replace_placeholder_in_xml(el, placeholder, new_text):
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W_R = f"{{{NS}}}r"; W_T = f"{{{NS}}}t"
    XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

    def _run_text(r): return "".join((t.text or "") for t in r.findall(W_T))

    def _set_run_text(r, text):
        t_els = r.findall(W_T)
        if not t_els:
            t = OxmlElement("w:t"); r.append(t); t_els = [t]
        t_els[0].text = text
        if text and (text[0]==" " or text[-1]==" "):
            t_els[0].set(XML_SPACE,"preserve")
        else:
            t_els[0].attrib.pop(XML_SPACE, None)
        for t in t_els[1:]: t.text = ""

    runs = list(el.iter(W_R))
    if not runs: return
    run_texts = [_run_text(r) for r in runs]
    full = "".join(run_texts)
    if placeholder not in full: return
    p_start = full.index(placeholder); p_end = p_start + len(placeholder)
    cursor = 0
    for r, rt in zip(runs, run_texts):
        r_start = cursor; r_end = cursor + len(rt)
        if r_end <= p_start or r_start >= p_end:
            cursor = r_end; continue
        overlap_start = max(r_start, p_start); overlap_end = min(r_end, p_end)
        prefix = rt[:overlap_start - r_start]; suffix = rt[overlap_end - r_start:]
        is_first = r_start <= p_start; is_last = r_end >= p_end
        if is_first:
            _set_run_text(r, prefix + new_text + (suffix if is_last else ""))
        else:
            _set_run_text(r, suffix if is_last else "")
        cursor = r_end


def _p_replace_in_para(paragraph, placeholder, new_text):
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W_T = f"{{{NS}}}t"; el = paragraph._element
    full = "".join((t.text or "") for r in el.iter(f"{{{NS}}}r") for t in r.findall(W_T))
    if placeholder not in full: return False
    _p_replace_placeholder_in_xml(el, placeholder, new_text); return True


def _p_strip_para_text_el(el):
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    dead = []
    for r in el.findall(f"{{{NS}}}r"):
        for t in r.findall(f"{{{NS}}}t"): r.remove(t)
        if len(list(r)) == 0: dead.append(r)
    for r in dead: el.remove(r)


@dataclass
class _POp:
    kind: str
    data: dict = field(default_factory=dict)


def _p_slot_type(text):
    t = text.lower()
    if t.startswith("introduction"):   return "introduction"
    if "course culminating task" in t: return "course_culm"
    if "module culminating task" in t: return "mod_culm"
    if "[task title]" in t:            return "task_slot"
    return "other"


def p_plan_operations(parsed, template_paragraphs, warnings):
    modules = parsed["modules"]; course_culm = parsed["course_culminating"]
    ops = []; mod_iter = iter(enumerate(modules))
    current_mod = None; task_queue = []; tmpl_task_slots = []; last_task_slot_xml = None

    tmpl_mod_count = sum(1 for p in template_paragraphs
                         if p.style.name == "Heading 1" and "[module title]" in p.text)
    if len(modules) > tmpl_mod_count:
        warnings.append(
            f"Markdown has {len(modules)} modules but template only has "
            f"{tmpl_mod_count} module slots — "
            f"{len(modules) - tmpl_mod_count} module(s) will be omitted.")

    def _flush_unused_slots(before_idx):
        for si in tmpl_task_slots:
            candidate = si + 1; bi = None
            if candidate < len(template_paragraphs):
                cp = template_paragraphs[candidate]
                if (cp.style.name == "Normal"
                        and not cp.style.name.startswith("Heading")
                        and not _p_has_page_break(cp._element)):
                    bi = candidate
            ops.append(_POp("delete_slot",{
                "heading_idx":si,"body_idx":bi,
                "heading_el":template_paragraphs[si]._element,
                "body_el":template_paragraphs[bi]._element if bi is not None else None}))
        tmpl_task_slots.clear()

    i = 0
    while i < len(template_paragraphs):
        para = template_paragraphs[i]; stripped = para.text.strip()

        if para.style.name == "Heading 1" and "[module title]" in stripped:
            _flush_unused_slots(i); task_queue.clear(); tmpl_task_slots.clear()
            entry = next(mod_iter, None)
            if entry is None:
                warnings.append(f"Template module slot at paragraph {i} has no corresponding markdown module.")
                i += 1; continue
            _, current_mod = entry
            ops.append(_POp("fill_title",{"para_idx":i,"placeholder":"[module title]","new_text":current_mod["title"]}))
            task_queue = list(current_mod["tasks"]); last_task_slot_xml = None; i += 1; continue

        if (para.style.name == "Normal"
                and "in this module you will have the opportunity to" in stripped.lower()
                and current_mod is not None):
            lp_indices = []; j = i + 1
            while j < len(template_paragraphs) and template_paragraphs[j].style.name == "List Paragraph":
                lp_indices.append(j); j += 1
            ops.append(_POp("fill_objectives",{"lp_indices":lp_indices,"obj_paras":current_mod["objectives"]}))
            i += 1; continue

        if (para.style.name == "Normal"
                and "in this course you will have the opportunity to" in stripped.lower()):
            lp_indices = []; j = i + 1
            while j < len(template_paragraphs) and template_paragraphs[j].style.name == "List Paragraph":
                lp_indices.append(j); j += 1
            ops.append(_POp("fill_objectives",{"lp_indices":lp_indices,"obj_paras":parsed["course_objectives"]}))
            i += 1; continue

        if para.style.name == "Heading 2":
            slot = _p_slot_type(stripped)

            if slot == "course_culm":
                body_idx = i+1 if (i+1 < len(template_paragraphs) and template_paragraphs[i+1].style.name=="Normal") else None
                if course_culm:
                    ops.append(_POp("fill_title",{"para_idx":i,"placeholder":"[task title]","new_text":course_culm["title"]}))
                    ops.append(_POp("fill_body",{
                        "heading_idx":i,"body_idx":body_idx,
                        "heading_el":template_paragraphs[i]._element,
                        "body_el":template_paragraphs[body_idx]._element if body_idx is not None else None,
                        "paras":course_culm["paras"]}))
                i += 1; continue

            if slot == "mod_culm" and current_mod is not None:
                non_intro = [t for t in task_queue if not t.get("is_intro")]
                if non_intro and last_task_slot_xml is None:
                    warnings.append(f"Module '{current_mod['title'][:40]}': {len(non_intro)} task(s) could not be inserted — no task-slot heading found to clone from.")
                for t in non_intro:
                    ops.append(_POp("insert_overflow_task",{
                        "before_idx":i,"before_el":template_paragraphs[i]._element,
                        "tmpl_h2_xml":deepcopy(last_task_slot_xml) if last_task_slot_xml is not None else None,
                        "task":t}))
                task_queue.clear(); _flush_unused_slots(i)
                culm = current_mod.get("module_culminating")
                body_idx = i+1 if (i+1 < len(template_paragraphs) and template_paragraphs[i+1].style.name=="Normal") else None
                if culm:
                    # Always emit fill_title (empty string title strips the placeholder cleanly)
                    ops.append(_POp("fill_title",{"para_idx":i,"placeholder":"[task title]","new_text":culm["title"]}))
                    ops.append(_POp("fill_body",{
                        "heading_idx":i,"body_idx":body_idx,
                        "heading_el":template_paragraphs[i]._element,
                        "body_el":template_paragraphs[body_idx]._element if body_idx is not None else None,
                        "paras":culm["paras"]}))
                else:
                    ops.append(_POp("delete_slot",{
                        "heading_idx":i,"body_idx":body_idx,
                        "heading_el":template_paragraphs[i]._element,
                        "body_el":template_paragraphs[body_idx]._element if body_idx is not None else None}))
                i += 1; continue

            if slot == "introduction":
                intro = next((t for t in task_queue if t.get("is_intro")), None)
                if intro: task_queue = [t for t in task_queue if t is not intro]
                body_idx = i+1 if (i+1 < len(template_paragraphs) and template_paragraphs[i+1].style.name=="Normal") else None
                if intro:
                    ops.append(_POp("fill_body",{
                        "heading_idx":i,"body_idx":body_idx,
                        "heading_el":template_paragraphs[i]._element,
                        "body_el":template_paragraphs[body_idx]._element if body_idx is not None else None,
                        "paras":intro["paras"]}))
                i += 1; continue

            if slot == "task_slot" and current_mod is not None:
                non_intro = [t for t in task_queue if not t.get("is_intro")]
                body_idx  = i+1 if (i+1 < len(template_paragraphs) and template_paragraphs[i+1].style.name=="Normal") else None
                if non_intro:
                    task = non_intro[0]; task_queue = [t for t in task_queue if t is not task]
                    last_task_slot_xml = deepcopy(template_paragraphs[i]._element)
                    ops.append(_POp("fill_title",{"para_idx":i,"placeholder":"[task title]","new_text":task["title"]}))
                    ops.append(_POp("fill_body",{
                        "heading_idx":i,"body_idx":body_idx,
                        "heading_el":template_paragraphs[i]._element,
                        "body_el":template_paragraphs[body_idx]._element if body_idx is not None else None,
                        "paras":task["paras"]}))
                else:
                    tmpl_task_slots.append(i)
                i += 1; continue

        i += 1

    _flush_unused_slots(len(template_paragraphs))
    return ops


def _p_exec_fill_objectives(paras, lp_indices, obj_paras, bullet_id, numbered_id, doc):
    def _write_spans(el, spans):
        for r_el in el.findall(qn("w:r")): el.remove(r_el)
        for span in spans:
            if not span.text.strip(): continue
            r = OxmlElement("w:r")
            if span.bold or span.italic:
                rPr = OxmlElement("w:rPr")
                if span.bold:   rPr.append(OxmlElement("w:b"))
                if span.italic: rPr.append(OxmlElement("w:i"))
                r.append(rPr)
            t = OxmlElement("w:t"); t.text = span.text
            if span.text and (span.text[0]==" " or span.text[-1]==" "):
                t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
            r.append(t); el.append(r)

    valid_objs = [p for p in obj_paras if p.plain_text().strip()]
    slot_els = [paras[idx]._element for idx in lp_indices]

    for k, op in enumerate(valid_objs):
        if k < len(slot_els):
            _write_spans(slot_els[k], op.spans)
        else:
            if not slot_els: break
            last_el = slot_els[-1]; new_el = deepcopy(last_el)
            last_el.addnext(new_el); _write_spans(new_el, op.spans)
            slot_els.append(new_el)

    for el in slot_els[len(valid_objs):]:
        try:
            parent = el.getparent()
            if parent is not None: parent.remove(el)
        except Exception:
            pass


def p_execute_operations(doc, ops, warnings):
    bullet_id, numbered_id = _p_get_numbering_ids(doc)
    paras = list(doc.paragraphs)

    def build_rich(rich_paras):
        return _p_build_rich_paras(rich_paras, bullet_id, numbered_id, doc)

    def insert_after(ref_el, elements):
        for el in reversed(elements): ref_el.addnext(el)

    # Pass 1: stable in-place mutations
    for op in ops:
        if op.kind == "fill_title":
            idx = op.data["para_idx"]
            if idx < len(paras):
                new_text = op.data["new_text"]
                placeholder = op.data["placeholder"]
                if new_text:
                    _p_replace_in_para(paras[idx], placeholder, new_text)
                else:
                    # Blank title: replace placeholder with "" then strip orphaned ": " runs
                    _p_replace_in_para(paras[idx], placeholder, "")
                    # Remove any runs that are purely ": " (separator left when title was empty)
                    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    el = paras[idx]._element
                    for r in list(el.findall(f"{{{NS}}}r")):
                        t_els = r.findall(f"{{{NS}}}t")
                        run_text = "".join((t.text or "") for t in t_els)
                        if run_text.strip() in (":", ":"):
                            el.remove(r)
                        elif run_text in (": ", " :", ": "):
                            el.remove(r)
        elif op.kind == "fill_objectives":
            _p_exec_fill_objectives(paras, op.data["lp_indices"], op.data["obj_paras"],
                                    bullet_id, numbered_id, doc)

    # Pass 2a: fill_body — reverse order
    fill_ops = [op for op in ops if op.kind == "fill_body"]
    for op in sorted(fill_ops, key=lambda o: o.data["heading_idx"], reverse=True):
        rich  = build_rich(op.data["paras"])
        if not rich: continue
        h_el  = op.data.get("heading_el")
        if h_el is None: h_el = paras[op.data["heading_idx"]]._element
        body_el = op.data.get("body_el")
        if body_el is not None and body_el.getparent() is not None:
            if _p_has_page_break(body_el):
                insert_after(h_el, rich); _p_strip_para_text_el(body_el)
            else:
                insert_after(body_el, rich); body_el.getparent().remove(body_el)
        else:
            insert_after(h_el, rich)

    # Pass 2b: overflow insertions
    overflow_ops = [op for op in ops if op.kind == "insert_overflow_task"]
    for op in sorted(overflow_ops, key=lambda o: o.data["before_idx"], reverse=True):
        task = op.data["task"]; before_el = op.data.get("before_el"); tmpl_xml = op.data.get("tmpl_h2_xml")
        if tmpl_xml is None:
            warnings.append(f"Could not find heading template for overflow task '{task['title'][:40]}' — skipped.")
            continue
        new_h2_xml = deepcopy(tmpl_xml)
        if before_el is not None and before_el.getparent() is not None:
            before_el.addprevious(new_h2_xml)
        else:
            warnings.append(f"Overflow task '{task['title'][:40]}' — insertion anchor lost, skipped.")
            continue
        _p_replace_placeholder_in_xml(new_h2_xml, "[task title]", task["title"])
        rich = build_rich(task["paras"]); insert_after(new_h2_xml, rich)

    # Pass 3: deletions
    for op in ops:
        if op.kind == "delete_slot":
            body_el = op.data.get("body_el"); heading_el = op.data.get("heading_el")
            if body_el is not None and body_el.getparent() is not None:
                try: body_el.getparent().remove(body_el)
                except Exception: pass
            if heading_el is not None and heading_el.getparent() is not None:
                try: heading_el.getparent().remove(heading_el)
                except Exception: pass


def p_populate_word_template(markdown_file, template_file, output_file):
    """Top-level entry point: convert and save. Returns list of warning strings."""
    parsed, warnings = p_parse_markdown(markdown_file)
    doc   = Document(template_file)
    paras = list(doc.paragraphs)
    ops = p_plan_operations(parsed, paras, warnings)
    p_execute_operations(doc, ops, warnings)
    doc.save(output_file)
    return warnings


# ═══════════════════════════════════════════════════════════════
#  WORD → BRIGHTSPACE  CONVERSION LOGIC  (unchanged from converter.py)
# ═══════════════════════════════════════════════════════════════

CONFIG_FILE   = Path.home() / ".brightspace_converter_config.json"
PRESETS_FILE  = Path.home() / ".brightspace_converter_presets.json"

DEFAULT_HEADING_MAP = {
    "Title":     "h1",
    "Heading 1": "h2",
    "Heading 2": "h3",
    "Heading 3": "h4",
    "Heading 4": "h5",
    "Heading 5": "h6",
    "Heading 6": "h6",
}

BLOCKQUOTE_STYLES = {"Quote", "Block Text", "Intense Quote"}
ACCORDION_HEADING = "Heading 4"

HEADING_OPTS = ["h1", "h2", "h3", "h4", "h5", "h6", "p", "(skip)"]
LIST_OPTS    = ["ul", "ol"]
BQ_OPTS      = ["blockquote", "p", 'div class="callout"', "(skip)"]


class Tooltip:
    def __init__(self, widget, text, delay=500):
        self.widget = widget; self.text = text; self.delay = delay
        self._id = None; self._tip_win = None
        widget.bind("<Enter>",    self._schedule)
        widget.bind("<Leave>",    self._cancel)
        widget.bind("<ButtonPress>", self._cancel)

    def _schedule(self, event=None):
        self._cancel()
        self._id = self.widget.after(self.delay, self._show)

    def _cancel(self, event=None):
        if self._id: self.widget.after_cancel(self._id); self._id = None
        self._hide()

    def _show(self):
        if self._tip_win: return
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 4
        self._tip_win = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True); tw.wm_geometry(f"+{x}+{y}")
        tw.attributes("-topmost", True)
        tk.Label(tw, text=self.text, justify="left",
                 background="#ffffcc", foreground="#222",
                 relief="solid", borderwidth=1,
                 font=("Segoe UI", 9),
                 wraplength=280, padx=6, pady=4).pack()

    def _hide(self):
        if self._tip_win: self._tip_win.destroy(); self._tip_win = None


def tip(widget, text):
    Tooltip(widget, text); return widget


def escape_html(text):
    if not text: return ""
    return (text.replace("&","&amp;").replace("<","&lt;")
            .replace(">","&gt;").replace('"',"&quot;"))


def style_name(para):
    try: return para.style.name or ""
    except Exception: return ""


def _get_num_pr(para):
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            nid = numPr.find(qn("w:numId")); ilv = numPr.find(qn("w:ilvl"))
            numId = int(nid.get(qn("w:val"),0)) if nid is not None else 0
            ilvl  = int(ilv.get(qn("w:val"),0)) if ilv is not None else 0
            if numId > 0: return numId, ilvl
    return None, None


def _lookup_num_fmt(para, numId, ilvl):
    try:
        np = para.part.numbering_part
        if np is None: return ""
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        abId = None
        for n in np._element.findall(f"{{{ns}}}num"):
            if int(n.get(qn("w:numId"),-1)) == numId:
                ref = n.find(qn("w:abstractNumId"))
                if ref is not None: abId = int(ref.get(qn("w:val"),-1))
                break
        if abId is None: return ""
        for ab in np._element.findall(f"{{{ns}}}abstractNum"):
            if int(ab.get(qn("w:abstractNumId"),-1)) == abId:
                lvls = ab.findall(qn("w:lvl"))
                if ilvl < len(lvls):
                    nf = lvls[ilvl].find(qn("w:numFmt"))
                    if nf is not None: return nf.get(qn("w:val"),"")
    except Exception: pass
    return ""


def is_list_para(para):
    numId, _ = _get_num_pr(para)
    if numId is not None: return True
    sn = style_name(para)
    return sn in ("List Bullet","List Bullet 2","List Bullet 3",
                  "List Number","List Number 2","List Number 3")


def is_ordered_para(para):
    numId, ilvl = _get_num_pr(para)
    if numId is not None:
        fmt = _lookup_num_fmt(para, numId, ilvl)
        if fmt: return fmt in ("decimal","lowerLetter","upperLetter",
                               "lowerRoman","upperRoman","ordinal",
                               "cardinalText","ordinalText")
    return "Number" in style_name(para)


def list_indent_level(para):
    _, ilvl = _get_num_pr(para)
    if ilvl is not None: return ilvl
    m = re.search(r"(\d+)$", style_name(para))
    return max(0, int(m.group(1))-1) if m else 0


def _collect_images(para, image_collector):
    NS_WP   = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    NS_BLIP = "http://schemas.openxmlformats.org/drawingml/2006/main"
    NS_REL  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    BLIP_TAG = f"{{{NS_BLIP}}}blip"; DOC_PR = f"{{{NS_WP}}}docPr"
    rId_to_info = {}
    try:
        part = para.part
        for drawing in para._p.iter(qn("w:drawing")):
            alt_text = ""
            for doc_pr in drawing.iter(DOC_PR):
                alt_text = (doc_pr.get("descr") or doc_pr.get("title") or "").strip(); break
            for blip in drawing.iter(BLIP_TAG):
                rId = blip.get(f"{{{NS_REL}}}embed")
                if not rId or rId in rId_to_info: continue
                try:
                    img_part = part.rels[rId].target_part; blob = img_part.blob
                    fname    = img_part.filename
                    if fname in image_collector and image_collector[fname] != blob:
                        ext = Path(fname).suffix; fname = f"{Path(fname).stem}_{rId}{ext}"
                    rId_to_info[rId] = (fname, alt_text); image_collector[fname] = blob
                except Exception: pass
    except Exception: pass
    return rId_to_info


def _extract_runs(el, hyperlink_map, image_collector=None, rId_to_fname=None):
    parts = []
    for child in el:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            parts.append(run_to_html(child))
        elif tag == "drawing":
            if rId_to_fname is not None:
                NS_BLIP = "http://schemas.openxmlformats.org/drawingml/2006/main"
                NS_REL  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                BLIP_TAG = f"{{{NS_BLIP}}}blip"
                for blip in child.iter(BLIP_TAG):
                    rId = blip.get(f"{{{NS_REL}}}embed")
                    info = rId_to_fname.get(rId) if rId else None
                    if info:
                        fname, alt_text = info
                        safe_src = escape_html(fname); safe_alt = escape_html(alt_text) if alt_text else safe_src
                        parts.append(f'<img src="images/{safe_src}" alt="{safe_alt}" style="max-width:100%;">')
        elif tag == "hyperlink":
            rid = child.get(qn("r:id")); url = hyperlink_map.get(rid,"#")
            inner = _extract_runs(child, hyperlink_map, image_collector, rId_to_fname)
            if inner: parts.append(f'<a href="{escape_html(url)}">{inner}</a>')
        elif tag in ("ins","del","smartTag","sdt","sdtContent"):
            parts.append(_extract_runs(child, hyperlink_map, image_collector, rId_to_fname))
    return "".join(parts)


def para_to_inline_html(para, image_collector=None):
    hyperlink_map = {}
    try:
        for rel in para.part.rels.values():
            if "hyperlink" in rel.reltype: hyperlink_map[rel.rId] = rel._target
    except Exception: pass
    rId_to_fname = None
    if image_collector is not None: rId_to_fname = _collect_images(para, image_collector)
    return _extract_runs(para._p, hyperlink_map, image_collector, rId_to_fname)


def _w_attr(el, local_name):
    v = el.get(qn(f"w:{local_name}"))
    if v is None: v = el.get(local_name)
    return v


def _rpr_flag(rpr, tag):
    if rpr is None: return False
    el = rpr.find(qn(tag))
    if el is None: return False
    val = _w_attr(el,"val")
    return val not in ("0","false","off")


def run_to_html(r):
    rpr = r.find(qn("w:rPr")); parts = []
    for child in r:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "t":
            parts.append(escape_html(child.text or ""))
        elif tag == "br":
            br_type = _w_attr(child,"type") or ""
            if br_type not in ("page","column"): parts.append("<br>")
    text = "".join(parts)
    if not text: return ""
    if rpr is not None:
        if _rpr_flag(rpr,"w:strike"): text = f"<del>{text}</del>"
        if _rpr_flag(rpr,"w:u"):      text = f"<u>{text}</u>"
        if _rpr_flag(rpr,"w:i"):      text = f"<em>{text}</em>"
        if _rpr_flag(rpr,"w:b"):      text = f"<strong>{text}</strong>"
    return text


def render_list(list_paras, force_tag=None, image_collector=None):
    lines = []; stack = []
    for para in list_paras:
        level = list_indent_level(para)
        tag   = force_tag if force_tag else ("ol" if is_ordered_para(para) else "ul")
        inline = para_to_inline_html(para, image_collector)
        while len(stack) <= level:
            lines.append("  " * len(stack) + f"<{tag}>"); stack.append((len(stack),tag))
        while len(stack) > level + 1:
            _, ct = stack.pop(); lines.append("  " * len(stack) + f"</{ct}>")
        lines.append("  " * (level+1) + f"<li>{inline}</li>")
    while stack:
        _, ct = stack.pop(); lines.append("  " * len(stack) + f"</{ct}>")
    return "\n".join(lines)


def _render_accordion_body(paragraphs, image_collector=None):
    html_parts = []; paras = list(paragraphs); i = 0
    while i < len(paras):
        para = paras[i]; sn = style_name(para)
        if sn in BLOCKQUOTE_STYLES:
            collected = []
            while i < len(paras) and style_name(paras[i]) in BLOCKQUOTE_STYLES:
                collected.append(para_to_inline_html(paras[i], image_collector)); i += 1
            html_parts.append("<blockquote>")
            for line in collected: html_parts.append(f"  <p>{line}</p>")
            html_parts.append("</blockquote>"); continue
        if is_list_para(para):
            list_paras = []
            while i < len(paras) and is_list_para(paras[i]):
                list_paras.append(paras[i]); i += 1
            html_parts.append(render_list(list_paras, image_collector=image_collector)); continue
        inline = para_to_inline_html(para, image_collector)
        if not inline.strip(): i += 1; continue
        html_parts.append(f"<p>{inline}</p>"); i += 1
    return html_parts


def render_accordion(cell_paragraphs, acc_heading, image_collector=None):
    cards = []; cur_title = None; cur_body = []
    for para in cell_paragraphs:
        if style_name(para) == acc_heading:
            if cur_title is not None: cards.append((cur_title, list(cur_body)))
            cur_title = para_to_inline_html(para, image_collector); cur_body = []
        else:
            cur_body.append(para)
    if cur_title is not None: cards.append((cur_title, list(cur_body)))
    lines = ['<div class="accordion">']
    for title, body_paras in cards:
        body_lines = _render_accordion_body(body_paras, image_collector)
        lines += ['  <div class="card">','    <div class="card-header">',
                  f'      <h2 class="card-title">{title}</h2>','    </div>',
                  '    <div class="collapse">','      <div class="card-body">']
        lines.extend(f"        {ln}" for ln in body_lines)
        lines += ['      </div>','    </div>','  </div>']
    lines.append('</div>'); return "\n".join(lines)


def is_accordion_table(table, acc_heading):
    try:
        cells = list({id(c): c for row in table.rows for c in row.cells}.values())
        if len(cells) != 1: return False
        return any(style_name(p) == acc_heading for p in cells[0].paragraphs)
    except Exception: return False


def render_table(table):
    lines = ["<table>"]
    for i, row in enumerate(table.rows):
        lines.append("  <tr>"); seen = set()
        for cell in row.cells:
            cid = id(cell)
            if cid in seen: continue
            seen.add(cid); tag = "th" if i == 0 else "td"
            content = " ".join(para_to_inline_html(p) for p in cell.paragraphs).strip()
            lines.append(f"    <{tag}>{content}</{tag}>")
        lines.append("  </tr>")
    lines.append("</table>"); return "\n".join(lines)


def _accept_tracked_changes(doc):
    body = doc.element.body
    ns   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    def _tag(local): return f"{{{ns}}}{local}"
    for _ in range(10):
        changed = False
        for el in body.iter():
            local = el.tag.split("}")[-1] if "}" in el.tag else el.tag
            if local == "ins":
                parent = el.getparent()
                if parent is not None:
                    idx = list(parent).index(el)
                    for i, child in enumerate(list(el)):
                        el.remove(child); parent.insert(idx+i, child)
                    parent.remove(el); changed = True; break
            elif local == "del":
                parent = el.getparent()
                if parent is not None: parent.remove(el); changed = True; break
            elif local in ("rPrChange","pPrChange","sectPrChange","tblPrChange","trPrChange","tcPrChange"):
                parent = el.getparent()
                if parent is not None: parent.remove(el); changed = True; break
        if not changed: break


_URL_RE = re.compile(r"^https?://\S+$", re.IGNORECASE)


def _extract_comments(doc):
    comment_map = {}
    try:
        cp = doc.part.comments_part
        if cp is None: return {}
        for cel in cp._element.iter(qn("w:comment")):
            cid = cel.get(qn("w:id"))
            if cid is None: continue
            texts = [t.text or "" for t in cel.iter(qn("w:t"))]
            comment_map[cid] = "".join(texts).strip()
    except Exception: pass
    return comment_map


def _collect_image_link_annotations(doc):
    comment_map = _extract_comments(doc)
    url_comments = {cid:txt for cid,txt in comment_map.items() if _URL_RE.match(txt)}
    if not url_comments: return []
    NS_BLIP = "http://schemas.openxmlformats.org/drawingml/2006/main"
    NS_REL  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    MOD_RE  = re.compile(r"Module\s+([1-9][0-9]?)\s*:(.+)", re.IGNORECASE|re.DOTALL)
    results = []; current_module = "Document"; open_ids = set()
    for el in doc.element.body:
        local = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if local != "p": continue
        txt = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        m = MOD_RE.match(txt)
        if m: current_module = "Module " + m.group(1) + ": " + m.group(2).strip()
        for crs in el.findall(qn("w:commentRangeStart")):
            cid = crs.get(qn("w:id"))
            if cid in url_comments: open_ids.add(cid)
        if open_ids:
            for drawing in el.iter(qn("w:drawing")):
                for blip in drawing.iter(f"{{{NS_BLIP}}}blip"):
                    rId = blip.get(f"{{{NS_REL}}}embed")
                    if rId:
                        try: fname = doc.part.rels[rId].target_part.filename
                        except Exception: fname = rId
                        for cid in sorted(open_ids):
                            results.append({"module":current_module,"image":fname,"url":url_comments[cid]})
        for cre in el.findall(qn("w:commentRangeEnd")):
            open_ids.discard(cre.get(qn("w:id")))
    return results


def _write_image_links_file(link_annotations, out_path):
    if not link_annotations: return False
    lines = ["Image Link Annotations","="*60,"Generated by Word -> Brightspace Converter",""]
    cur_mod = None
    for entry in link_annotations:
        mod = entry["module"]
        if mod != cur_mod:
            if cur_mod is not None: lines.append("")
            lines.append(f"[ {mod} ]"); cur_mod = mod
        lines.append(f"  {entry['image']}"); lines.append(f"    {entry['url']}")
    out_path.write_text("\n".join(lines), encoding="utf-8"); return True


def convert_docx(docx_path, settings):
    hmap        = settings.get("heading_map", DEFAULT_HEADING_MAP)
    ul_out      = settings.get("ul_transform","ul")
    ol_out      = settings.get("ol_transform","ol")
    bq_out      = settings.get("bq_transform","blockquote")
    bq_hr       = settings.get("bq_hr",False)
    acc_head    = settings.get("accordion_heading",ACCORDION_HEADING)
    strip_style = settings.get("strip_style",True)
    log = []
    def _log(level,msg): log.append({"level":level,"msg":msg})

    try:
        with tempfile.NamedTemporaryFile(suffix=".docx",delete=False) as tmp:
            tmp_path = tmp.name
        shutil.copy2(docx_path, tmp_path); doc = Document(tmp_path)
    except Exception:
        doc = Document(docx_path)
    _accept_tracked_changes(doc)

    image_collector = {}; html_parts = []
    n_headings=0; n_paras=0; n_lists=0; n_blockquotes=0; n_tables=0; n_accordions=0; n_skipped=0
    unknown_styles = {}

    para_iter  = iter(doc.paragraphs); table_iter = iter(doc.tables)
    next_para  = next(para_iter,None); next_table = next(table_iter,None)
    blocks = []
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag=="p" and next_para is not None:
            blocks.append(("para",next_para)); next_para = next(para_iter,None)
        elif tag=="tbl" and next_table is not None:
            blocks.append(("table",next_table)); next_table = next(table_iter,None)

    i = 0
    while i < len(blocks):
        kind,obj = blocks[i]
        if kind == "table":
            if is_accordion_table(obj, acc_head):
                cells = list({id(c):c for row in obj.rows for c in row.cells}.values())
                html_parts.append(render_accordion(cells[0].paragraphs, acc_head, image_collector=image_collector))
                n_accordions += 1
            else:
                html_parts.append(render_table(obj)); n_tables += 1
            i += 1; continue

        para = obj; sn = style_name(para)
        if sn in hmap:
            out_tag = hmap[sn]
            if out_tag != "(skip)":
                inline = para_to_inline_html(para, image_collector)
                html_parts.append(f"<p>{inline}</p>" if out_tag=="p" else f"<{out_tag}>{inline}</{out_tag}>")
                n_headings += 1
            else: n_skipped += 1
            i += 1; continue

        if sn in BLOCKQUOTE_STYLES:
            if bq_out == "(skip)": n_skipped += 1; i += 1; continue
            open_tag,close_tag = _bq_tags(bq_out); collected = []
            while (i<len(blocks) and blocks[i][0]=="para" and style_name(blocks[i][1]) in BLOCKQUOTE_STYLES):
                collected.append(para_to_inline_html(blocks[i][1], image_collector)); i += 1
            html_parts.append(open_tag)
            if bq_hr and bq_out=="blockquote": html_parts.append("  <hr>")
            for line in collected: html_parts.append(f"  <p>{line}</p>")
            if bq_hr and bq_out=="blockquote": html_parts.append("  <hr>")
            html_parts.append(close_tag); n_blockquotes += len(collected); continue

        if is_list_para(para):
            list_paras = []
            while i<len(blocks) and blocks[i][0]=="para" and is_list_para(blocks[i][1]):
                list_paras.append(blocks[i][1]); i += 1
            all_ord = all(is_ordered_para(p) for p in list_paras)
            all_bul = all(not is_ordered_para(p) for p in list_paras)
            force   = ol_out if all_ord else (ul_out if all_bul else None)
            html_parts.append(render_list(list_paras,force_tag=force,image_collector=image_collector))
            n_lists += len(list_paras); continue

        inline = para_to_inline_html(para, image_collector)
        if not inline.strip(): i += 1; continue
        if sn and sn not in ("Normal","Body Text","Default Paragraph Font","No Spacing",""):
            unknown_styles[sn] = unknown_styles.get(sn,0)+1
        html_parts.append(f"<p>{inline}</p>"); n_paras += 1; i += 1

    parts_summary = []
    if n_headings:    parts_summary.append(f"{n_headings} heading{'s' if n_headings!=1 else ''}")
    if n_paras:       parts_summary.append(f"{n_paras} paragraph{'s' if n_paras!=1 else ''}")
    if n_lists:       parts_summary.append(f"{n_lists} list item{'s' if n_lists!=1 else ''}")
    if n_blockquotes: parts_summary.append(f"{n_blockquotes} blockquote line{'s' if n_blockquotes!=1 else ''}")
    if n_tables:      parts_summary.append(f"{n_tables} table{'s' if n_tables!=1 else ''}")
    if n_accordions:  parts_summary.append(f"{n_accordions} accordion{'s' if n_accordions!=1 else ''}")
    _log("info","Converted: "+(",".join(parts_summary) if parts_summary else "no content found"))
    if image_collector:
        _log("info",f"{len(image_collector)} image{'s' if len(image_collector)!=1 else ''} extracted: "+", ".join(image_collector.keys()))
    else:
        _log("info","No images found in document")
    if n_skipped: _log("warn",f"{n_skipped} element{'s' if n_skipped!=1 else ''} skipped (heading or blockquote set to '(skip)' in Settings)")
    for sn,count in sorted(unknown_styles.items(),key=lambda x:-x[1]):
        _log("warn",f"Unrecognised style '{sn}' ({count}×) — converted as plain <p>")

    body = "\n".join(html_parts)
    if strip_style: body = re.sub(r' *style="[^"]*"',"",body)
    link_annotations = _collect_image_link_annotations(doc)
    return body, image_collector, log, link_annotations


def _bq_tags(bq_out):
    if bq_out=="blockquote": return "<blockquote>","</blockquote>"
    if bq_out=="p":          return "<p>","</p>"
    tag_name = bq_out.split()[0]
    return f"<{bq_out}>",f"</{tag_name}>"


def wrap_html(body_html, title="", css=""):
    css_block = f"  <style>\n{css}\n  </style>" if css else ""
    return (
        "<!DOCTYPE html>\n"
        '<html lang="en">\n'
        "<head>\n"
        '  <meta charset="UTF-8">\n'
        '  <meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
        f"  <title>{escape_html(title)}</title>\n"
        f"{css_block}\n"
        "</head>\n"
        "<body>\n"
        f"{body_html}\n"
        "</body>\n"
        "</html>"
    )


MODULE_RE = re.compile(
    r'<(h[1-6])>(\s*Module\s+([1-9][0-9]?)\s*:.*?)</\1>',
    re.IGNORECASE|re.DOTALL)


def split_modules(body_html):
    matches = list(MODULE_RE.finditer(body_html))
    if not matches: return []
    modules = []
    pre = body_html[:matches[0].start()].strip()
    if pre: modules.append({"number":0,"title":"Introduction","html":pre})
    for i,m in enumerate(matches):
        start = m.start(); end = matches[i+1].start() if i+1<len(matches) else len(body_html)
        title = re.sub(r"<[^>]+>","",m.group(2)).strip(); num = int(m.group(3))
        chunk = body_html[start:end].strip()
        modules.append({"number":num,"title":title,"html":chunk})
    return modules


# ═══════════════════════════════════════════════════════════════
#  UI CONSTANTS  (shared dark palette)
# ═══════════════════════════════════════════════════════════════

PAD    = 12
FONT   = ("Segoe UI", 10)
FONT_B = ("Segoe UI", 10, "bold")
FONT_S = ("Segoe UI", 9)
MONO   = ("Cascadia Code", 9)

BG        = "#1a1d27"
BG2       = "#22263a"
BG3       = "#2a2f45"
BORDER    = "#383d56"
ACCENT    = "#6c8ef5"
ACCENT2   = "#a78bfa"
FG        = "#e2e8f0"
FG2       = "#94a3b8"
FG3       = "#64748b"
SUCCESS   = "#4ade80"
WARN      = "#fbbf24"
ERR       = "#f87171"
DROP_BG   = "#252942"
DROP_HL   = "#6c8ef5"
PRE_BG    = "#0f1117"
PRE_FG    = "#e2e8f0"

DEFAULT_PREVIEW_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Lato:ital,wght@0,400;0,700;1,400&display=swap');
html { background: #f5f5f5; min-height: 100%; }
body { font-family: 'Lato', 'Segoe UI', system-ui, -apple-system, sans-serif;
       font-size: 19px; line-height: 1.6; color: #212121; background: #ffffff;
       max-width: 860px; margin: 28px auto; padding: 32px 40px 48px;
       box-shadow: 0 1px 4px rgba(0,0,0,.14); border-radius: 2px; }
h1,h2,h3,h4,h5,h6 { color:#212121; font-weight:700; line-height:1.3; margin:1.5em 0 0.4em; }
h1{font-size:1.875em} h2{font-size:1.5em;border-bottom:1px solid #e0e0e0;padding-bottom:.25em}
h3{font-size:1.25em} h4{font-size:1.1em} h5{font-size:1em} h6{font-size:.9em;color:#616161}
p{margin:0 0 1em} a{color:#006fbf;text-decoration:underline} a:hover{color:#004a8f}
ul,ol{margin:0 0 1em 1.5em;padding:0} li{margin-bottom:.3em}
blockquote{border-left:4px solid #006fbf;margin:1em 0;padding:.5em 1em;background:#f0f6fb;color:#424242;border-radius:0 3px 3px 0}
table{border-collapse:collapse;width:100%;margin:1em 0 1.5em;font-size:.95em}
th,td{border:1px solid #bdbdbd;padding:9px 13px;text-align:left;vertical-align:top}
th{background:#eeeeee;font-weight:700} tr:nth-child(even) td{background:#fafafa}
img{max-width:100%;height:auto;display:block;margin:.5em 0} hr{border:none;border-top:1px solid #e0e0e0;margin:1.5em 0}
.accordion{border:1px solid #d3d9de;border-radius:4px;overflow:hidden;margin:1em 0 1.5em}
.card{border-bottom:1px solid #d3d9de} .card:last-child{border-bottom:none}
.card-header{background:#f2f3f5;padding:12px 18px} .card-header:hover{background:#e8eaed}
.card-title{margin:0;font-size:1em;font-weight:700;color:#006fbf}
.card-body{padding:14px 18px;background:#ffffff} .card-body p:last-child{margin-bottom:0}
div.callout{background:#fff8e1;border-left:4px solid #f9a825;padding:10px 16px;margin:1em 0;border-radius:0 3px 3px 0;color:#424242}
"""


# ═══════════════════════════════════════════════════════════════
#  MAIN APPLICATION
# ═══════════════════════════════════════════════════════════════

class ConverterApp(tk.Tk):

    def __init__(self):
        super().__init__()
        self.title("Course File Converter")
        self.configure(bg=BG)
        self.resizable(True, True)
        self.minsize(1020, 680)

        self.selected_file   = None
        self.custom_css      = ""
        self.custom_css_path = None
        self.presets         = {}
        self.preset_var      = tk.StringVar(value="")

        self.heading_vars    = {}
        self.ul_var          = tk.StringVar(value="ul")
        self.ol_var          = tk.StringVar(value="ol")
        self.bq_var          = tk.StringVar(value="blockquote")
        self.bq_hr_var       = tk.BooleanVar(value=False)
        self.acc_head_var    = tk.StringVar(value=ACCORDION_HEADING)
        self.full_html_var   = tk.BooleanVar(value=False)
        self.save_next_var   = tk.BooleanVar(value=True)
        self.strip_style_var = tk.BooleanVar(value=True)
        self.para_font_size  = tk.IntVar(value=19)
        self.preview_mode    = tk.StringVar(value="source")
        self._source_wrap    = tk.BooleanVar(value=False)
        self._last_file_dir  = None

        self._p_preview_mode = tk.StringVar(value="rendered")  # Padlet tab preview mode
        self._p_last_md_path    = None
        self._p_last_docx_path  = None

        self.split_modules_var = tk.BooleanVar(value=True)
        self._modules          = []
        self._module_idx       = 0

        self._mode          = tk.StringVar(value="single")
        self._batch_dir     = Path.cwd()
        self._batch_files   = []
        self._batch_running = False

        self._apply_dark_theme()
        self._build_ui()
        self._load_config()
        self._load_presets_file()

    # ── Dark theme ────────────────────────────────────────────

    def _apply_dark_theme(self):
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure(".", background=BG2, foreground=FG,
                        fieldbackground=BG3, troughcolor=BG3,
                        bordercolor=BORDER, darkcolor=BG2, lightcolor=BG2,
                        selectbackground=ACCENT, selectforeground=FG, font=FONT)
        style.configure("TNotebook", background=BG, borderwidth=0)
        style.configure("TNotebook.Tab", background=BG3, foreground=FG2,
                        padding=[14,7], borderwidth=0, font=FONT)
        style.map("TNotebook.Tab",
                  background=[("selected",BG2),("active",BORDER)],
                  foreground=[("selected",FG),("active",FG)])
        style.configure("TCombobox", fieldbackground=BG3, background=BG3,
                        foreground=FG, arrowcolor=FG2,
                        selectbackground=ACCENT, selectforeground=FG,
                        bordercolor=BORDER, insertcolor=FG)
        style.map("TCombobox", fieldbackground=[("readonly",BG3)],
                  foreground=[("readonly",FG)], bordercolor=[("focus",ACCENT)])
        SCROLL_TRACK = "#1e2235"; SCROLL_THUMB = "#4a5175"
        SCROLL_HOVER = "#6c75a8"; SCROLL_PRESS = "#8891c4"
        for orient in ("Vertical","Horizontal"):
            style.configure(f"{orient}.TScrollbar",
                            background=SCROLL_THUMB, troughcolor=SCROLL_TRACK,
                            bordercolor=SCROLL_TRACK, darkcolor=SCROLL_THUMB,
                            lightcolor=SCROLL_THUMB, arrowcolor=SCROLL_TRACK,
                            relief="flat", gripcount=0, width=10)
            style.map(f"{orient}.TScrollbar",
                      background=[("active",SCROLL_HOVER),("pressed",SCROLL_PRESS)])
        self.option_add("*TCombobox*Listbox.background", BG3)
        self.option_add("*TCombobox*Listbox.foreground", FG)
        self.option_add("*TCombobox*Listbox.selectBackground", ACCENT)
        self.option_add("*TCombobox*Listbox.selectForeground", FG)

    # ── Top-level layout ──────────────────────────────────────

    def _build_ui(self):
        hdr = tk.Frame(self, bg=BG2, pady=0); hdr.pack(fill="x")
        tk.Frame(hdr, bg=ACCENT, width=4).pack(side="left", fill="y")
        tk.Label(hdr, text="  Course File Converter",
                 bg=BG2, fg=FG, font=("Segoe UI",13,"bold"), pady=11).pack(side="left")
        tk.Frame(hdr, bg=BORDER, height=1).pack(side="bottom", fill="x")

        pane = tk.PanedWindow(self, orient="horizontal", bg=BG,
                              sashwidth=6, sashrelief="flat", sashpad=0)
        pane.pack(fill="both", expand=True)

        left = tk.Frame(pane, bg=BG2, width=380); left.pack_propagate(False)
        pane.add(left, minsize=320)

        self.nb = ttk.Notebook(left)
        self.nb.pack(fill="both", expand=True)

        tc = tk.Frame(self.nb, bg=BG2, padx=PAD, pady=PAD)
        tp = tk.Frame(self.nb, bg=BG2, padx=PAD, pady=PAD)   # ← Padlet tab
        ts = tk.Frame(self.nb, bg=BG2)

        self.nb.add(tc, text="  Word to Brightspace  ")
        self.nb.add(tp, text="  Padlet to Word  ")
        self.nb.add(ts, text="  Settings  ")

        self._build_convert_tab(tc)
        self._build_padlet_tab(tp)   # ← new
        self._build_settings_tab(ts)

        right = tk.Frame(pane, bg=BG)
        pane.add(right, minsize=380)

        # Both preview panels live in the same right frame; tab-change swaps them
        self._word_preview_frame   = tk.Frame(right, bg=BG)
        self._padlet_preview_frame = tk.Frame(right, bg=BG)
        self._word_preview_frame.pack(fill="both", expand=True)   # visible by default

        self._build_preview_panel(self._word_preview_frame)
        self._build_padlet_preview(self._padlet_preview_frame)

        # Swap right panel when the user switches tabs
        def _on_tab_change(event=None):
            idx = self.nb.index(self.nb.select())
            if idx == 2:   # Settings tab — hide all previews
                self._word_preview_frame.pack_forget()
                self._padlet_preview_frame.pack_forget()
            elif idx == 1:   # Padlet → Word tab
                self._word_preview_frame.pack_forget()
                self._padlet_preview_frame.pack(fill="both", expand=True)
            else:            # Word to Brightspace tab
                self._padlet_preview_frame.pack_forget()
                self._word_preview_frame.pack(fill="both", expand=True)
        self.nb.bind("<<NotebookTabChanged>>", _on_tab_change)

    # ════════════════════════════════════════════════════════
    #  PADLET → WORD  TAB
    # ════════════════════════════════════════════════════════

    def _build_padlet_tab(self, p):
        # Controls fill the tab frame directly; preview lives in the shared right panel
        self._build_padlet_controls(p)

    def _build_padlet_controls(self, p):
        # ── Header description ────────────────────────────────
        desc = tk.Frame(p, bg=BG3, highlightthickness=1, highlightbackground=BORDER)
        desc.pack(fill="x", pady=(0, 16))
        tk.Label(desc,
                 text="Convert a Padlet-exported Markdown curriculum map into a\n"
                      "populated Word document by filling placeholders in a .docx template.",
                 font=FONT_S, bg=BG3, fg=FG2, justify="left",
                 padx=10, pady=8, wraplength=330).pack(anchor="w")

        # ── File rows ─────────────────────────────────────────
        tk.Label(p, text="INPUT FILES", font=("Segoe UI",8,"bold"),
                 bg=BG2, fg=FG3).pack(anchor="w", pady=(0,4))

        self._p_md_var   = tk.StringVar()
        self._p_tmpl_var = tk.StringVar()
        self._p_out_var  = tk.StringVar()

        # Load MD preview as soon as the file path is set (browse or typed)
        self._p_md_var.trace_add("write", lambda *_: self.after(50, self._p_load_md_preview))
        # Load template docx preview as soon as its path is set (Fix 3)
        self._p_tmpl_var.trace_add("write", lambda *_: self.after(50, self._p_load_tmpl_preview))

        self._p_md_row   = self._padlet_file_row(p, "Curriculum map (.md)",
                                                  self._p_md_var,   mode="open",
                                                  filetypes=[("Markdown","*.md"),("All files","*.*")])
        self._p_tmpl_row = self._padlet_file_row(p, "Word template (.docx)",
                                                  self._p_tmpl_var, mode="open",
                                                  filetypes=[("Word document","*.docx")])

        tk.Label(p, text="OUTPUT FILE", font=("Segoe UI",8,"bold"),
                 bg=BG2, fg=FG3).pack(anchor="w", pady=(8,4))

        self._p_out_row  = self._padlet_file_row(p, "Save output as (.docx)",
                                                  self._p_out_var,  mode="save",
                                                  filetypes=[("Word document","*.docx")])

        # ── Status bar ────────────────────────────────────────
        self._p_status_var = tk.StringVar(value="")
        self._p_status_lbl = tk.Label(p, textvariable=self._p_status_var,
                                      bg=BG2, fg=SUCCESS, font=FONT,
                                      wraplength=320, justify="left")
        self._p_status_lbl.pack(anchor="w", pady=(10, 0))

        # ── Warnings log (collapsed by default) ───────────────
        self._p_log_expanded = tk.BooleanVar(value=False)
        log_hdr = tk.Frame(p, bg=BG2, cursor="hand2"); log_hdr.pack(fill="x", pady=(10,0))
        self._p_log_arrow = tk.Label(log_hdr, text="▶", font=FONT_S, bg=BG2, fg=FG3)
        self._p_log_arrow.pack(side="left")
        self._p_log_hdr_lbl = tk.Label(log_hdr, text="CONVERSION LOG",
                                        font=("Segoe UI",8,"bold"), bg=BG2, fg=FG3)
        self._p_log_hdr_lbl.pack(side="left", padx=(4,0))
        self._p_log_badge = tk.Label(log_hdr, text="", font=FONT_S, bg=BG2, fg=FG3)
        self._p_log_badge.pack(side="left", padx=(6,0))
        for w in (log_hdr, self._p_log_arrow, self._p_log_hdr_lbl, self._p_log_badge):
            w.bind("<Button-1>", lambda e: self._p_toggle_log())

        self._p_log_body = tk.Frame(p, bg=BG3)
        self._p_log_text = tk.Text(
            self._p_log_body, font=("Cascadia Code",8), wrap="word",
            bg=BG3, fg=FG2, relief="flat", bd=0,
            padx=8, pady=6, height=8, state="disabled",
            selectbackground=ACCENT, selectforeground=FG, cursor="arrow")
        log_vsb = ttk.Scrollbar(self._p_log_body, orient="vertical", command=self._p_log_text.yview)
        self._p_log_text.configure(yscrollcommand=log_vsb.set)
        log_vsb.pack(side="right", fill="y"); self._p_log_text.pack(fill="both", expand=True)
        self._p_log_text.tag_configure("info",  foreground=FG2)
        self._p_log_text.tag_configure("warn",  foreground=WARN)
        self._p_log_text.tag_configure("error", foreground=ERR)

        # ── Generate button ───────────────────────────────────
        tk.Frame(p, bg=BORDER, height=1).pack(fill="x", pady=(10,0))
        self._p_btn = tk.Button(
            p, text="Generate  →", font=FONT_B,
            bg=ACCENT, fg=FG, relief="flat", bd=0,
            padx=16, pady=10, cursor="hand2",
            activebackground=ACCENT2, activeforeground=FG,
            command=self._p_run)
        self._p_btn.pack(fill="x", pady=(6,0))
        tip(self._p_btn, "Fill the Word template with content from the Padlet markdown file.")

    def _padlet_file_row(self, parent, label, var, mode="open", filetypes=None):
        """A file-pick row matching the existing converter style."""
        filetypes = filetypes or [("All files","*.*")]
        frame = tk.Frame(parent, bg=BG2); frame.pack(fill="x", pady=(0,10))
        frame.columnconfigure(1, weight=1)

        tk.Label(frame, text=label, font=("Segoe UI",9,"bold"),
                 bg=BG2, fg=FG2, width=22, anchor="w").grid(row=0, column=0, padx=(0,8), sticky="w")

        entry = tk.Entry(frame, textvariable=var,
                         font=MONO, bg=BG3, fg=FG, insertbackground=FG,
                         relief="flat", bd=4)
        entry.grid(row=0, column=1, sticky="ew", padx=(0,8))

        def _browse():
            if mode == "save":
                path = filedialog.asksaveasfilename(
                    defaultextension=".docx", filetypes=filetypes,
                    initialdir=self._last_file_dir or str(Path.home()))
            else:
                path = filedialog.askopenfilename(
                    filetypes=filetypes,
                    initialdir=self._last_file_dir or str(Path.home()))
            if path:
                var.set(path)

        tk.Button(frame, text="Browse", font=FONT_S,
                  bg=BG3, fg=FG, relief="flat", bd=0,
                  padx=10, pady=5, cursor="hand2",
                  activebackground=BORDER, activeforeground=FG,
                  command=_browse).grid(row=0, column=2)
        return frame

    def _p_toggle_log(self):
        if self._p_log_expanded.get():
            self._p_log_body.pack_forget()
            self._p_log_expanded.set(False)
            self._p_log_arrow.config(text="▶")
        else:
            self._p_log_body.pack(fill="x", pady=(2,0))
            self._p_log_expanded.set(True)
            self._p_log_arrow.config(text="▼")
            self._p_log_text.see("end")

    def _p_update_log(self, warnings):
        t = self._p_log_text
        t.config(state="normal"); t.delete("1.0","end")
        if not warnings:
            t.insert("end"," ·  No warnings — clean run.\n","info")
            self._p_log_badge.config(text="  OK", fg=SUCCESS)
            if self._p_log_expanded.get(): self._p_toggle_log()
        else:
            for w in warnings:
                t.insert("end"," ⚠  ","warn"); t.insert("end", w+"\n","warn")
            self._p_log_badge.config(text=f"  {len(warnings)} warning{'s' if len(warnings)!=1 else ''}",fg=WARN)
            if not self._p_log_expanded.get(): self._p_toggle_log()
        t.config(state="disabled")

    def _p_run(self):
        md   = self._p_md_var.get().strip()
        tmpl = self._p_tmpl_var.get().strip()
        out  = self._p_out_var.get().strip()

        if not md or not tmpl or not out:
            self._p_status_var.set("⚠  Please fill in all three file paths.")
            self._p_status_lbl.config(fg=ERR); return
        if not os.path.isfile(md):
            self._p_status_var.set("⚠  Markdown file not found.")
            self._p_status_lbl.config(fg=ERR); return
        if not os.path.isfile(tmpl):
            self._p_status_var.set("⚠  Template file not found.")
            self._p_status_lbl.config(fg=ERR); return

        self._p_btn.configure(state="disabled", text="Working…")
        self._p_status_var.set("Converting…"); self._p_status_lbl.config(fg=ACCENT)

        def worker():
            try:
                warns = p_populate_word_template(md, tmpl, out)
                self.after(0, self._p_on_success, out, warns)
            except Exception:
                import traceback
                self.after(0, self._p_on_error, traceback.format_exc())

        threading.Thread(target=worker, daemon=True).start()

    def _p_on_success(self, out_path, warnings):
        short = os.path.basename(out_path)
        self._p_update_log(warnings)
        if warnings:
            self._p_status_var.set(f"✓  Saved — {short}  ({len(warnings)} warning(s))")
            self._p_status_lbl.config(fg=WARN)
        else:
            self._p_status_var.set(f"✓  Saved — {short}")
            self._p_status_lbl.config(fg=SUCCESS)
        self._p_btn.configure(state="normal", text="Generate  →")
        self._p_refresh_preview(out_path)

    def _p_on_error(self, tb):
        first = [l.strip() for l in tb.splitlines() if l.strip()][-1]
        self._p_status_var.set(f"✗  {first}"); self._p_status_lbl.config(fg=ERR)
        self._p_btn.configure(state="normal", text="Generate  →")
        # Show full traceback in an expandable detail window
        win = tk.Toplevel(self); win.title("Error details")
        win.geometry("660x320"); win.configure(fg_color=BG) if hasattr(win,"fg_color") else win.configure(bg=BG)
        tk.Label(win, text="Error:", font=FONT_B, fg=ERR,
                 bg=BG).pack(anchor="w", padx=20, pady=(16,6))
        box = tk.Text(win, font=MONO, bg=BG3, fg=FG,
                      relief="flat", bd=4)
        box.pack(fill="both", expand=True, padx=20, pady=(0,20))
        box.insert("end", tb); box.configure(state="disabled")

    def _build_padlet_preview(self, parent):
        """Right-side preview panel for the Padlet → Word tab."""
        # ── Toolbar ──────────────────────────────────────────
        bar = tk.Frame(parent, bg=BG2); bar.pack(fill="x")
        tk.Frame(bar, bg=BORDER, height=1).pack(side="bottom", fill="x")
        tk.Label(bar, text="  Preview", font=FONT_B, bg=BG2, fg=FG, pady=8).pack(side="left")

        def _mode_btn(text, value):
            rb = tk.Radiobutton(bar, text=f"  {text}  ", variable=self._p_preview_mode, value=value,
                                bg=BG2, fg=FG2, font=FONT,
                                activebackground=BG2, activeforeground=FG,
                                selectcolor=BG3, indicatoron=False,
                                relief="flat", bd=0, padx=8, pady=5, cursor="hand2",
                                command=self._p_on_mode_btn)
            rb.pack(side="right", padx=4, pady=4); return rb

        _mode_btn("Side by Side", "split")
        _mode_btn("MD Source", "md_source")
        _mode_btn("Rendered", "rendered")

        # ── Preview area ──────────────────────────────────────
        self._p_preview_outer = tk.Frame(parent, bg=PRE_BG)
        self._p_preview_outer.pack(fill="both", expand=True)

        D2L_SHELL = "#f5f5f5"

        # Single container (rendered or md_source)
        self._p_single_container = tk.Frame(self._p_preview_outer, bg=PRE_BG)
        self._p_single_container.pack(fill="both", expand=True)

        # MD source view
        self._p_source_frame = tk.Frame(self._p_single_container, bg=PRE_BG)
        self._p_source_text = tk.Text(
            self._p_source_frame, font=("Cascadia Code", 9), wrap="none",
            bg=PRE_BG, fg=PRE_FG, insertbackground=FG, relief="flat", bd=0,
            padx=12, pady=12, selectbackground=ACCENT, selectforeground=FG,
            state="disabled")
        _vsb = ttk.Scrollbar(self._p_source_frame, orient="vertical", command=self._p_source_text.yview)
        _hsb = ttk.Scrollbar(self._p_source_frame, orient="horizontal", command=self._p_source_text.xview)
        self._p_source_text.configure(yscrollcommand=_vsb.set, xscrollcommand=_hsb.set)
        _vsb.pack(side="right", fill="y"); _hsb.pack(side="bottom", fill="x")
        self._p_source_text.pack(fill="both", expand=True)

        # Rendered (Word-like) view
        self._p_rendered_frame = tk.Frame(self._p_single_container, bg=D2L_SHELL)
        _card = tk.Frame(self._p_rendered_frame, bg="#ffffff",
                         highlightbackground="#d0d0d0", highlightthickness=1)
        _card.pack(fill="both", expand=True)
        self._p_rendered_text = tk.Text(
            _card, font=("Lato", 11), wrap="word",
            bg="#ffffff", fg="#212121", relief="flat", bd=0,
            padx=32, pady=28, spacing1=0, spacing2=0, spacing3=0,
            selectbackground="#b3d4f5", selectforeground="#212121",
            cursor="xterm", state="normal", insertwidth=0)
        _rvsb = ttk.Scrollbar(_card, orient="vertical", command=self._p_rendered_text.yview)
        self._p_rendered_text.configure(yscrollcommand=_rvsb.set)
        _rvsb.pack(side="right", fill="y"); self._p_rendered_text.pack(fill="both", expand=True)
        self._p_rendered_text.bind("<Key>", lambda e: "break")
        self._configure_rendered_tags(target=self._p_rendered_text)

        # Split pane
        self._p_split_pane = tk.PanedWindow(self._p_preview_outer, orient="horizontal",
                                             bg=BORDER, sashwidth=5, sashrelief="flat", sashpad=0)
        _sp_left = tk.Frame(self._p_split_pane, bg=PRE_BG)
        self._p_split_pane.add(_sp_left, minsize=180, stretch="always")
        self._p_split_source_text = tk.Text(
            _sp_left, font=("Cascadia Code", 9), wrap="none",
            bg=PRE_BG, fg=PRE_FG, insertbackground=FG, relief="flat", bd=0,
            padx=12, pady=12, selectbackground=ACCENT, selectforeground=FG,
            state="disabled")
        _sp_vsb = ttk.Scrollbar(_sp_left, orient="vertical", command=self._p_split_source_text.yview)
        _sp_hsb = ttk.Scrollbar(_sp_left, orient="horizontal", command=self._p_split_source_text.xview)
        self._p_split_source_text.configure(yscrollcommand=_sp_vsb.set, xscrollcommand=_sp_hsb.set)
        _sp_vsb.pack(side="right", fill="y"); _sp_hsb.pack(side="bottom", fill="x")
        self._p_split_source_text.pack(fill="both", expand=True)

        _sp_right = tk.Frame(self._p_split_pane, bg=D2L_SHELL)
        self._p_split_pane.add(_sp_right, minsize=180, stretch="always")
        _sp_card = tk.Frame(_sp_right, bg="#ffffff", highlightbackground="#d0d0d0", highlightthickness=1)
        _sp_card.pack(fill="both", expand=True)
        self._p_split_rendered_text = tk.Text(
            _sp_card, font=("Lato", 11), wrap="word",
            bg="#ffffff", fg="#212121", relief="flat", bd=0,
            padx=32, pady=28, spacing1=0, spacing2=0, spacing3=0,
            selectbackground="#b3d4f5", selectforeground="#212121",
            cursor="xterm", state="normal", insertwidth=0)
        _sp_rvsb = ttk.Scrollbar(_sp_card, orient="vertical", command=self._p_split_rendered_text.yview)
        self._p_split_rendered_text.configure(yscrollcommand=_sp_rvsb.set)
        _sp_rvsb.pack(side="right", fill="y"); self._p_split_rendered_text.pack(fill="both", expand=True)
        self._p_split_rendered_text.bind("<Key>", lambda e: "break")
        self._configure_rendered_tags(target=self._p_split_rendered_text)

        # Status bar
        bot = tk.Frame(parent, bg=BG2); bot.pack(fill="x")
        tk.Frame(bot, bg=BORDER, height=1).pack(side="top", fill="x")
        self._p_preview_status = tk.Label(bot, text="Generate a .docx to see a preview here.",
                                          font=FONT_S, bg=BG2, fg=FG3, pady=6, padx=10, anchor="w")
        self._p_preview_status.pack(side="left", fill="x", expand=True)

        # Show initial mode
        self._p_show_pane("rendered")
        self._p_set_preview("Generate a document to preview it here.\n\nThe Rendered view shows a Word-style layout.\nMD Source shows the raw Markdown file.")

    def _p_show_pane(self, mode):
        """Switch Padlet preview between rendered, md_source, split."""
        if mode == "split":
            self._p_single_container.pack_forget()
            self._p_split_pane.pack(fill="both", expand=True)
        else:
            self._p_split_pane.pack_forget()
            self._p_single_container.pack(fill="both", expand=True)
            self._p_source_frame.pack_forget()
            self._p_rendered_frame.pack_forget()
            if mode == "md_source":
                self._p_source_frame.pack(fill="both", expand=True)
            else:
                self._p_rendered_frame.pack(fill="both", expand=True)

    def _p_on_mode_btn(self):
        mode = self._p_preview_mode.get()
        self._p_show_pane(mode)
        # Re-populate if we have content
        if hasattr(self, "_p_last_md_path") and self._p_last_md_path:
            self._p_refresh_preview(self._p_last_docx_path)

    def _p_set_preview(self, text):
        """Set placeholder text in all Padlet source/split-source panes."""
        for tw in (self._p_source_text, self._p_split_source_text):
            tw.config(state="normal"); tw.delete("1.0", "end"); tw.insert("1.0", text)
            tw.config(state="disabled")

    def _p_load_md_preview(self):
        """Show MD source immediately when a .md file is selected, before generation."""
        md_path = self._p_md_var.get().strip()
        if not md_path or not Path(md_path).exists():
            return
        self._p_last_md_path = md_path
        try:
            md_text = Path(md_path).read_text(encoding="utf-8", errors="replace")
        except Exception:
            return
        # Populate source panes
        for tw in (self._p_source_text, self._p_split_source_text):
            tw.config(state="normal"); tw.delete("1.0", "end"); tw.insert("1.0", md_text)
            tw.config(state="disabled")
        # If no docx yet, rendered pane shows a prompt; switch to md_source mode automatically
        if not (hasattr(self, "_p_last_docx_path") and self._p_last_docx_path
                and Path(self._p_last_docx_path).exists()):
            self._p_preview_mode.set("md_source")
            self._p_show_pane("md_source")
        fname = Path(md_path).name
        self._p_preview_status.config(text=f"MD loaded: {fname}", fg=FG2)

    def _p_load_tmpl_preview(self):
        """Render the template .docx in the Rendered pane as soon as it is selected (Fix 3)."""
        tmpl_path = self._p_tmpl_var.get().strip()
        if not tmpl_path or not Path(tmpl_path).exists():
            return
        # Only act if no generated output is already showing
        if (hasattr(self, "_p_last_docx_path") and self._p_last_docx_path
                and Path(self._p_last_docx_path).exists()):
            return
        # Reuse the same rendering path as post-generation, passing the template as the docx
        self._p_last_docx_path = tmpl_path
        self._p_preview_mode.set("rendered")
        self._p_show_pane("rendered")
        self._p_refresh_preview(tmpl_path)
        fname = Path(tmpl_path).name
        self._p_preview_status.config(text=f"Template preview: {fname}", fg=FG2)

    def _p_refresh_preview(self, docx_path=None):
        """Populate the Padlet preview pane after generation."""
        mode = self._p_preview_mode.get()
        self._p_show_pane(mode)

        # Load MD source
        md_path = self._p_md_var.get().strip()
        if md_path and Path(md_path).exists():
            self._p_last_md_path = md_path
            try:
                md_text = Path(md_path).read_text(encoding="utf-8", errors="replace")
            except Exception:
                md_text = "(Could not read Markdown file)"
        else:
            md_text = "(No Markdown file loaded)"

        # Load rendered content from the output docx
        rendered_html = ""
        if docx_path and Path(docx_path).exists():
            self._p_last_docx_path = docx_path
            try:
                from docx import Document as _Doc
                doc = _Doc(docx_path)
                html_parts = []

                # ── counters for auto-numbered headings (Fix 1) ──────────
                _mod_counter  = 0
                _task_counter = 0  # resets per module
                _is_first_para = True   # Fix 2: first non-empty para → h1

                # ── list-state tracking (Fix 4) ──────────────────────────
                _in_list      = False   # are we currently inside a <ul>/<ol>?
                _list_is_ol   = False

                def _close_list():
                    nonlocal _in_list
                    if _in_list:
                        html_parts.append("</ol>" if _list_is_ol else "</ul>")
                        _in_list = False

                def _open_list(ordered):
                    nonlocal _in_list, _list_is_ol
                    _close_list()
                    html_parts.append("<ol>" if ordered else "<ul>")
                    _in_list  = True
                    _list_is_ol = ordered

                # ── inline HTML extractor (handles bold/italic/links) ─────
                def _para_inline(para):
                    """Return inline HTML for a paragraph, preserving bold/
                    italic formatting and making hyperlinks clickable (Fix 5)."""
                    hmap = {}
                    try:
                        for rel in para.part.rels.values():
                            if "hyperlink" in rel.reltype:
                                hmap[rel.rId] = rel._target
                    except Exception:
                        pass

                    def _runs(el):
                        out = []
                        for child in el:
                            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                            if tag == "r":
                                rpr  = child.find(qn("w:rPr"))
                                bold = rpr is not None and rpr.find(qn("w:b")) is not None
                                ital = rpr is not None and rpr.find(qn("w:i")) is not None
                                txt  = "".join(
                                    (t.text or "") for t in child.findall(qn("w:t"))
                                )
                                txt = escape_html(txt)
                                if bold and ital: txt = f"<strong><em>{txt}</em></strong>"
                                elif bold:        txt = f"<strong>{txt}</strong>"
                                elif ital:        txt = f"<em>{txt}</em>"
                                out.append(txt)
                            elif tag == "hyperlink":
                                rid  = child.get(qn("r:id"))
                                url  = hmap.get(rid, "#")
                                inner = "".join(_runs(child))
                                if inner:
                                    out.append(f'<a href="{escape_html(url)}">{inner}</a>')
                            elif tag in ("ins", "smartTag", "sdt", "sdtContent"):
                                out.extend(_runs(child))
                        return out

                    return "".join(_runs(para._p))

                for para in doc.paragraphs:
                    sn   = para.style.name if para.style else ""
                    text = para.text.strip()

                    # ── list paragraphs (Fix 4) ──────────────────────────
                    is_list = (
                        sn in ("List Paragraph", "List Bullet", "List Bullet 2",
                               "List Bullet 3", "List Number", "List Number 2",
                               "List Number 3")
                        or is_list_para(para)
                    )
                    if is_list:
                        ordered = is_ordered_para(para)
                        if not _in_list or _list_is_ol != ordered:
                            _open_list(ordered)
                        inline = _para_inline(para)
                        html_parts.append(f"  <li>{inline}</li>")
                        continue

                    # close any open list before a non-list element
                    _close_list()

                    if not text:
                        continue

                    inline = _para_inline(para)

                    # ── Fix 2: very first non-empty paragraph → h1 ───────
                    if _is_first_para:
                        _is_first_para = False
                        # Treat as Heading 1 regardless of Word style
                        sn = "Heading 1"

                    # ── headings with auto-numbering (Fixes 1 & 2) ───────
                    if sn.startswith("Heading 1"):
                        # Heading 1 = Module title — inject "Module N: " prefix
                        _mod_counter += 1
                        _task_counter = 0
                        # If the text already carries a "Module N:" prefix
                        # (written literally into the docx), use it as-is;
                        # otherwise synthesise one.
                        if not re.match(r"Module\s+\d+\s*:", text, re.IGNORECASE):
                            inline = f"Module {_mod_counter}: {inline}"
                        html_parts.append(f"<h1>{inline}</h1>")

                    elif sn.startswith("Heading 2"):
                        # Heading 2 = Task / section title — inject "Task N: " prefix
                        # except for well-known non-task headings
                        NON_TASK = ("introduction", "module culminating", "course culminating")
                        is_non_task = any(k in text.lower() for k in NON_TASK)
                        if not is_non_task:
                            _task_counter += 1
                            if not re.match(r"Task\s+\d+\s*:", text, re.IGNORECASE):
                                inline = f"Task {_task_counter}: {inline}"
                        html_parts.append(f"<h2>{inline}</h2>")

                    elif sn.startswith("Heading 3"):
                        html_parts.append(f"<h3>{inline}</h3>")

                    elif sn.startswith("Heading 4"):
                        html_parts.append(f"<h4>{inline}</h4>")

                    else:
                        html_parts.append(f"<p>{inline}</p>")

                _close_list()   # close any trailing list
                rendered_html = "\n".join(html_parts)
            except Exception as exc:
                rendered_html = f"<p>(Preview error: {exc})</p>"
        else:
            rendered_html = "<p>(No output document yet — click Generate first.)</p>"

        # Populate source panes
        if mode in ("md_source", "split"):
            for tw in (self._p_source_text, self._p_split_source_text):
                tw.config(state="normal"); tw.delete("1.0", "end"); tw.insert("1.0", md_text)
                tw.config(state="disabled")

        # Populate rendered panes
        if mode in ("rendered", "split"):
            self._render_html_to_text(rendered_html, _target=self._p_rendered_text)
            if mode == "split":
                self._render_html_to_text(rendered_html, _target=self._p_split_rendered_text)

        fname = Path(docx_path).name if docx_path and Path(docx_path).exists() else "—"
        self._p_preview_status.config(text=f"Preview: {fname}", fg=FG2)

    # ════════════════════════════════════════════════════════
    #  CONVERT TAB  (Word → Brightspace, unchanged)
    # ════════════════════════════════════════════════════════

    def _build_convert_tab(self, p):
        pill = tk.Frame(p, bg=BG3, bd=0, highlightthickness=1, highlightbackground=BORDER)
        pill.pack(fill="x", pady=(0,14))

        def _pill_btn(text, value, side):
            rb = tk.Radiobutton(pill, text=text, variable=self._mode, value=value,
                                bg=BG3, fg=FG2, font=FONT_B,
                                activebackground=ACCENT, activeforeground="#ffffff",
                                selectcolor=ACCENT, indicatoron=False,
                                relief="flat", bd=0, padx=0, pady=7, cursor="hand2",
                                command=self._on_mode_change)
            rb.pack(side=side, fill="x", expand=True)
            def _sync(*_): rb.config(fg="#ffffff" if self._mode.get()==value else FG2)
            self._mode.trace_add("write", lambda *_: _sync())
            return rb

        btn_single = _pill_btn("  Single File  ","single","left")
        tk.Frame(pill, bg=BORDER, width=1).pack(side="left", fill="y")
        _pill_btn("  Batch  ","batch","left")
        btn_single.config(fg="#ffffff")

        self._single_frame = tk.Frame(p, bg=BG2)
        self._batch_frame  = tk.Frame(p, bg=BG2)
        self._build_single_frame(self._single_frame)
        self._build_batch_frame(self._batch_frame)
        self._single_frame.pack(fill="both", expand=True)

    def _on_mode_change(self):
        if self._mode.get() == "single":
            self._batch_frame.pack_forget()
            self._single_frame.pack(fill="both", expand=True)
        else:
            self._single_frame.pack_forget()
            self._batch_frame.pack(fill="both", expand=True)
            self._batch_scan_dir()

    def _build_single_frame(self, p):
        tk.Label(p, text="INPUT FILE", font=("Segoe UI",8,"bold"),
                 bg=BG2, fg=FG3, pady=0).pack(anchor="w", pady=(0,4))

        self.drop_zone = tk.Frame(p, bg=DROP_BG, cursor="hand2",
                                  highlightthickness=2, highlightbackground=BORDER,
                                  highlightcolor=DROP_HL)
        self.drop_zone.pack(fill="x", pady=(0,16))
        self.file_icon  = tk.Label(self.drop_zone, text="\U0001f4c2", bg=DROP_BG, font=("Segoe UI",22), pady=6)
        self.file_icon.pack()
        self.file_label = tk.Label(self.drop_zone, text="Click or drag & drop a .docx file",
                                   bg=DROP_BG, fg=FG2, font=FONT, pady=4)
        self.file_label.pack()
        self.file_sub   = tk.Label(self.drop_zone, text="", bg=DROP_BG, fg=FG3, font=FONT_S, pady=2)
        self.file_sub.pack()

        for w in (self.drop_zone, self.file_icon, self.file_label, self.file_sub):
            w.bind("<Button-1>", lambda e: self._pick_file())
            w.bind("<Enter>",    lambda e: self.drop_zone.config(highlightbackground=DROP_HL))
            w.bind("<Leave>",    lambda e: self.drop_zone.config(
                highlightbackground=ACCENT if self.selected_file else BORDER))
        self._setup_drag_drop()

        tk.Label(p, text="OUTPUT OPTIONS", font=("Segoe UI",8,"bold"),
                 bg=BG2, fg=FG3).pack(anchor="w", pady=(0,6))
        opts_container = tk.Frame(p, bg=BG3, highlightthickness=1, highlightbackground=BORDER)
        opts_container.pack(fill="x", pady=(0,4))

        def _dark_cb(text, var, tip_text, parent=opts_container):
            f = tk.Frame(parent, bg=BG3); f.pack(fill="x", padx=10, pady=3)
            cb = tk.Checkbutton(f, text=text, variable=var, bg=BG3, fg=FG,
                                activebackground=BG3, activeforeground=FG,
                                selectcolor=BG2, font=FONT, bd=0, highlightthickness=0, cursor="hand2")
            cb.pack(anchor="w"); tip(cb, tip_text); return cb

        _dark_cb("Save output alongside source .docx", self.save_next_var,
                 "Saves MyDoc.html next to MyDoc.docx.")
        _dark_cb("Split into one file per Module", self.split_modules_var,
                 "When headings like 'Module 1:', 'Module 2:' … are detected,\neach module is saved as its own HTML file.")
        _dark_cb('Strip inline style="" attributes', self.strip_style_var,
                 'Removes style="..." attributes from every HTML element.')
        tk.Frame(opts_container, bg=BORDER, height=1).pack(fill="x", padx=10, pady=(4,0))
        _dark_cb("Wrap in full HTML document", self.full_html_var,
                 "Adds <!DOCTYPE html>, <html>, <head>, <body> tags.")

        self.status_var = tk.StringVar(value="")
        self.status_lbl = tk.Label(p, textvariable=self.status_var,
                                   bg=BG2, fg=SUCCESS, font=FONT, wraplength=320, justify="left")
        self.status_lbl.pack(anchor="w", pady=(8,0))
        self.module_info_var = tk.StringVar(value="")
        self.module_info_lbl = tk.Label(p, textvariable=self.module_info_var,
                                        bg=BG2, fg=ACCENT2, font=FONT_S, wraplength=320, justify="left")
        self.module_info_lbl.pack(anchor="w", pady=(2,0))

        self._log_expanded = tk.BooleanVar(value=False)
        log_header = tk.Frame(p, bg=BG2, cursor="hand2"); log_header.pack(fill="x", pady=(14,0))
        self._log_arrow = tk.Label(log_header, text="▶", font=FONT_S, bg=BG2, fg=FG3)
        self._log_arrow.pack(side="left")
        self._log_header_lbl = tk.Label(log_header, text="CONVERSION LOG",
                                         font=("Segoe UI",8,"bold"), bg=BG2, fg=FG3)
        self._log_header_lbl.pack(side="left", padx=(4,0))
        self._log_badge = tk.Label(log_header, text="", font=FONT_S, bg=BG2, fg=FG3)
        self._log_badge.pack(side="left", padx=(6,0))
        for w in (log_header, self._log_arrow, self._log_header_lbl, self._log_badge):
            w.bind("<Button-1>", lambda e: self._toggle_log())

        self._log_body = tk.Frame(p, bg=BG3)
        self._log_text = tk.Text(self._log_body, font=("Cascadia Code",8), wrap="word",
                                  bg=BG3, fg=FG2, relief="flat", bd=0,
                                  padx=8, pady=6, height=8, state="disabled",
                                  selectbackground=ACCENT, selectforeground=FG, cursor="arrow")
        log_vsb = ttk.Scrollbar(self._log_body, orient="vertical", command=self._log_text.yview)
        self._log_text.configure(yscrollcommand=log_vsb.set)
        log_vsb.pack(side="right", fill="y"); self._log_text.pack(fill="both", expand=True)
        self._log_text.tag_configure("info",  foreground=FG2)
        self._log_text.tag_configure("warn",  foreground=WARN)
        self._log_text.tag_configure("error", foreground=ERR)
        self._log_text.tag_configure("dim",   foreground=FG3)

        tk.Frame(p, bg=BORDER, height=1).pack(fill="x", pady=(10,0))
        btn_convert = tk.Button(p, text="Convert  →", font=FONT_B,
                                bg=ACCENT, fg=FG, relief="flat", bd=0,
                                padx=16, pady=10, cursor="hand2",
                                activebackground=ACCENT2, activeforeground=FG,
                                command=self._run_convert)
        btn_convert.pack(fill="x", pady=(6,0))
        tip(btn_convert, "Convert the selected .docx and save the HTML output.")

    def _build_batch_frame(self, p):
        tk.Label(p, text="FOLDER", font=("Segoe UI",8,"bold"),
                 bg=BG2, fg=FG3).pack(anchor="w", pady=(0,4))
        dir_row = tk.Frame(p, bg=BG2); dir_row.pack(fill="x", pady=(0,8))
        self._batch_dir_var = tk.StringVar(value=str(self._batch_dir))
        dir_entry = tk.Entry(dir_row, textvariable=self._batch_dir_var, font=FONT_S,
                             bg=BG3, fg=FG2, insertbackground=FG, relief="flat", bd=4, state="readonly")
        dir_entry.pack(side="left", fill="x", expand=True, padx=(0,6))
        tk.Button(dir_row, text="Browse…", font=FONT_S, bg=BG3, fg=FG,
                  relief="flat", bd=0, padx=10, pady=5, cursor="hand2",
                  activebackground=BORDER, activeforeground=FG,
                  command=self._batch_browse).pack(side="left")

        sel_row = tk.Frame(p, bg=BG2); sel_row.pack(fill="x", pady=(0,4))
        self._batch_count_lbl = tk.Label(sel_row, text="", font=FONT_S, bg=BG2, fg=FG3)
        self._batch_count_lbl.pack(side="left")
        def _sel_lnk(text, cmd):
            lbl = tk.Label(sel_row, text=text, font=FONT_S, bg=BG2, fg=ACCENT, cursor="hand2")
            lbl.pack(side="right", padx=(6,0)); lbl.bind("<Button-1>", lambda e: cmd()); return lbl
        _sel_lnk("None", self._batch_select_none); _sel_lnk("All", self._batch_select_all)
        tk.Label(sel_row, text="Select:", font=FONT_S, bg=BG2, fg=FG3).pack(side="right", padx=(0,4))

        list_outer = tk.Frame(p, bg=BG3, highlightthickness=1, highlightbackground=BORDER)
        list_outer.pack(fill="both", expand=True, pady=(0,8))
        self._batch_canvas = tk.Canvas(list_outer, bg=BG3, highlightthickness=0)
        list_vsb = ttk.Scrollbar(list_outer, orient="vertical", command=self._batch_canvas.yview)
        self._batch_canvas.configure(yscrollcommand=list_vsb.set)
        list_vsb.pack(side="right", fill="y"); self._batch_canvas.pack(side="left", fill="both", expand=True)
        self._batch_list_frame = tk.Frame(self._batch_canvas, bg=BG3)
        self._batch_list_win = self._batch_canvas.create_window((0,0), window=self._batch_list_frame, anchor="nw")
        def _on_list_resize(e=None):
            self._batch_canvas.configure(scrollregion=self._batch_canvas.bbox("all"))
            self._batch_canvas.itemconfig(self._batch_list_win, width=self._batch_canvas.winfo_width())
        self._batch_list_frame.bind("<Configure>", lambda e: _on_list_resize())
        self._batch_canvas.bind("<Configure>", lambda e: _on_list_resize())
        self._batch_canvas.bind_all("<MouseWheel>",
            lambda e: self._batch_canvas.yview_scroll(int(-1*(e.delta/120)),"units"))

        self._batch_progress_frame = tk.Frame(p, bg=BG2)
        self._batch_progress_var   = tk.DoubleVar(value=0)
        self._batch_progress_lbl_var = tk.StringVar(value="")
        tk.Label(self._batch_progress_frame, textvariable=self._batch_progress_lbl_var,
                 font=FONT_S, bg=BG2, fg=FG2).pack(anchor="w")
        self._batch_bar = ttk.Progressbar(self._batch_progress_frame,
                                           variable=self._batch_progress_var, maximum=100, length=300)
        self._batch_bar.pack(fill="x", pady=(2,0))

        self._batch_btn = tk.Button(p, text="Convert All  →", font=FONT_B,
                                    bg=ACCENT, fg=FG, relief="flat", bd=0,
                                    padx=16, pady=10, cursor="hand2",
                                    activebackground=ACCENT2, activeforeground=FG,
                                    command=self._run_batch)
        self._batch_btn.pack(fill="x", pady=(4,0))
        tip(self._batch_btn, "Convert all checked .docx files using the current settings.")

    def _batch_browse(self):
        d = filedialog.askdirectory(title="Select folder containing .docx files",
                                    initialdir=str(self._batch_dir))
        if d:
            self._batch_dir = Path(d); self._batch_dir_var.set(str(self._batch_dir))
            self._batch_scan_dir()

    def _batch_scan_dir(self):
        for w in self._batch_list_frame.winfo_children(): w.destroy()
        self._batch_files = []
        files = sorted(self._batch_dir.glob("*.docx"), key=lambda p: p.name.lower())
        if not files:
            tk.Label(self._batch_list_frame, text="No .docx files found in this folder.",
                     font=FONT_S, bg=BG3, fg=FG3, padx=8, pady=10).pack(anchor="w")
            self._batch_count_lbl.config(text="0 files"); return
        for i, path in enumerate(files):
            row_bg = BG3 if i%2==0 else "#262b40"
            row = tk.Frame(self._batch_list_frame, bg=row_bg); row.pack(fill="x")
            var = tk.BooleanVar(value=True); status_var = tk.StringVar(value="")
            cb = tk.Checkbutton(row, variable=var, bg=row_bg, activebackground=row_bg,
                                selectcolor=BG2, bd=0, highlightthickness=0, cursor="hand2")
            cb.pack(side="left", padx=(6,2))
            name_lbl = tk.Label(row, text=path.name, font=FONT_S, bg=row_bg, fg=FG,
                                anchor="w", cursor="hand2")
            name_lbl.pack(side="left", fill="x", expand=True, pady=5)
            status_lbl = tk.Label(row, textvariable=status_var, font=FONT_S, bg=row_bg,
                                  fg=FG3, width=6, anchor="e")
            status_lbl.pack(side="right", padx=(0,8))
            entry = {"path":path,"var":var,"status_var":status_var,"row":row,
                     "name_lbl":name_lbl,"row_bg":row_bg}
            self._batch_files.append(entry)
            name_lbl.bind("<Button-1>", lambda e, p=path: self._batch_preview_file(p))
        self._batch_count_lbl.config(text=f"{len(files)} file{'s' if len(files)!=1 else ''} found")

    def _batch_select_all(self):
        for f in self._batch_files: f["var"].set(True)

    def _batch_select_none(self):
        for f in self._batch_files: f["var"].set(False)

    def _batch_preview_file(self, path):
        self._load_file(str(path))

    def _run_batch(self):
        if self._batch_running: return
        checked = [f for f in self._batch_files if f["var"].get()]
        if not checked:
            messagebox.showinfo("Nothing selected","Check at least one file to convert."); return
        for f in self._batch_files:
            f["status_var"].set(""); f["name_lbl"].config(fg=FG)
        self._batch_running = True
        self._batch_btn.config(state="disabled", text="Converting…")
        self._batch_progress_frame.pack(fill="x", pady=(0,6), before=self._batch_btn)
        self._batch_progress_var.set(0); total = len(checked)

        def _convert_one(idx):
            if idx >= total:
                self._batch_running = False
                self._batch_btn.config(state="normal", text="Convert All  →")
                self._batch_progress_lbl_var.set(f"Done — {total} file{'s' if total!=1 else ''} converted")
                return
            entry = checked[idx]
            entry["status_var"].set("⏳"); entry["name_lbl"].config(fg=FG)
            self._batch_progress_lbl_var.set(f"Converting {idx+1} / {total}: {entry['path'].name}")
            self._batch_progress_var.set((idx/total)*100); self.update_idletasks()
            try:
                settings = self._collect_settings()
                body_html, img_map, lg, _bl = convert_docx(str(entry["path"]), settings)
                css = self.custom_css or DEFAULT_PREVIEW_CSS; src_path = entry["path"]
                def _write_images(out_dir):
                    if not img_map: return
                    img_dir = out_dir/"images"; img_dir.mkdir(exist_ok=True)
                    for fname, blob in img_map.items(): (img_dir/fname).write_bytes(blob)
                if self.split_modules_var.get(): modules = split_modules(body_html)
                else: modules = []
                if modules:
                    for mod in modules:
                        content = mod["html"]
                        if self.full_html_var.get(): content = wrap_html(content,title=mod["title"],css=css)
                        fname = src_path.parent / f"{src_path.stem}_module_{mod['number']:02d}.html"
                        fname.write_text(content, encoding="utf-8")
                    _write_images(src_path.parent); entry["status_var"].set(f"✅ ×{len(modules)}")
                else:
                    output = (wrap_html(body_html,title=src_path.stem,css=css)
                              if self.full_html_var.get() else body_html)
                    out_path = src_path.with_suffix(".html")
                    out_path.write_text(output, encoding="utf-8")
                    _write_images(src_path.parent); entry["status_var"].set("✅")
                entry["name_lbl"].config(fg=SUCCESS)
            except Exception as exc:
                entry["status_var"].set("❌"); entry["name_lbl"].config(fg=ERR)
                tip(entry["name_lbl"], f"Error: {exc}")
            self._batch_progress_var.set(((idx+1)/total)*100)
            self.after(10, lambda: _convert_one(idx+1))

        self.after(10, lambda: _convert_one(0))

    def _setup_drag_drop(self):
        try:
            from tkinterdnd2 import DND_FILES
            drop_widgets = (self.drop_zone, self.file_icon, self.file_label, self.file_sub)
            for w in drop_widgets:
                w.drop_target_register(DND_FILES); w.dnd_bind("<<Drop>>", self._on_drop)
            self.drop_target_register(DND_FILES)
            self.dnd_bind("<<DragEnter>>", self._on_drag_enter)
            self.dnd_bind("<<DragLeave>>", self._on_drag_leave)
            self.dnd_bind("<<Drop>>", self._on_drop)
            for w in drop_widgets:
                w.dnd_bind("<<DragEnter>>", self._on_drag_enter)
                w.dnd_bind("<<DragLeave>>", self._on_drag_leave)
        except Exception: pass

    def _on_drag_enter(self, event=None):
        if hasattr(self,"_drag_leave_id") and self._drag_leave_id:
            self.after_cancel(self._drag_leave_id); self._drag_leave_id = None
        self.drop_zone.config(highlightbackground=DROP_HL, bg="#2e3350")
        for w in (self.file_icon, self.file_label, self.file_sub): w.config(bg="#2e3350")

    def _on_drag_leave(self, event=None):
        if hasattr(self,"_drag_leave_id") and self._drag_leave_id:
            self.after_cancel(self._drag_leave_id)
        self._drag_leave_id = self.after(50, self._do_drag_leave)

    def _do_drag_leave(self):
        self._drag_leave_id = None
        self.drop_zone.config(highlightbackground=BORDER if not self.selected_file else ACCENT, bg=DROP_BG)
        for w in (self.file_icon, self.file_label, self.file_sub): w.config(bg=DROP_BG)

    def _on_drop(self, event):
        raw = event.data.strip()
        path = raw[1:-1] if (raw.startswith("{") and raw.endswith("}")) else raw.split()[0]
        if hasattr(self,"_drag_leave_id") and self._drag_leave_id:
            self.after_cancel(self._drag_leave_id); self._drag_leave_id = None
        self._do_drag_leave(); self._load_file(path)

    def _pick_file(self):
        path = filedialog.askopenfilename(
            title="Select Word Document",
            initialdir=self._last_file_dir or Path.home(),
            filetypes=[("Word Documents","*.docx"),("All files","*.*")])
        if path: self._load_file(path)

    def _load_file(self, path):
        if not path.lower().endswith(".docx"):
            messagebox.showwarning("Wrong file type","Please select a .docx Word document."); return
        self.selected_file = path; self._last_file_dir = str(Path(path).parent); self._save_config()
        if hasattr(self,"_d2l_title_var"): self._d2l_title_var.set(Path(path).stem)
        name = Path(path).name
        self.file_icon.config(text="\U0001f4c4"); self.file_label.config(text=name, fg=FG)
        self.file_sub.config(text=str(Path(path).parent), fg=FG3)
        self.drop_zone.config(highlightbackground=ACCENT); self.status_var.set("")
        self._refresh_preview()

    # ── Settings tab ──────────────────────────────────────────

    def _build_settings_tab(self, parent):
        canvas = tk.Canvas(parent, bg=BG2, highlightthickness=0)
        vsb    = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y"); canvas.pack(side="left", fill="both", expand=True)
        inner  = tk.Frame(canvas, bg=BG2, padx=PAD, pady=PAD)
        win_id = canvas.create_window((0,0), window=inner, anchor="nw")
        def _resize(e=None):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(win_id, width=canvas.winfo_width())
        inner.bind("<Configure>",  lambda e: _resize())
        canvas.bind("<Configure>", lambda e: _resize())
        canvas.bind_all("<MouseWheel>",
            lambda e: canvas.yview_scroll(int(-1*(e.delta/120)),"units"))

        # ── Brightspace section header ────────────────────────
        bs_hdr = tk.Frame(inner, bg=BG3, highlightthickness=1, highlightbackground=BORDER)
        bs_hdr.pack(fill="x", pady=(0, 10))
        tk.Label(bs_hdr, text="  Word → Brightspace", font=("Segoe UI", 11, "bold"),
                 bg=BG3, fg=ACCENT2, padx=4, pady=8).pack(side="left")

        self._sep(inner,"Presets","Save current settings as a named preset.")
        pf = tk.Frame(inner, bg=BG2); pf.pack(fill="x", pady=(0,8))
        self.preset_combo = ttk.Combobox(pf, textvariable=self.preset_var, values=[], state="readonly", width=22)
        self.preset_combo.grid(row=0, column=0, sticky="w", padx=(0,6))
        self.preset_combo.bind("<<ComboboxSelected>>", lambda e: self._load_preset())
        def _icon_btn(parent, text, cmd, tip_text, col):
            b = tk.Button(parent, text=text, font=FONT, bg=BG3, fg=FG, relief="flat", bd=0,
                          padx=9, pady=4, cursor="hand2",
                          activebackground=BORDER, activeforeground=FG, command=cmd)
            b.grid(row=0, column=col, padx=(0,4)); tip(b, tip_text); return b
        _icon_btn(pf,"Save…",   self._save_preset_dialog,  "Save current settings as a new preset.", 1)
        _icon_btn(pf,"Rename…", self._rename_preset_dialog, "Rename the selected preset.", 2)
        _icon_btn(pf,"Delete",  self._delete_preset,        "Delete the selected preset.", 3)
        io_row = tk.Frame(inner, bg=BG2); io_row.pack(fill="x", pady=(4,8))
        def _io_btn(text, cmd, tip_text):
            b = tk.Button(io_row, text=text, font=FONT_S, bg=BG3, fg=FG, relief="flat", bd=0,
                          padx=9, pady=4, cursor="hand2",
                          activebackground=BORDER, activeforeground=FG, command=cmd)
            b.pack(side="left", padx=(0,6)); tip(b, tip_text)
        _io_btn("Export profiles…", self._export_presets,  "Save all profiles to a .json file.")
        _io_btn("Import profiles…", self._import_presets, "Load profiles from a .json file.")

        self._sep(inner,"Heading Transform","Map Word heading levels to HTML tags.")
        hf = tk.Frame(inner, bg=BG2); hf.pack(fill="x", pady=(0,8))
        tk.Label(hf, text="Word Style", font=("Segoe UI",9,"bold"), fg=FG3, bg=BG2).grid(row=0,column=0,sticky="w",padx=(0,8))
        tk.Label(hf, text="Output Tag", font=("Segoe UI",9,"bold"), fg=FG3, bg=BG2).grid(row=0,column=2,sticky="w")
        word_styles = ["Title"]+[f"Heading {n}" for n in range(1,7)]
        defaults    = list(DEFAULT_HEADING_MAP.values())
        for idx,(ws,dfl) in enumerate(zip(word_styles,defaults)):
            tk.Label(hf, text=ws, font=FONT, bg=BG2, fg=FG).grid(row=idx+1,column=0,sticky="w",pady=3,padx=(0,6))
            tk.Label(hf, text="→", font=FONT, bg=BG2, fg=FG3).grid(row=idx+1,column=1,padx=8)
            var = tk.StringVar(value=dfl); self.heading_vars[ws] = var
            cb = ttk.Combobox(hf, textvariable=var, values=HEADING_OPTS, state="readonly", width=10)
            cb.grid(row=idx+1,column=2,sticky="w",pady=3)
            cb.bind("<<ComboboxSelected>>", lambda e: self._refresh_preview())

        self._sep(inner,"List Transform","Choose HTML list element for bullet/numbered lists.")
        lf = tk.Frame(inner, bg=BG2); lf.pack(fill="x", pady=(0,8))
        list_rows = [("Bullet list (UL) →",self.ul_var,"Force bullet lists to ul or ol."),
                     ("Numbered list (OL) →",self.ol_var,"Force numbered lists to ol or ul.")]
        for row_i,(lbl,var,tip_text) in enumerate(list_rows):
            tk.Label(lf, text=lbl, font=FONT, bg=BG2, fg=FG).grid(row=row_i,column=0,sticky="w",pady=3,padx=(0,8))
            cb = ttk.Combobox(lf, textvariable=var, values=LIST_OPTS, state="readonly", width=10)
            cb.grid(row=row_i,column=1,sticky="w")
            cb.bind("<<ComboboxSelected>>", lambda e: self._refresh_preview()); tip(cb, tip_text)

        self._sep(inner,"Blockquote Transform","Control how Word 'Quote' styles are output.")
        bqf = tk.Frame(inner, bg=BG2); bqf.pack(fill="x", pady=(0,8))
        tk.Label(bqf, text="Quote style →", font=FONT, bg=BG2, fg=FG).grid(row=0,column=0,sticky="w",padx=(0,8),pady=2)
        cb_bq = ttk.Combobox(bqf, textvariable=self.bq_var, values=BQ_OPTS, state="readonly", width=22)
        cb_bq.grid(row=0,column=1,sticky="w"); cb_bq.bind("<<ComboboxSelected>>", lambda e: self._refresh_preview())
        cb_hr = tk.Checkbutton(bqf, text="Add <hr> dividers inside blockquote",
                               variable=self.bq_hr_var, bg=BG2, fg=FG,
                               activebackground=BG2, activeforeground=FG,
                               selectcolor=BG3, font=FONT, bd=0, highlightthickness=0,
                               command=self._refresh_preview)
        cb_hr.grid(row=1,column=0,columnspan=2,sticky="w",pady=(4,0))

        self._sep(inner,"Accordion Card Title Style","Single-cell tables → D2L accordions.")
        af = tk.Frame(inner, bg=BG2); af.pack(fill="x", pady=(0,8))
        tk.Label(af, text="Trigger heading →", font=FONT, bg=BG2, fg=FG).grid(row=0,column=0,sticky="w",padx=(0,8))
        cb_acc = ttk.Combobox(af, textvariable=self.acc_head_var,
                              values=[f"Heading {n}" for n in range(1,7)],
                              state="readonly", width=12)
        cb_acc.grid(row=0,column=1,sticky="w")
        cb_acc.bind("<<ComboboxSelected>>", lambda e: self._refresh_preview())

        self._sep(inner,"Preview CSS (optional)","Load a CSS file to style the preview panel.")
        cf = tk.Frame(inner, bg=BG2); cf.pack(fill="x", pady=(0,8))
        self.css_label = tk.Label(cf, text="No CSS loaded — default preview styles",
                                   font=FONT_S, fg=FG3, bg=BG2, wraplength=300, justify="left")
        self.css_label.pack(anchor="w")
        br = tk.Frame(cf, bg=BG2); br.pack(anchor="w", pady=(6,0))
        def _dark_btn(parent, text, cmd, tip_text):
            b = tk.Button(parent, text=text, font=FONT, bg=BG3, fg=FG, relief="flat", bd=0,
                          padx=10, pady=5, cursor="hand2",
                          activebackground=BORDER, activeforeground=FG, command=cmd)
            b.pack(side="left", padx=(0,6)); tip(b, tip_text); return b
        _dark_btn(br,"Upload CSS…",self._pick_css,"Load a .css file for the preview.")
        _dark_btn(br,"Clear",self._clear_css,"Revert to default preview styles.")

        self._sep(inner,"Paragraph Font Size","Font size for body paragraphs in preview.")
        pff = tk.Frame(inner, bg=BG2); pff.pack(fill="x", pady=(0,8))
        tk.Label(pff, text="Body text size →", font=FONT, bg=BG2, fg=FG).grid(row=0,column=0,sticky="w",padx=(0,8))
        pf_spin = tk.Spinbox(pff, from_=10, to=36, increment=1, textvariable=self.para_font_size,
                             width=5, font=FONT, bg=BG3, fg=FG, insertbackground=FG,
                             buttonbackground=BG3, relief="flat", command=self._on_font_size_change)
        pf_spin.grid(row=0,column=1,sticky="w")
        pf_spin.bind("<FocusOut>", lambda e: self._on_font_size_change())
        pf_spin.bind("<Return>",   lambda e: self._on_font_size_change())
        tk.Label(pff, text="px", font=FONT, bg=BG2, fg=FG2).grid(row=0,column=2,sticky="w",padx=(4,0))

        # ── Padlet section header ─────────────────────────────
        tk.Frame(inner, bg=BORDER, height=1).pack(fill="x", pady=(18, 0))
        padlet_hdr = tk.Frame(inner, bg=BG3, highlightthickness=1, highlightbackground=BORDER)
        padlet_hdr.pack(fill="x", pady=(10, 10))
        tk.Label(padlet_hdr, text="  Padlet \u2192 Word", font=("Segoe UI", 11, "bold"),
                 bg=BG3, fg=ACCENT2, padx=4, pady=8).pack(side="left")

        # ── Remove Padlet Instructions ────────────────────────
        self._sep(inner, "Remove Padlet Instructions",
                  "Strip all bracketed content (e.g. [see instructions]) from the "
                  "Markdown file on import, so those internal notes don\u2019t appear "
                  "in the generated Word document.")
        rpi_frame = tk.Frame(inner, bg=BG2); rpi_frame.pack(fill="x", pady=(0, 8))
        rpi_cb = tk.Checkbutton(
            rpi_frame,
            text="Remove bracketed content on import",
            variable=tk.BooleanVar(value=False),
            state="disabled",
            bg=BG2, fg=FG3,
            activebackground=BG2, activeforeground=FG,
            selectcolor=BG3, disabledforeground=FG3,
            font=FONT, bd=0, highlightthickness=0)
        rpi_cb.pack(anchor="w")
        tip(rpi_cb,
            "When enabled, any text enclosed in square brackets — such as "
            "[placeholder], [see earlier note], or [optional] — will be removed "
            "from the Markdown source before it is parsed. Useful for keeping "
            "internal Padlet annotations out of the final Word document.")

        # ── Link to Attached Files ────────────────────────────
        self._sep(inner, "Attached Files",
                  "Control whether files attached to Padlet cards are included "
                  "in the output document.")
        laf_frame = tk.Frame(inner, bg=BG2); laf_frame.pack(fill="x", pady=(0, 8))
        laf_cb = tk.Checkbutton(
            laf_frame,
            text="Link to attached files",
            variable=tk.BooleanVar(value=True),
            state="disabled",
            bg=BG2, fg=FG3,
            activebackground=BG2, activeforeground=FG,
            selectcolor=BG3, disabledforeground=FG3,
            font=FONT, bd=0, highlightthickness=0)
        laf_cb.pack(anchor="w")
        tip(laf_cb,
            "When ON (current behaviour): attachment URLs from each Padlet card "
            "are included as hyperlinks in the Word document.\n\n"
            "When OFF: all attached files are silently omitted from the output, "
            "keeping the document free of external links.")

        # ── Import Course Name ────────────────────────────────
        self._sep(inner, "Import Course Name",
                  "Detect and reformat the course title from the Padlet heading.")
        icn_frame = tk.Frame(inner, bg=BG2); icn_frame.pack(fill="x", pady=(0, 8))
        icn_cb = tk.Checkbutton(
            icn_frame,
            text="Detect and reformat course name",
            variable=tk.BooleanVar(value=False),
            state="disabled",
            bg=BG2, fg=FG3,
            activebackground=BG2, activeforeground=FG,
            selectcolor=BG3, disabledforeground=FG3,
            font=FONT, bd=0, highlightthickness=0)
        icn_cb.pack(anchor="w")
        tip(icn_cb,
            "Reads the Padlet document title (e.g. \u201cThe Learning Environment "
            "CONT 806 course planner\u201d) and converts it into a formatted heading "
            "at the top of the Word document.\n\n"
            "Format: \u201cCONT806: The Learning Environment\u201d \u2014 the course code is "
            "extracted and moved to the front, followed by the descriptive title. "
            "Replaces the current placeholder \u201cCONT###:\u201d heading in the template.")

    # ── Preset management ─────────────────────────────────────

    def _on_font_size_change(self):
        try: sz = max(10, min(36, int(self.para_font_size.get())))
        except (ValueError, tk.TclError): sz = 19
        self.para_font_size.set(sz); pt = max(8, round(sz*0.75))
        for tw in (self.rendered_text, self.split_rendered_text):
            tw.tag_configure("p", font=("Lato",pt))
        self._refresh_preview()

    def _preset_snapshot(self):
        return {"headings":{ws:v.get() for ws,v in self.heading_vars.items()},
                "ul":self.ul_var.get(),"ol":self.ol_var.get(),"bq":self.bq_var.get(),
                "bq_hr":self.bq_hr_var.get(),"acc_heading":self.acc_head_var.get(),
                "css_path":self.custom_css_path or ""}

    def _apply_preset_dict(self, d):
        for ws,var in self.heading_vars.items():
            if ws in d.get("headings",{}): var.set(d["headings"][ws])
        self.ul_var.set(d.get("ul","ul")); self.ol_var.set(d.get("ol","ol"))
        self.bq_var.set(d.get("bq","blockquote")); self.bq_hr_var.set(d.get("bq_hr",False))
        self.acc_head_var.set(d.get("acc_heading",ACCORDION_HEADING))
        css_path = d.get("css_path","")
        if css_path and Path(css_path).exists():
            self.custom_css = Path(css_path).read_text(encoding="utf-8",errors="replace")
            self.custom_css_path = css_path
            self.css_label.config(text=f"✅  {Path(css_path).name}", fg=SUCCESS)
        elif not css_path:
            self.custom_css = ""; self.custom_css_path = None
            self.css_label.config(text="No CSS loaded — default preview styles", fg=FG3)
        self._refresh_preview()

    def _refresh_preset_combo(self, select=None):
        names = sorted(self.presets.keys())
        self.preset_combo.configure(values=names)
        if select and select in names: self.preset_var.set(select)
        elif names: self.preset_var.set(names[0])
        else: self.preset_var.set("")

    def _load_preset(self):
        name = self.preset_var.get()
        if name and name in self.presets:
            self._apply_preset_dict(self.presets[name]); self._save_config()

    def _save_preset_dialog(self):
        dlg = tk.Toplevel(self); dlg.title("Save Preset"); dlg.configure(bg=BG2)
        dlg.resizable(False,False); dlg.grab_set()
        tk.Label(dlg, text="Preset name:", font=FONT, bg=BG2, fg=FG, padx=14, pady=(14,0)).pack(anchor="w")
        existing = sorted(self.presets.keys()); name_var = tk.StringVar(value=self.preset_var.get() or "")
        combo = ttk.Combobox(dlg, textvariable=name_var, values=existing, font=FONT, width=28)
        combo.pack(padx=14, pady=6); combo.focus_set()
        msg = tk.Label(dlg, text="", font=FONT_S, bg=BG2, fg=WARN, padx=14, wraplength=240); msg.pack(anchor="w")
        def _do_save():
            name = name_var.get().strip()
            if not name: msg.config(text="Please enter a name."); return
            overwriting = name in self.presets
            self.presets[name] = self._preset_snapshot(); self._save_presets_file()
            self._refresh_preset_combo(select=name); dlg.destroy()
            self.status_var.set(f"✅  Preset \"{name}\" {'updated' if overwriting else 'saved'}.")
        bf = tk.Frame(dlg, bg=BG2); bf.pack(pady=(4,14), padx=14, anchor="e")
        tk.Button(bf, text="Save", font=FONT, bg=ACCENT, fg="#fff", relief="flat", bd=0,
                  padx=12, pady=5, cursor="hand2",
                  activebackground=ACCENT2, activeforeground="#fff", command=_do_save).pack(side="left",padx=(0,6))
        tk.Button(bf, text="Cancel", font=FONT, bg=BG3, fg=FG, relief="flat", bd=0,
                  padx=12, pady=5, cursor="hand2",
                  activebackground=BORDER, activeforeground=FG, command=dlg.destroy).pack(side="left")
        dlg.bind("<Return>", lambda e: _do_save()); dlg.bind("<Escape>", lambda e: dlg.destroy())

    def _rename_preset_dialog(self):
        name = self.preset_var.get()
        if not name: messagebox.showinfo("No preset selected","Select a preset first."); return
        dlg = tk.Toplevel(self); dlg.title("Rename Preset"); dlg.configure(bg=BG2)
        dlg.resizable(False,False); dlg.grab_set()
        tk.Label(dlg, text=f"Rename \"{name}\" to:", font=FONT, bg=BG2, fg=FG, padx=14, pady=(14,0)).pack(anchor="w")
        new_var = tk.StringVar(value=name)
        entry = tk.Entry(dlg, textvariable=new_var, font=FONT, bg=BG3, fg=FG,
                         insertbackground=FG, relief="flat", bd=4, width=28)
        entry.pack(padx=14,pady=6); entry.focus_set(); entry.select_range(0,"end")
        msg = tk.Label(dlg, text="", font=FONT_S, bg=BG2, fg=WARN, padx=14); msg.pack(anchor="w")
        def _do_rename():
            new_name = new_var.get().strip()
            if not new_name: msg.config(text="Name cannot be empty."); return
            if new_name == name: dlg.destroy(); return
            if new_name in self.presets: msg.config(text="Name already exists."); return
            self.presets[new_name] = self.presets.pop(name); self._save_presets_file()
            self._refresh_preset_combo(select=new_name); dlg.destroy()
            self.status_var.set(f'✅  Renamed to "{new_name}".')
        bf = tk.Frame(dlg, bg=BG2); bf.pack(pady=(4,14), padx=14, anchor="e")
        tk.Button(bf, text="Rename", font=FONT, bg=ACCENT, fg="#fff", relief="flat", bd=0,
                  padx=12, pady=5, cursor="hand2",
                  activebackground=ACCENT2, activeforeground="#fff", command=_do_rename).pack(side="left",padx=(0,6))
        tk.Button(bf, text="Cancel", font=FONT, bg=BG3, fg=FG, relief="flat", bd=0,
                  padx=12, pady=5, cursor="hand2",
                  activebackground=BORDER, activeforeground=FG, command=dlg.destroy).pack(side="left")
        dlg.bind("<Return>", lambda e: _do_rename()); dlg.bind("<Escape>", lambda e: dlg.destroy())

    def _delete_preset(self):
        name = self.preset_var.get()
        if not name: messagebox.showinfo("No preset selected","Select a preset first."); return
        if not messagebox.askyesno("Delete preset",f'Delete preset "{name}"?',icon="warning"): return
        self.presets.pop(name,None); self._save_presets_file(); self._refresh_preset_combo()
        self.status_var.set(f'❌  Preset "{name}" deleted.')

    def _export_presets(self):
        if not self.presets: messagebox.showinfo("No profiles","No profiles to export."); return
        path = filedialog.asksaveasfilename(title="Export profiles", defaultextension=".json",
                                            filetypes=[("JSON files","*.json"),("All files","*.*")],
                                            initialfile="brightspace_profiles.json")
        if not path: return
        try:
            Path(path).write_text(json.dumps(self.presets,indent=2), encoding="utf-8")
            n = len(self.presets)
            self.status_var.set(f"✅  Exported {n} profile{'s' if n!=1 else ''} to {Path(path).name}")
        except Exception as exc: messagebox.showerror("Export failed", str(exc))

    def _import_presets(self):
        path = filedialog.askopenfilename(title="Import profiles",
                                          filetypes=[("JSON files","*.json"),("All files","*.*")])
        if not path: return
        try:
            data = json.loads(Path(path).read_text(encoding="utf-8"))
            if not isinstance(data,dict): raise ValueError("File does not contain a profiles dictionary.")
            added = sum(1 for k,v in data.items()
                        if isinstance(k,str) and isinstance(v,dict)
                        and not self.presets.__setitem__(k,v))
            if added == 0:
                messagebox.showwarning("Nothing imported","No valid profiles found."); return
            self._save_presets_file(); self._refresh_preset_combo()
            self.status_var.set(f"✅  Imported {added} profile{'s' if added!=1 else ''} from {Path(path).name}")
        except Exception as exc: messagebox.showerror("Import failed", str(exc))

    def _save_presets_file(self):
        try: PRESETS_FILE.write_text(json.dumps(self.presets,indent=2), encoding="utf-8")
        except Exception: pass

    def _load_presets_file(self):
        try: self.presets = json.loads(PRESETS_FILE.read_text(encoding="utf-8"))
        except Exception: self.presets = {}
        if not self.presets:
            self.presets["Profile 1"] = self._preset_snapshot(); self._save_presets_file()
        last = getattr(self,"_last_preset_name","")
        if last and last in self.presets: self._refresh_preset_combo(select=last)
        else: self._refresh_preset_combo()

    # ── Conversion log (Word → Brightspace) ──────────────────

    def _toggle_log(self):
        if self._log_expanded.get():
            self._log_body.pack_forget(); self._log_expanded.set(False); self._log_arrow.config(text="▶")
        else:
            self._log_body.pack(fill="x", pady=(2,0)); self._log_expanded.set(True)
            self._log_arrow.config(text="▼"); self._log_text.see("end")

    def _update_log(self, log_entries):
        t = self._log_text; t.config(state="normal"); t.delete("1.0","end")
        n_warn  = sum(1 for e in log_entries if e["level"]=="warn")
        n_error = sum(1 for e in log_entries if e["level"]=="error")
        ICONS = {"info":"·","warn":"⚠","error":"✕"}
        for entry in log_entries:
            lvl = entry["level"]; icon = ICONS.get(lvl,"·")
            t.insert("end",f" {icon} ",lvl); t.insert("end",entry["msg"]+"\n",lvl)
        t.config(state="disabled")
        if n_error:
            self._log_badge.config(text=f"  {n_error} error{'s' if n_error!=1 else ''}",fg=ERR)
        elif n_warn:
            self._log_badge.config(text=f"  {n_warn} warning{'s' if n_warn!=1 else ''}",fg=WARN)
        else:
            self._log_badge.config(text="  OK",fg=SUCCESS)
        if n_warn or n_error:
            if not self._log_expanded.get(): self._toggle_log()
        else:
            if self._log_expanded.get(): self._toggle_log()

    def _log_append(self, level, msg):
        t = self._log_text; t.config(state="normal")
        ICONS = {"info":"·","warn":"⚠","error":"✕"}
        icon = ICONS.get(level,"·")
        t.insert("end",f" {icon} ",level); t.insert("end",msg+"\n",level)
        t.config(state="disabled"); t.see("end")

    def _sep(self, parent, label, tooltip_text=None):
        lbl = tk.Label(parent, text=label, font=("Segoe UI",10,"bold"), bg=BG2, fg=ACCENT2)
        lbl.pack(anchor="w", pady=(14,3))
        if tooltip_text: tip(lbl, tooltip_text)
        tk.Frame(parent, bg=BORDER, height=1).pack(fill="x", pady=(0,8))

    # ── Preview panel ─────────────────────────────────────────

    def _build_preview_panel(self, parent):
        bar = tk.Frame(parent, bg=BG2); bar.pack(fill="x")
        tk.Frame(bar, bg=BORDER, height=1).pack(side="bottom", fill="x")
        tk.Label(bar, text="  Preview", font=FONT_B, bg=BG2, fg=FG, pady=8).pack(side="left")

        def _mode_btn(text, value):
            rb = tk.Radiobutton(bar, text=f"  {text}  ", variable=self.preview_mode, value=value,
                                bg=BG2, fg=FG2, font=FONT,
                                activebackground=BG2, activeforeground=FG,
                                selectcolor=BG3, indicatoron=False,
                                relief="flat", bd=0, padx=8, pady=5, cursor="hand2",
                                command=self._refresh_preview)
            rb.pack(side="right", padx=4, pady=4); return rb

        _mode_btn("Rendered","rendered"); _mode_btn("HTML Source","source"); _mode_btn("Side by Side","split")

        def _toggle_wrap():
            wrap_on = self._source_wrap.get(); wrap_val = "word" if wrap_on else "none"
            self.preview_text.config(wrap=wrap_val)
            if hasattr(self,"split_source_text"): self.split_source_text.config(wrap=wrap_val)
            if wrap_on: self._src_hsb.pack_forget()
            else: self._src_hsb.pack(side="bottom",fill="x"); self._src_hsb.lift(self.preview_text)

        self._wrap_btn = tk.Checkbutton(bar, text=" ⇌ Wrap  ", variable=self._source_wrap,
                                        font=FONT, bg=BG2, fg=FG3,
                                        activebackground=BG2, activeforeground=FG,
                                        selectcolor=BG3, indicatoron=False,
                                        relief="flat", bd=0, padx=8, pady=5, cursor="hand2",
                                        command=_toggle_wrap)
        self._wrap_btn.pack(side="left", padx=(2,0), pady=4)

        self.mod_nav_frame = tk.Frame(bar, bg=BG2)
        self._mod_prev_btn = tk.Button(self.mod_nav_frame, text="◀", font=FONT, bg=BG3, fg=FG,
                                       relief="flat", bd=0, padx=8, pady=4, cursor="hand2",
                                       activebackground=BORDER, activeforeground=FG,
                                       command=self._module_prev)
        self._mod_prev_btn.pack(side="left")
        self._mod_label_var = tk.StringVar(value="")
        tk.Label(self.mod_nav_frame, textvariable=self._mod_label_var, bg=BG2, fg=FG, font=FONT, padx=8).pack(side="left")
        self._mod_next_btn = tk.Button(self.mod_nav_frame, text="▶", font=FONT, bg=BG3, fg=FG,
                                       relief="flat", bd=0, padx=8, pady=4, cursor="hand2",
                                       activebackground=BORDER, activeforeground=FG,
                                       command=self._module_next)
        self._mod_next_btn.pack(side="left")

        self.search_frame = tk.Frame(parent, bg=BG2)
        tk.Frame(self.search_frame, bg=BORDER, height=1).pack(side="bottom", fill="x")
        self._search_var     = tk.StringVar()
        self._search_matches = []
        self._search_idx     = 0
        sf_inner = tk.Frame(self.search_frame, bg=BG2, padx=8, pady=5); sf_inner.pack(fill="x")
        tk.Label(sf_inner, text="Find:", font=FONT, bg=BG2, fg=FG2).pack(side="left")
        self._search_entry = tk.Entry(sf_inner, textvariable=self._search_var, font=FONT,
                                      bg=BG3, fg=FG, insertbackground=FG, relief="flat", bd=4, width=28)
        self._search_entry.pack(side="left", padx=(6,4))
        self._search_entry.bind("<Return>",       lambda e: self._search_next())
        self._search_entry.bind("<Shift-Return>", lambda e: self._search_prev())
        self._search_entry.bind("<Escape>",       lambda e: self._hide_search())
        self._search_var.trace_add("write", lambda *_: self._search_run())
        self._search_count_lbl = tk.Label(sf_inner, text="", font=FONT_S, bg=BG2, fg=FG3, width=12, anchor="w")
        self._search_count_lbl.pack(side="left", padx=(2,8))
        tk.Button(sf_inner, text="▲", font=FONT_S, bg=BG3, fg=FG, relief="flat", bd=0,
                  padx=6, pady=2, cursor="hand2", activebackground=BORDER, activeforeground=FG,
                  command=self._search_prev).pack(side="left", padx=(0,2))
        tk.Button(sf_inner, text="▼", font=FONT_S, bg=BG3, fg=FG, relief="flat", bd=0,
                  padx=6, pady=2, cursor="hand2", activebackground=BORDER, activeforeground=FG,
                  command=self._search_next).pack(side="left", padx=(0,8))
        tk.Button(sf_inner, text="✕", font=FONT_S, bg=BG2, fg=FG3, relief="flat", bd=0,
                  padx=6, pady=2, cursor="hand2", activebackground=BG2, activeforeground=FG,
                  command=self._hide_search).pack(side="left")
        self._search_visible = False

        self.preview_outer = tk.Frame(parent, bg=PRE_BG); self.preview_outer.pack(fill="both", expand=True)

        D2L_NAV = "#1a2632"; D2L_NAV_FG = "#ffffff"; D2L_SEP = "#3a4a58"; D2L_SHELL = "#f5f5f5"
        bs_chrome = tk.Frame(self.preview_outer, bg=D2L_NAV); bs_chrome.pack(fill="x")
        bc_row = tk.Frame(bs_chrome, bg=D2L_NAV, padx=14, pady=8); bc_row.pack(fill="x")
        def _crumb(text, fg="#78909c", bold=False):
            font = ("Segoe UI",9,"bold") if bold else ("Segoe UI",9)
            tk.Label(bc_row, text=text, font=font, bg=D2L_NAV, fg=fg, pady=0).pack(side="left")
        _crumb("🏠  Course Home"); _crumb("  ›  ",fg="#546e7a"); _crumb("Content"); _crumb("  ›  ",fg="#546e7a")
        self._d2l_title_var = tk.StringVar(value="(no file loaded)")
        tk.Label(bc_row, textvariable=self._d2l_title_var, font=("Segoe UI",9,"bold"),
                 bg=D2L_NAV, fg=D2L_NAV_FG).pack(side="left")
        tk.Frame(self.preview_outer, bg=D2L_SEP, height=1).pack(fill="x")

        self._single_container = tk.Frame(self.preview_outer, bg=PRE_BG)
        self._single_container.pack(fill="both", expand=True)

        self.source_frame = tk.Frame(self._single_container, bg=PRE_BG)
        self.preview_text = tk.Text(self.source_frame, font=("Cascadia Code",9), wrap="none",
                                    bg=PRE_BG, fg=PRE_FG, insertbackground=FG, relief="flat", bd=0,
                                    padx=12, pady=12, selectbackground=ACCENT, selectforeground=FG)
        vsb = ttk.Scrollbar(self.source_frame, orient="vertical", command=self.preview_text.yview)
        self._src_hsb = ttk.Scrollbar(self.source_frame, orient="horizontal", command=self.preview_text.xview)
        self.preview_text.configure(yscrollcommand=vsb.set, xscrollcommand=self._src_hsb.set)
        vsb.pack(side="right", fill="y"); self._src_hsb.pack(side="bottom", fill="x")
        self.preview_text.pack(fill="both", expand=True)

        self.rendered_frame = tk.Frame(self._single_container, bg=D2L_SHELL)
        card_frame = tk.Frame(self.rendered_frame, bg="#ffffff",
                              highlightbackground="#d0d0d0", highlightthickness=1)
        card_frame.pack(fill="both", expand=True)
        self.rendered_text = tk.Text(card_frame, font=("Lato",11), wrap="word",
                                     bg="#ffffff", fg="#212121", relief="flat", bd=0,
                                     padx=32, pady=28, spacing1=0, spacing2=0, spacing3=0,
                                     selectbackground="#b3d4f5", selectforeground="#212121",
                                     cursor="xterm", state="normal", insertwidth=0)
        rvsb = ttk.Scrollbar(card_frame, orient="vertical", command=self.rendered_text.yview)
        self.rendered_text.configure(yscrollcommand=rvsb.set)
        rvsb.pack(side="right", fill="y"); self.rendered_text.pack(fill="both", expand=True)
        self._configure_rendered_tags()
        self.rendered_text.bind("<Key>", lambda e: "break")
        self.rendered_text.bind("<Control-a>",
            lambda e: (self.rendered_text.tag_add("sel","1.0","end"),"break"))
        self.rendered_text.bind("<Control-c>", lambda e: None)
        self._acc_toggles = {}
        def _rt_click(e, w=self.rendered_text):
            idx = w.index(f"@{e.x},{e.y}")
            for tag in w.tag_names(idx):
                if "_title" in tag or "_arrow" in tag:
                    click_tag = tag.rsplit("_",1)[0]+"_click"
                    cb = self._acc_toggles.get(click_tag)
                    if cb: cb(e); return "break"
        self.rendered_text.bind("<Button-1>", _rt_click)

        self._split_pane = tk.PanedWindow(self.preview_outer, orient="horizontal",
                                           bg=BORDER, sashwidth=5, sashrelief="flat", sashpad=0)
        _sp_left = tk.Frame(self._split_pane, bg=PRE_BG)
        self._split_pane.add(_sp_left, minsize=200, stretch="always")
        self.split_source_text = tk.Text(_sp_left, font=("Cascadia Code",9), wrap="none",
                                          bg=PRE_BG, fg=PRE_FG, insertbackground=FG, relief="flat", bd=0,
                                          padx=12, pady=12, selectbackground=ACCENT, selectforeground=FG,
                                          state="disabled")
        _sp_vsb = ttk.Scrollbar(_sp_left, orient="vertical", command=self.split_source_text.yview)
        _sp_hsb = ttk.Scrollbar(_sp_left, orient="horizontal", command=self.split_source_text.xview)
        self.split_source_text.configure(yscrollcommand=_sp_vsb.set, xscrollcommand=_sp_hsb.set)
        _sp_vsb.pack(side="right",fill="y"); _sp_hsb.pack(side="bottom",fill="x")
        self.split_source_text.pack(fill="both", expand=True)

        _sp_right = tk.Frame(self._split_pane, bg=D2L_SHELL)
        self._split_pane.add(_sp_right, minsize=200, stretch="always")
        _sp_card = tk.Frame(_sp_right, bg="#ffffff", highlightbackground="#d0d0d0", highlightthickness=1)
        _sp_card.pack(fill="both", expand=True)
        self.split_rendered_text = tk.Text(_sp_card, font=("Lato",11), wrap="word",
                                            bg="#ffffff", fg="#212121", relief="flat", bd=0,
                                            padx=32, pady=28, spacing1=0, spacing2=0, spacing3=0,
                                            selectbackground="#b3d4f5", selectforeground="#212121",
                                            cursor="xterm", state="normal", insertwidth=0)
        _sp_rvsb = ttk.Scrollbar(_sp_card, orient="vertical", command=self.split_rendered_text.yview)
        self.split_rendered_text.configure(yscrollcommand=_sp_rvsb.set)
        _sp_rvsb.pack(side="right", fill="y"); self.split_rendered_text.pack(fill="both", expand=True)
        self._configure_rendered_tags(target=self.split_rendered_text)
        self.split_rendered_text.bind("<Key>", lambda e: "break")
        self.split_rendered_text.bind("<Control-a>",
            lambda e: (self.split_rendered_text.tag_add("sel","1.0","end"),"break"))
        self.split_rendered_text.bind("<Control-c>", lambda e: None)

        bot = tk.Frame(parent, bg=BG2); bot.pack(fill="x")
        tk.Frame(bot, bg=BORDER, height=1).pack(side="top", fill="x")
        btn_browser = tk.Button(bot, text="Open in Browser  🌐", font=FONT, bg=BG2, fg=FG2,
                                relief="flat", bd=0, padx=10, pady=6, cursor="hand2",
                                activebackground=BG3, activeforeground=FG,
                                command=self._open_in_browser)
        btn_browser.pack(side="right", padx=8, pady=4)
        self._zip_btn = tk.Button(bot, text="Save as ZIP  📦", font=FONT, bg=BG2, fg=FG2,
                                  relief="flat", bd=0, padx=10, pady=6, cursor="hand2",
                                  activebackground=BG3, activeforeground=FG, command=self._save_zip)
        self._zip_btn.pack(side="right", padx=8, pady=4)
        self._copy_btn = tk.Button(bot, text="Copy HTML  📋", font=FONT, bg=BG2, fg=FG2,
                                   relief="flat", bd=0, padx=10, pady=6, cursor="hand2",
                                   activebackground=BG3, activeforeground=FG, command=self._copy_html)
        self._copy_btn.pack(side="left", padx=8, pady=4)

        self._show_pane("source")
        self._set_preview("Select a .docx file to see a preview here.")

        self.bind_all("<Control-f>",       lambda e: self._toggle_search())
        self.bind_all("<Control-F>",       lambda e: self._toggle_search())
        self.bind_all("<Control-Shift-c>", lambda e: self._copy_html())
        self.bind_all("<Control-Shift-C>", lambda e: self._copy_html())

    def _configure_rendered_tags(self, target=None):
        t = target if target is not None else self.rendered_text
        t.tag_configure("h1", font=("Lato",24,"bold"), foreground="#212121", spacing1=8, spacing3=3)
        t.tag_configure("h2", font=("Lato",19,"bold"), foreground="#212121", spacing1=7, spacing3=3)
        t.tag_configure("h3", font=("Lato",16,"bold"), foreground="#212121", spacing1=5, spacing3=2)
        t.tag_configure("h4", font=("Lato",14,"bold"), foreground="#212121", spacing1=4, spacing3=2)
        t.tag_configure("h5", font=("Lato",12,"bold"), foreground="#424242", spacing1=3, spacing3=1)
        t.tag_configure("h6", font=("Lato",11,"bold"), foreground="#616161", spacing1=3, spacing3=1)
        try: _px = int(self.para_font_size.get())
        except Exception: _px = 19
        _pt = max(8, round(_px*0.75))
        t.tag_configure("p",  font=("Lato",_pt), foreground="#212121", spacing1=0, spacing3=4)
        t.tag_configure("li", font=("Lato",_pt),  foreground="#212121", spacing1=2, spacing3=2, lmargin1=12, lmargin2=28)
        t.tag_configure("blockquote", font=("Lato",_pt,"italic"), foreground="#424242",
                        lmargin1=28, lmargin2=28, spacing1=2, spacing3=2)
        t.tag_configure("bold",          font=("Lato",_pt,"bold"))
        t.tag_configure("italic",        font=("Lato",_pt,"italic"))
        t.tag_configure("bold_italic",   font=("Lato",_pt,"bold italic"))
        t.tag_configure("underline",     font=("Lato",_pt,"underline"))
        t.tag_configure("strikethrough", font=("Lato",_pt), overstrike=True, foreground="#757575")
        t.tag_configure("link", foreground="#006fbf", font=("Lato",_pt,"underline"))
        t.tag_configure("hr",  foreground="#e0e0e0", spacing1=4, spacing3=4)
        t.tag_configure("img", foreground="#006fbf", font=("Lato",10,"italic"), spacing1=4, spacing3=4, lmargin1=4)
        t.tag_configure("acc_arrow",   font=("Lato",11), foreground="#aaaaaa", background="#ffffff")
        t.tag_configure("acc_body",    font=("Lato",11), foreground="#333333", background="#f5f5f5",
                        spacing1=4, spacing3=4, lmargin1=16, lmargin2=16)
        t.tag_configure("acc_divider", font=("Lato",1),  background="#e0e0e0", foreground="#e0e0e0",
                        spacing1=0, spacing3=0)
        t.tag_configure("tbl_rule",   foreground="#bdbdbd", font=("Cascadia Code",10), spacing1=0, spacing3=0)
        t.tag_configure("tbl_th",     foreground="#212121", font=("Cascadia Code",10,"bold"),
                        background="#f5f5f5", spacing1=1, spacing3=1)
        t.tag_configure("tbl_td",     foreground="#212121", font=("Cascadia Code",10), spacing1=1, spacing3=1)
        t.tag_configure("tbl_border", foreground="#bdbdbd", font=("Cascadia Code",10))

    def _toggle_search(self):
        if self._search_visible: self._hide_search()
        else: self._show_search()

    def _show_search(self):
        if self._search_visible: return
        self.search_frame.pack(fill="x", before=self.preview_outer)
        self._search_visible = True; self._search_entry.focus_set()
        self._search_entry.select_range(0,"end"); self._search_run()

    def _hide_search(self):
        if not self._search_visible: return
        self.search_frame.pack_forget(); self._search_visible = False
        self._search_clear_highlights(); self._search_matches = []; self._search_count_lbl.config(text="")
        if self.preview_mode.get() in ("source","split"): self.preview_text.focus_set()
        else: self.rendered_text.focus_set()

    def _active_text_widget(self):
        if self.preview_mode.get() in ("source","split"): return self.preview_text
        return self.rendered_text

    def _search_clear_highlights(self):
        for tw in (self.preview_text, self.rendered_text, self.split_source_text, self.split_rendered_text):
            tw.tag_remove("search_match","1.0","end"); tw.tag_remove("search_current","1.0","end")

    def _search_run(self):
        self._search_clear_highlights(); self._search_matches = []; self._search_idx = 0
        query = self._search_var.get()
        if not query: self._search_count_lbl.config(text=""); return
        tw = self._active_text_widget()
        tw.tag_configure("search_match",   background="#fbbf24", foreground="#1e293b")
        tw.tag_configure("search_current", background="#f97316", foreground="#fff")
        start = "1.0"
        while True:
            pos = tw.search(query, start, stopindex="end", nocase=True)
            if not pos: break
            end = f"{pos}+{len(query)}c"
            tw.tag_add("search_match", pos, end); self._search_matches.append((pos,end)); start = end
        count = len(self._search_matches)
        if count == 0: self._search_count_lbl.config(text="No results", fg=ERR)
        else: self._search_idx = 0; self._search_highlight_current()

    def _search_highlight_current(self):
        if not self._search_matches: return
        tw = self._active_text_widget(); tw.tag_remove("search_current","1.0","end")
        pos,end = self._search_matches[self._search_idx]
        tw.tag_add("search_current", pos, end); tw.see(pos)
        total = len(self._search_matches)
        self._search_count_lbl.config(text=f"{self._search_idx+1} / {total}", fg=FG2)

    def _search_next(self):
        if not self._search_matches: return
        self._search_idx = (self._search_idx+1)%len(self._search_matches)
        self._search_highlight_current()

    def _search_prev(self):
        if not self._search_matches: return
        self._search_idx = (self._search_idx-1)%len(self._search_matches)
        self._search_highlight_current()

    def _render_html_to_text(self, body_html, _target=None):
        import html as html_lib
        import html.parser

        t = _target if _target is not None else self.rendered_text
        t.config(state="normal"); t.delete("1.0","end")
        is_primary = (t is self.rendered_text)
        if is_primary: self._acc_toggles.clear()

        class _TableExtractor(html.parser.HTMLParser):
            def __init__(self):
                super().__init__()
                self.tables=[]; self._cur_tbl=None; self._cur_row=None
                self._cell_buf=[]; self._cell_is_th=False; self._cell_depth=0
            def handle_starttag(self,tag,attrs):
                if tag=="table": self._cur_tbl=[]
                elif tag=="tr" and self._cur_tbl is not None: self._cur_row=[]
                elif tag in ("th","td") and self._cur_row is not None:
                    self._cell_buf=[]; self._cell_is_th=(tag=="th"); self._cell_depth=1
                elif self._cell_depth>0: self._cell_depth+=1
            def handle_endtag(self,tag):
                if tag=="table":
                    if self._cur_tbl is not None: self.tables.append(self._cur_tbl)
                    self._cur_tbl=None
                elif tag=="tr" and self._cur_tbl is not None:
                    if self._cur_row is not None: self._cur_tbl.append(self._cur_row)
                    self._cur_row=None
                elif tag in ("th","td") and self._cell_depth==1:
                    text=html_lib.unescape("".join(self._cell_buf)).strip()
                    if self._cur_row is not None: self._cur_row.append((text,self._cell_is_th))
                    self._cell_depth=0
                elif self._cell_depth>1: self._cell_depth-=1
            def handle_data(self,data):
                if self._cell_depth>=1: self._cell_buf.append(data)

        ex=_TableExtractor(); ex.feed(body_html)
        table_data=ex.tables; table_idx=[0]

        def _col_widths(rows):
            if not rows: return []
            n=max(len(r) for r in rows); w=[0]*n
            for row in rows:
                for ci,(txt,_) in enumerate(row):
                    if ci<n: w[ci]=max(w[ci],len(txt))
            return w

        def _render_table(rows):
            if not rows: return
            widths=_col_widths(rows)
            if not widths: return
            def _rule(l,m,r,fill="─"): return l+m.join(fill*(w+2) for w in widths)+r
            t.insert("end","\n"); t.insert("end",_rule("┌","┬","┐"),("tbl_rule",))
            for row in rows:
                is_hdr=any(is_th for _,is_th in row); cell_tag="tbl_th" if is_hdr else "tbl_td"
                t.insert("end","\n"); t.insert("end","│",("tbl_border",))
                for ci,w in enumerate(widths):
                    txt,_=row[ci] if ci<len(row) else ("",False)
                    t.insert("end",f" {txt:<{w}} ",(cell_tag,)); t.insert("end","│",("tbl_border",))
                if is_hdr: t.insert("end","\n"); t.insert("end",_rule("├","┼","┤"),("tbl_rule",))
            t.insert("end","\n"); t.insert("end",_rule("└","┴","┘"),("tbl_rule",)); t.insert("end","\n")

        link_counter=[0]; card_counter=[0]
        acc_toggles = self._acc_toggles if is_primary else {}
        try: _render_pt = max(8, round(int(self.para_font_size.get()) * 0.75))
        except Exception: _render_pt = 11

        class _Parser(html.parser.HTMLParser):
            def __init__(self,widget):
                super().__init__()
                self.w=widget; self.tag_stack=[]; self.list_stack=[]; self.skip=False
                self._pending_nl=0; self._link_tag=None; self._in_table=False
                self._in_accordion=False; self._in_card_hdr=False; self._in_card_body=False
                self._cur_elide_tag=None; self._hdr_title=""; self._div_depth=0
                self._acc_depth=0; self._hdr_depth=0; self._body_depth=0

            def _flush_nl(self,elide_tag=None):
                if self._pending_nl:
                    tags=(elide_tag,) if elide_tag else ()
                    self.w.insert("end","\n"*self._pending_nl,tags); self._pending_nl=0

            def _nl(self,n=1): self._pending_nl=max(self._pending_nl,n)

            def _insert(self,text,*tags):
                if self.skip or self._in_table or not text: return
                elide=self._cur_elide_tag if self._in_card_body else None
                self._flush_nl(elide_tag=elide); self.w.insert("end",text,tags)

            def _inline_tags(self):
                cur=self.tag_stack
                bold="strong" in cur or "b" in cur; italic="em" in cur or "i" in cur
                underline="u" in cur; strike="del" in cur or "s" in cur
                in_link="a" in cur; in_li="li" in cur; in_bq="blockquote" in cur
                in_hdr=any(h in cur for h in ("h1","h2","h3","h4","h5","h6"))
                tags=[]
                if self._in_card_body and self._cur_elide_tag: tags.append(self._cur_elide_tag)
                if in_hdr: tags.append(next(h for h in cur if h in ("h1","h2","h3","h4","h5","h6")))
                elif in_li:
                    depth=max(sum(1 for x in cur if x in ("ul","ol")),1); tags.append(f"li_d{depth}")
                elif in_bq: tags.append("blockquote")
                elif self._in_card_body: tags.append("acc_body")
                else: tags.append("p")
                if bold and italic: tags.append("bold_italic")
                elif bold: tags.append("bold")
                elif italic: tags.append("italic")
                if underline: tags.append("underline")
                if strike: tags.append("strikethrough")
                if in_link and self._link_tag: tags.append("link"); tags.append(self._link_tag)
                return tuple(tags) if tags else ("p",)

            def handle_starttag(self,tag,attrs):
                self.tag_stack.append(tag); attrs=dict(attrs); cls=attrs.get("class","")
                if tag=="div":
                    self._div_depth+=1
                    if cls=="accordion":
                        self._in_accordion=True; self._acc_depth=self._div_depth; self._flush_nl()
                    elif cls=="card-header" and self._in_accordion:
                        self._in_card_hdr=True; self._hdr_depth=self._div_depth; self._hdr_title=""; self._flush_nl()
                        card_counter[0]+=1; self._cur_elide_tag=f"acc_elide_{card_counter[0]}"
                        self.w.tag_configure(self._cur_elide_tag,elide=True)
                    elif cls=="card-body" and self._in_accordion:
                        self._in_card_body=True; self._body_depth=self._div_depth
                    return
                if tag=="table":
                    self._in_table=True; self._nl(1); self._flush_nl()
                    idx=table_idx[0]
                    if idx<len(table_data): _render_table(table_data[idx]); table_idx[0]+=1
                    return
                if self._in_table: return
                if tag in ("h1","h2","h3","h4","h5","h6"):
                    if not self._in_card_hdr:
                        if self.w.index("end-1c")!="1.0": self._nl(2)
                elif tag=="ul": self.list_stack.append(("ul",[0])); self._nl(1)
                elif tag=="ol": self.list_stack.append(("ol",[0])); self._nl(1)
                elif tag=="li":
                    self._flush_nl(); depth=len(self.list_stack)
                    if self.list_stack:
                        kind,ctr=self.list_stack[-1]; ctr[0]+=1
                        bullet=f"  {ctr[0]}. " if kind=="ol" else "  • "
                        lm1=20+(depth-1)*20; lm2=36+(depth-1)*20; li_tag=f"li_d{depth}"
                        self.w.tag_configure(li_tag,font=("Lato",_render_pt),foreground="#333333",
                                             spacing1=2,spacing3=2,lmargin1=lm1,lmargin2=lm2)
                        bullet_tags=[li_tag]
                        if self._in_card_body and self._cur_elide_tag: bullet_tags=[self._cur_elide_tag,li_tag]
                        self.w.insert("end",bullet,tuple(bullet_tags))
                elif tag=="a":
                    href=attrs.get("href",""); link_counter[0]+=1; ltag=f"link_{link_counter[0]}"
                    self.w.tag_configure(ltag,foreground="#006fbf",font=("Lato",_render_pt,"underline"))
                    self.w.tag_bind(ltag,"<Enter>",lambda e: self.w.config(cursor="hand2"))
                    self.w.tag_bind(ltag,"<Leave>",lambda e: self.w.config(cursor="xterm"))
                    self.w.tag_bind(ltag,"<Button-1>",lambda e,u=href: webbrowser.open(u))
                    self._link_tag=ltag
                elif tag=="blockquote": self._nl(1)
                elif tag=="hr":
                    self._nl(1); self._flush_nl(); self.w.insert("end","─"*40,("hr",)); self._nl(1)
                elif tag=="img":
                    src=attrs.get("src",""); alt=attrs.get("alt","")
                    label=alt or (Path(src).name if src else "image")
                    self._nl(1); self._flush_nl(); self.w.insert("end",f"  🖼️  [{label}]",("img",)); self._nl(1)
                elif tag in ("script","style","head"): self.skip=True

            def handle_endtag(self,tag):
                if tag in ("script","style","head"): self.skip=False
                if self.tag_stack and self.tag_stack[-1]==tag: self.tag_stack.pop()
                if tag=="div":
                    d=self._div_depth; self._div_depth-=1
                    if self._in_card_hdr and d==self._hdr_depth:
                        self._in_card_hdr=False; self._flush_nl()
                        elide_tag=self._cur_elide_tag
                        arrow_tag=f"{elide_tag}_arrow"; title_tag=f"{elide_tag}_title"
                        click_tag=f"{elide_tag}_click"; title=self._hdr_title.strip()
                        self.w.insert("end","\n",("acc_divider",))
                        line_start=self.w.index("end")
                        self.w.tag_configure(arrow_tag,font=("Lato",11),foreground="#aaaaaa",
                                             background="#ffffff",spacing1=8,spacing3=8,lmargin1=10,lmargin2=10)
                        self.w.tag_configure(title_tag,font=("Lato",11,"bold"),foreground="#333333",
                                             background="#ffffff",spacing1=8,spacing3=8)
                        self.w.insert("end","▷  ",(arrow_tag,)); self.w.insert("end",title,(title_tag,))
                        self.w.insert("end","\n",(title_tag,))
                        card_open=[False]
                        self.w.tag_configure(click_tag)
                        self.w.tag_add(click_tag,line_start,"end")
                        self.w.tag_bind(click_tag,"<Enter>",lambda e: self.w.config(cursor="hand2"))
                        self.w.tag_bind(click_tag,"<Leave>",lambda e: self.w.config(cursor="xterm"))
                        def _make_toggle(et,at,tt,st,w):
                            def _toggle(e):
                                st[0]=not st[0]; now_open=st[0]
                                w.tag_configure(et,elide=not now_open)
                                ar=w.tag_ranges(at)
                                if ar: w.delete(ar[0],ar[1]); w.insert(ar[0],"∨  " if now_open else "▷  ",(at,))
                                tr=w.tag_ranges(tt)
                                if tr: w.tag_configure(tt,foreground="#1a73e8" if now_open else "#333333")
                                return "break"
                            return _toggle
                        card_open=[False]
                        toggle_fn=_make_toggle(elide_tag,arrow_tag,title_tag,card_open,self.w)
                        acc_toggles[click_tag]=toggle_fn
                        self.w.tag_bind(click_tag,"<Enter>",lambda e: self.w.config(cursor="hand2"))
                        self.w.tag_bind(click_tag,"<Leave>",lambda e: self.w.config(cursor="xterm"))
                        self.w.insert("end","\n",("acc_divider",elide_tag))
                    elif self._in_card_body and d==self._body_depth:
                        self._in_card_body=False
                        self._flush_nl(elide_tag=self._cur_elide_tag)
                        self.w.insert("end","\n",("acc_divider",self._cur_elide_tag))
                    elif self._in_accordion and d==self._acc_depth:
                        self._in_accordion=False; self._flush_nl()
                        self.w.insert("end","\n",("acc_divider",)); self._nl(1)
                    return
                if tag=="table": self._in_table=False; return
                if self._in_table: return
                if tag=="a": self._link_tag=None
                elif tag in ("h1","h2","h3","h4","h5","h6"): self._nl(1)
                elif tag=="p": self._nl(2)
                elif tag=="li": self._nl(1)
                elif tag in ("ul","ol"):
                    if self.list_stack: self.list_stack.pop()
                    self._nl(1)
                elif tag=="blockquote": self._nl(1)

            def handle_data(self,data):
                if self.skip or self._in_table: return
                text=html_lib.unescape(data).strip("\n\r")
                if not text: return
                if self._in_card_hdr: self._hdr_title+=text; return
                self._insert(text,*self._inline_tags())

        _Parser(t).feed(body_html)
        if hasattr(self,"split_rendered_text") and t is self.rendered_text:
            self._render_html_to_text(body_html, _target=self.split_rendered_text)

    def _show_pane(self, mode):
        if mode=="split":
            self._single_container.pack_forget()
            self._split_pane.pack(fill="both", expand=True)
        else:
            self._split_pane.pack_forget()
            self._single_container.pack(fill="both", expand=True)
            self.source_frame.pack_forget(); self.rendered_frame.pack_forget()
            if mode=="source":
                self.source_frame.pack(fill="both", expand=True)
                wrap_on=self._source_wrap.get()
                self.preview_text.config(wrap="word" if wrap_on else "none")
                if wrap_on: self._src_hsb.pack_forget()
                elif not self._src_hsb.winfo_ismapped():
                    self._src_hsb.pack(side="bottom", fill="x")
            else:
                self.rendered_frame.pack(fill="both", expand=True)

    def _update_module_nav(self):
        if self._modules and self.split_modules_var.get():
            total=len(self._modules); mod=self._modules[self._module_idx]
            mod_num=mod["number"]; n_named=sum(1 for m in self._modules if m["number"]>0)
            self._mod_label_var.set(f"Module {mod_num} of {n_named}")
            self._mod_prev_btn.config(state="normal" if self._module_idx>0 else "disabled",
                                      fg=FG if self._module_idx>0 else FG3)
            self._mod_next_btn.config(state="normal" if self._module_idx<total-1 else "disabled",
                                      fg=FG if self._module_idx<total-1 else FG3)
            if not self.mod_nav_frame.winfo_ismapped():
                self.mod_nav_frame.pack(side="left", padx=(8,0))
            n_files=total
            self.module_info_var.set(
                f"📦  {n_named} module{'s' if n_named!=1 else ''} detected "
                f"— will save {n_files} file{'s' if n_files!=1 else ''}")
        else:
            self.mod_nav_frame.pack_forget(); self._mod_label_var.set(""); self.module_info_var.set("")

    def _module_prev(self):
        if self._module_idx>0: self._module_idx-=1; self._show_module()

    def _module_next(self):
        if self._module_idx<len(self._modules)-1: self._module_idx+=1; self._show_module()

    def _show_module(self):
        self._update_module_nav()
        if not self._modules: return
        html = self._modules[self._module_idx]["html"]; mode=self.preview_mode.get()
        self._show_pane(mode)
        if mode=="source": self._set_preview(html)
        elif mode=="rendered": self._render_html_to_text(html)
        else: self._set_preview(html); self._render_html_to_text(html)
        if self._search_visible: self._search_run()

    def _pick_css(self):
        path = filedialog.askopenfilename(title="Select Brightspace CSS File",
                                          filetypes=[("CSS files","*.css"),("All files","*.*")])
        if path:
            try:
                self.custom_css = Path(path).read_text(encoding="utf-8",errors="replace")
                self.custom_css_path = path
                self.css_label.config(text=f"✅  {Path(path).name}", fg=SUCCESS)
                self._refresh_preview()
            except Exception as e: messagebox.showerror("CSS Error", str(e))

    def _clear_css(self):
        self.custom_css=""; self.custom_css_path=None
        self.css_label.config(text="No CSS loaded — default preview styles", fg=FG3)
        self._refresh_preview()

    def _collect_settings(self):
        hmap = {ws:v.get() for ws,v in self.heading_vars.items() if v.get()!="(skip)"}
        return {"heading_map":hmap,"ul_transform":self.ul_var.get(),"ol_transform":self.ol_var.get(),
                "bq_transform":self.bq_var.get(),"bq_hr":self.bq_hr_var.get(),
                "accordion_heading":self.acc_head_var.get(),"strip_style":self.strip_style_var.get()}

    def _refresh_preview(self, *_):
        if not self.selected_file: return
        try:
            settings = self._collect_settings()
            body_html, _imgs, log, _links = convert_docx(self.selected_file, settings)
            self._update_log(log)
            if self.split_modules_var.get(): self._modules = split_modules(body_html)
            else: self._modules = []
            if self._modules:
                self._module_idx = max(0, min(self._module_idx, len(self._modules)-1))
                self._update_module_nav(); self._show_module(); return
            else: self._update_module_nav()
            mode=self.preview_mode.get(); self._show_pane(mode)
            if mode=="source": self._set_preview(body_html)
            elif mode=="rendered": self._render_html_to_text(body_html)
            else: self._set_preview(body_html); self._render_html_to_text(body_html)
        except Exception as exc:
            self._show_pane("source"); self._set_preview(f"Preview error:\n{exc}")
        if self._search_visible: self._search_run()

    def _set_preview(self, text):
        for tw in (self.preview_text, self.split_source_text):
            tw.config(state="normal"); tw.delete("1.0","end"); tw.insert("1.0",text)
            tw.config(state="disabled")

    def _open_in_browser(self):
        if not self.selected_file:
            messagebox.showinfo("No file","Please select a .docx file first."); return
        try:
            settings = self._collect_settings()
            body_html, img_map, _log, _links = convert_docx(self.selected_file, settings)
            css=self.custom_css or DEFAULT_PREVIEW_CSS
            full=wrap_html(body_html, title=Path(self.selected_file).stem, css=css)
            tmp=tempfile.NamedTemporaryFile(mode="w",suffix=".html",delete=False,encoding="utf-8")
            tmp.write(full); tmp.close()
            if img_map:
                tmp_img_dir=Path(tmp.name).parent/"images"; tmp_img_dir.mkdir(exist_ok=True)
                for fname,blob in img_map.items(): (tmp_img_dir/fname).write_bytes(blob)
            webbrowser.open(f"file://{tmp.name}")
        except Exception as exc: messagebox.showerror("Error", str(exc))

    def _save_zip(self):
        if not self.selected_file:
            messagebox.showinfo("No file","Please select a .docx file first."); return
        try:
            settings = self._collect_settings()
            body_html, img_map, _, _links = convert_docx(self.selected_file, settings)
            css=self.custom_css or DEFAULT_PREVIEW_CSS; src_path=Path(self.selected_file); stem=src_path.stem
            entries=[]
            if self.split_modules_var.get(): modules=split_modules(body_html)
            else: modules=[]
            if modules:
                for mod in modules:
                    html=mod["html"]
                    if self.full_html_var.get(): html=wrap_html(html,title=mod["title"],css=css)
                    arc_name=f"{stem}_module_{mod['number']:02d}.html"
                    entries.append((arc_name,html.encode("utf-8")))
            else:
                if self.full_html_var.get(): output=wrap_html(body_html,title=stem,css=css)
                else: output=body_html
                entries.append((stem+".html",output.encode("utf-8")))
            for fname,blob in (img_map or {}).items(): entries.append(("images/"+fname,blob))
            zip_path=filedialog.asksaveasfilename(title="Save ZIP archive",defaultextension=".zip",
                                                   filetypes=[("ZIP archives","*.zip"),("All files","*.*")],
                                                   initialfile=stem+".zip",
                                                   initialdir=self._last_file_dir or str(Path.home()))
            if not zip_path: return
            import io; buf=io.BytesIO()
            with zipfile.ZipFile(buf,"w",compression=zipfile.ZIP_DEFLATED) as zf:
                for arc_name,data in entries: zf.writestr(arc_name,data)
            Path(zip_path).write_bytes(buf.getvalue())
            n_html=sum(1 for n,_ in entries if n.endswith(".html")); n_images=len(img_map or {})
            parts=[str(n_html)+" HTML file"+("s" if n_html!=1 else "")]
            if n_images: parts.append(str(n_images)+" image"+("s" if n_images!=1 else ""))
            self.status_var.set("✅  Saved "+Path(zip_path).name+" ("+", ".join(parts)+")")
        except Exception as exc: messagebox.showerror("ZIP error", str(exc))

    def _copy_html(self):
        if not self.selected_file:
            self.status_var.set("⚠  No file loaded — nothing to copy."); return
        try:
            if self._modules and self.split_modules_var.get():
                html=self._modules[self._module_idx]["html"]; label=f"Module {self._module_idx+1}"
            else:
                settings=self._collect_settings()
                html,_,_l,_links=convert_docx(self.selected_file,settings); label="HTML"
            self.clipboard_clear(); self.clipboard_append(html)
            self._copy_btn.config(text="✓ Copied!", fg=SUCCESS)
            self.after(1500, lambda: self._copy_btn.config(text="Copy HTML  📋",fg=FG2))
            self.status_var.set(f"📋  {label} copied to clipboard.")
        except Exception as exc: messagebox.showerror("Copy error", str(exc))

    def _run_convert(self):
        if not self.selected_file:
            messagebox.showwarning("No file","Please select a .docx file first."); return
        try:
            settings = self._collect_settings()
            body_html, img_map, log, link_annotations = convert_docx(self.selected_file, settings)
            css=self.custom_css or DEFAULT_PREVIEW_CSS; src_path=Path(self.selected_file)
            self._update_log(log)

            def _write_images(output_dir):
                if not img_map: return 0
                img_dir=output_dir/"images"; img_dir.mkdir(exist_ok=True)
                for fname,blob in img_map.items(): (img_dir/fname).write_bytes(blob)
                return len(img_map)

            def _img_note(n):
                if n==0: return ""
                return f"  ({n} image{'s' if n!=1 else ''} → images/)"

            if self.split_modules_var.get(): modules=split_modules(body_html)
            else: modules=[]

            if modules:
                if self.save_next_var.get(): out_dir=src_path.parent
                else:
                    out_dir_str=filedialog.askdirectory(title="Choose folder to save module files")
                    if not out_dir_str: self.status_var.set("Save cancelled."); return
                    out_dir=Path(out_dir_str)
                stem=src_path.stem; saved=[]
                for mod in modules:
                    num=mod["number"]; content=mod["html"]
                    if self.full_html_var.get(): content=wrap_html(content,title=mod["title"],css=css)
                    fname=out_dir/f"{stem}_module_{num:02d}.html"
                    fname.write_text(content,encoding="utf-8"); saved.append(fname.name)
                    self._log_append("info",f"Saved {fname.name} ({fname.stat().st_size//1024} KB)")
                n_imgs=_write_images(out_dir)
                if n_imgs: self._log_append("info",f"Images written to {out_dir/'images'}")
                if link_annotations:
                    lp=out_dir/f"{stem}_image_links.txt"
                    if _write_image_links_file(link_annotations,lp):
                        self._log_append("info",f"Image links saved to {lp.name}")
                self._save_config()
                short=", ".join(saved[:3])+(" …" if len(saved)>3 else "")
                self.status_var.set(f"✅  Saved {len(saved)} module file{'s' if len(saved)!=1 else ''}: {short}{_img_note(n_imgs)}")
                self._refresh_preview(); return

            if self.full_html_var.get(): output=wrap_html(body_html,title=src_path.stem,css=css)
            else: output=body_html

            if self.save_next_var.get(): out_path=src_path.with_suffix(".html")
            else:
                out_path=filedialog.asksaveasfilename(title="Save HTML as",defaultextension=".html",
                                                       filetypes=[("HTML files","*.html"),("All files","*.*")],
                                                       initialfile=src_path.stem+".html")
                if not out_path: self.status_var.set("Save cancelled."); return
                out_path=Path(out_path)

            out_path.write_text(output,encoding="utf-8")
            n_imgs=_write_images(out_path.parent)
            self._log_append("info",f"Saved {out_path.name} ({out_path.stat().st_size//1024} KB)")
            if n_imgs: self._log_append("info",f"Images written to {out_path.parent/'images'}")
            if link_annotations:
                lp=out_path.parent/f"{src_path.stem}_image_links.txt"
                if _write_image_links_file(link_annotations,lp):
                    self._log_append("info",f"Image links saved to {lp.name}")
            self._save_config()
            self.status_var.set(f"✅  Saved: {out_path.name}{_img_note(n_imgs)}")
            self._refresh_preview()
        except Exception as exc:
            messagebox.showerror("Conversion error", str(exc))
            self.status_var.set(f"❌  Error: {exc}")

    # ── Config persistence ────────────────────────────────────

    def _save_config(self):
        cfg = {"headings":{ws:v.get() for ws,v in self.heading_vars.items()},
               "ul":self.ul_var.get(),"ol":self.ol_var.get(),"bq":self.bq_var.get(),
               "bq_hr":self.bq_hr_var.get(),"acc_heading":self.acc_head_var.get(),
               "full_html":self.full_html_var.get(),"save_next":self.save_next_var.get(),
               "split_modules":self.split_modules_var.get(),"strip_style":self.strip_style_var.get(),
               "para_font_size":self.para_font_size.get(),"css_path":self.custom_css_path or "",
               "last_preset":self.preset_var.get(),"last_file_dir":self._last_file_dir or ""}
        try: CONFIG_FILE.write_text(json.dumps(cfg,indent=2), encoding="utf-8")
        except Exception: pass

    def _load_config(self):
        try:
            cfg = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
            for ws,var in self.heading_vars.items():
                if ws in cfg.get("headings",{}): var.set(cfg["headings"][ws])
            self.ul_var.set(cfg.get("ul","ul")); self.ol_var.set(cfg.get("ol","ol"))
            self.bq_var.set(cfg.get("bq","blockquote")); self.bq_hr_var.set(cfg.get("bq_hr",False))
            self.acc_head_var.set(cfg.get("acc_heading",ACCORDION_HEADING))
            self.full_html_var.set(cfg.get("full_html",False))
            self.save_next_var.set(cfg.get("save_next",True))
            self.split_modules_var.set(cfg.get("split_modules",True))
            self.strip_style_var.set(cfg.get("strip_style",True))
            self.para_font_size.set(cfg.get("para_font_size",19))
            lfd=cfg.get("last_file_dir","")
            if lfd and Path(lfd).is_dir(): self._last_file_dir=lfd
            css_path=cfg.get("css_path","")
            if css_path and Path(css_path).exists():
                self.custom_css=Path(css_path).read_text(encoding="utf-8",errors="replace")
                self.custom_css_path=css_path
                self.css_label.config(text=f"✅  {Path(css_path).name}", fg=SUCCESS)
            self._last_preset_name=cfg.get("last_preset","")
        except Exception: self._last_preset_name=""


# ═══════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    try:
        import tkinterdnd2
        ConverterApp.__bases__ = (tkinterdnd2.TkinterDnD.Tk,)
    except ImportError:
        pass

    app = ConverterApp()
    app.mainloop()
