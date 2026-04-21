"""
Word → Brightspace HTML Converter
──────────────────────────────────
Two-tab GUI:
  • Convert  — file picker, output options, live preview (HTML source / rendered)
  • Settings — element transform dropdowns, blockquote HR toggle, CSS upload

Accordion detection: a single-cell Word table whose cell contains one or more
H4 paragraphs is treated as an accordion group.  Each H4 → card title; the
normal paragraphs that follow it (until the next H4 or end of cell) → card body.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import json, re, tempfile, webbrowser, shutil, subprocess, sys
from pathlib import Path


# ═══════════════════════════════════════════════════════════════
#  DEPENDENCY BOOTSTRAP
#  Runs before any third-party import so the file always loads.
#  Shows a GUI installer window if required packages are missing.
# ═══════════════════════════════════════════════════════════════

# Each entry: (import_name, pip_package_name, required)
# required=False → installed if present, skipped gracefully if not
_DEPS = [
    ("docx",       "python-docx",  True),
    ("tkinterdnd2","tkinterdnd2",  False),
]

def _check_and_install_deps():
    """Return True if all required deps are satisfied (after install if needed)."""
    import importlib

    missing_required = []
    missing_optional = []
    for import_name, pip_name, required in _DEPS:
        try:
            importlib.import_module(import_name)
        except ImportError:
            (missing_required if required else missing_optional).append(
                (import_name, pip_name))

    if not missing_required and not missing_optional:
        return True   # nothing to do

    to_install = missing_required + missing_optional

    # ── Build the installer dialog ────────────────────────────────────────
    root = tk.Tk()
    root.title("Install Required Packages")
    root.resizable(False, False)
    root.configure(bg="#1a2632")

    PAD = 16
    W   = 520

    tk.Label(root, text="📦  Missing Python Packages",
             font=("Segoe UI", 13, "bold"),
             bg="#1a2632", fg="#ffffff",
             pady=14).pack(fill="x", padx=PAD)

    # Description
    if missing_required:
        req_names = ", ".join(p for _, p in missing_required)
        desc = (f"The following package{'s are' if len(missing_required) > 1 else ' is'} "
                f"required to run the converter:\n\n  {req_names}")
    else:
        desc = ""
    if missing_optional:
        opt_names = ", ".join(p for _, p in missing_optional)
        opt_note  = (f"\nThe following optional package adds drag-and-drop support:\n\n"
                     f"  {opt_names}")
        desc += opt_note

    tk.Label(root, text=desc.strip(),
             font=("Segoe UI", 10), justify="left",
             bg="#1a2632", fg="#b0bec5",
             wraplength=W - PAD * 2).pack(anchor="w", padx=PAD)

    tk.Frame(root, bg="#3a4a58", height=1).pack(fill="x", padx=PAD, pady=(12, 0))

    # Output log area
    log_frame = tk.Frame(root, bg="#0f1117")
    log_frame.pack(fill="both", padx=PAD, pady=(0, 0))

    log_text = tk.Text(
        log_frame, height=10, width=62,
        font=("Cascadia Code", 9),
        bg="#0f1117", fg="#e2e8f0",
        relief="flat", bd=0, padx=10, pady=8,
        state="disabled", wrap="word")
    log_vsb = ttk.Scrollbar(log_frame, orient="vertical",
                             command=log_text.yview)
    log_text.configure(yscrollcommand=log_vsb.set)
    log_vsb.pack(side="right", fill="y")
    log_text.pack(fill="both", expand=True)

    # Status / button row
    btn_frame = tk.Frame(root, bg="#1a2632")
    btn_frame.pack(fill="x", padx=PAD, pady=12)

    status_var = tk.StringVar(
        value=f"Ready to install {len(to_install)} package"
              f"{'s' if len(to_install) > 1 else ''}.")
    tk.Label(btn_frame, textvariable=status_var,
             font=("Segoe UI", 9), bg="#1a2632", fg="#78909c",
             anchor="w").pack(side="left", expand=True, fill="x")

    install_btn  = tk.Button(btn_frame)
    skip_btn     = tk.Button(btn_frame)
    _result      = {"ok": False}

    def _log(line, colour="#e2e8f0"):
        log_text.config(state="normal")
        log_text.insert("end", line + "\n", (colour,))
        log_text.tag_configure(colour, foreground=colour)
        log_text.see("end")
        log_text.config(state="disabled")
        root.update_idletasks()

    def _run_install():
        install_btn.config(state="disabled", text="Installing…")
        skip_btn.config(state="disabled")
        status_var.set("Installing — please wait…")
        root.update_idletasks()

        all_ok = True
        for import_name, pip_name in to_install:
            _log(f"▶  pip install {pip_name}", "#78909c")
            try:
                proc = subprocess.Popen(
                    [sys.executable, "-m", "pip", "install", pip_name],
                    stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                    text=True, bufsize=1)
                for line in proc.stdout:
                    _log("   " + line.rstrip(), "#94a3b8")
                proc.wait()
                if proc.returncode == 0:
                    _log(f"✅  {pip_name} installed successfully.", "#4ade80")
                else:
                    _log(f"❌  {pip_name} install failed (exit {proc.returncode}).",
                         "#f87171")
                    if any(n == import_name for n, _ in missing_required):
                        all_ok = False
            except Exception as exc:
                _log(f"❌  Error: {exc}", "#f87171")
                all_ok = False

        if all_ok:
            status_var.set("All packages installed. Starting the app…")
            _log("\n✅  Done! The converter will now open.", "#4ade80")
            _result["ok"] = True
            root.after(1200, root.destroy)
        else:
            status_var.set("Some required packages failed to install.")
            install_btn.config(state="normal", text="Retry")
            skip_btn.config(state="normal")

    def _skip():
        if missing_required:
            if not messagebox.askyesno(
                "Skip installation?",
                "Required packages are missing. The app will not work correctly "
                "without them.\n\nSkip anyway?",
                parent=root):
                return
        root.destroy()

    install_btn.config(
        text=f"Install {len(to_install)} Package{'s' if len(to_install) > 1 else ''}",
        font=("Segoe UI", 10, "bold"),
        bg="#006fbf", fg="#ffffff", relief="flat", bd=0,
        padx=14, pady=6, cursor="hand2",
        activebackground="#004a8f", activeforeground="#ffffff",
        command=_run_install)
    install_btn.pack(side="right", padx=(8, 0))

    skip_btn.config(
        text="Skip",
        font=("Segoe UI", 10),
        bg="#2a2f45", fg="#94a3b8", relief="flat", bd=0,
        padx=10, pady=6, cursor="hand2",
        activebackground="#383d56", activeforeground="#e2e8f0",
        command=_skip)
    skip_btn.pack(side="right")

    # Centre the window on screen
    root.update_idletasks()
    sw, sh = root.winfo_screenwidth(), root.winfo_screenheight()
    rw, rh = root.winfo_width(), root.winfo_height()
    root.geometry(f"+{(sw - rw) // 2}+{(sh - rh) // 2}")

    root.mainloop()
    return _result["ok"]


# ── Run bootstrap before any third-party import ──────────────────────────────
# When run directly, check deps and offer to install before importing them.
# When imported as a module, attempt the import directly (may raise ImportError).
if __name__ == "__main__":
    if not _check_and_install_deps():
        sys.exit(1)

# Third-party imports — guaranteed present if bootstrap ran successfully.
from docx import Document
from docx.oxml.ns import qn



# ═══════════════════════════════════════════════════════════════
#  DEFAULTS
# ═══════════════════════════════════════════════════════════════

CONFIG_FILE   = Path.home() / ".brightspace_converter_config.json"
PRESETS_FILE  = Path.home() / ".brightspace_converter_presets.json"

DEFAULT_HEADING_MAP = {
    "Heading 1": "h2",
    "Heading 2": "h3",
    "Heading 3": "h4",
    "Heading 4": "h5",
    "Heading 5": "h6",
    "Heading 6": "h6",
}

BLOCKQUOTE_STYLES = {"Quote", "Block Text", "Intense Quote"}
ACCORDION_HEADING = "Heading 4"

HEADING_OPTS = ["h1", "h2", "h3", "h4", "h5", "h6", "p", "(skip)"]
LIST_OPTS    = ["ul", "ol"]
BQ_OPTS      = ["blockquote", "p", 'div class="callout"', "(skip)"]


# ═══════════════════════════════════════════════════════════════
#  TOOLTIP
# ═══════════════════════════════════════════════════════════════

class Tooltip:
    """A simple hover tooltip for any tkinter widget."""

    def __init__(self, widget, text, delay=500):
        self.widget  = widget
        self.text    = text
        self.delay   = delay
        self._id     = None
        self._tip_win = None
        widget.bind("<Enter>",    self._schedule)
        widget.bind("<Leave>",    self._cancel)
        widget.bind("<ButtonPress>", self._cancel)

    def _schedule(self, event=None):
        self._cancel()
        self._id = self.widget.after(self.delay, self._show)

    def _cancel(self, event=None):
        if self._id:
            self.widget.after_cancel(self._id)
            self._id = None
        self._hide()

    def _show(self):
        if self._tip_win:
            return
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 4
        self._tip_win = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        tw.attributes("-topmost", True)
        lbl = tk.Label(
            tw, text=self.text, justify="left",
            background="#ffffcc", foreground="#222",
            relief="solid", borderwidth=1,
            font=("Segoe UI", 9),
            wraplength=280, padx=6, pady=4,
        )
        lbl.pack()

    def _hide(self):
        if self._tip_win:
            self._tip_win.destroy()
            self._tip_win = None


def tip(widget, text):
    """Convenience wrapper — attach a tooltip and return the widget."""
    Tooltip(widget, text)
    return widget


# ═══════════════════════════════════════════════════════════════
#  LOW-LEVEL HELPERS
# ═══════════════════════════════════════════════════════════════

def escape_html(text):
    if not text:
        return ""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def style_name(para):
    try:
        return para.style.name or ""
    except Exception:
        return ""


def _get_num_pr(para):
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            nid = numPr.find(qn("w:numId"))
            ilv = numPr.find(qn("w:ilvl"))
            numId = int(nid.get(qn("w:val"), 0)) if nid is not None else 0
            ilvl  = int(ilv.get(qn("w:val"), 0)) if ilv is not None else 0
            if numId > 0:
                return numId, ilvl
    return None, None


def _lookup_num_fmt(para, numId, ilvl):
    try:
        np = para.part.numbering_part
        if np is None:
            return ""
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        abId = None
        for n in np._element.findall(f"{{{ns}}}num"):
            if int(n.get(qn("w:numId"), -1)) == numId:
                ref = n.find(qn("w:abstractNumId"))
                if ref is not None:
                    abId = int(ref.get(qn("w:val"), -1))
                break
        if abId is None:
            return ""
        for ab in np._element.findall(f"{{{ns}}}abstractNum"):
            if int(ab.get(qn("w:abstractNumId"), -1)) == abId:
                lvls = ab.findall(qn("w:lvl"))
                if ilvl < len(lvls):
                    nf = lvls[ilvl].find(qn("w:numFmt"))
                    if nf is not None:
                        return nf.get(qn("w:val"), "")
    except Exception:
        pass
    return ""


def is_list_para(para):
    numId, _ = _get_num_pr(para)
    if numId is not None:
        return True
    sn = style_name(para)
    return sn in ("List Bullet", "List Bullet 2", "List Bullet 3",
                  "List Number", "List Number 2", "List Number 3")


def is_ordered_para(para):
    numId, ilvl = _get_num_pr(para)
    if numId is not None:
        fmt = _lookup_num_fmt(para, numId, ilvl)
        if fmt:
            return fmt in ("decimal", "lowerLetter", "upperLetter",
                           "lowerRoman", "upperRoman", "ordinal",
                           "cardinalText", "ordinalText")
    return "Number" in style_name(para)


def list_indent_level(para):
    _, ilvl = _get_num_pr(para)
    if ilvl is not None:
        return ilvl
    m = re.search(r"(\d+)$", style_name(para))
    return max(0, int(m.group(1)) - 1) if m else 0


# ═══════════════════════════════════════════════════════════════
#  IMAGE EXTRACTION
# ═══════════════════════════════════════════════════════════════

def _collect_images(para, image_collector):
    """Find all w:drawing elements in *para* and populate *image_collector*.

    *image_collector* is a dict: {output_filename: bytes}

    Returns a dict: {rId: (output_filename, alt_text)} for use by
    _extract_runs.  alt_text is read from wp:docPr/@descr (Word's
    "Alt Text" panel); falls back to empty string if absent.
    """
    NS_WP    = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    NS_BLIP  = "http://schemas.openxmlformats.org/drawingml/2006/main"
    NS_REL   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    BLIP_TAG = f"{{{NS_BLIP}}}blip"
    DOC_PR   = f"{{{NS_WP}}}docPr"

    rId_to_info = {}   # rId -> (fname, alt_text)
    try:
        part = para.part
        for drawing in para._p.iter(qn("w:drawing")):
            alt_text = ""
            for doc_pr in drawing.iter(DOC_PR):
                alt_text = (doc_pr.get("descr") or
                            doc_pr.get("title") or "").strip()
                break
            for blip in drawing.iter(BLIP_TAG):
                rId = blip.get(f"{{{NS_REL}}}embed")
                if not rId or rId in rId_to_info:
                    continue
                try:
                    img_part = part.rels[rId].target_part
                    blob     = img_part.blob
                    fname    = img_part.filename
                    if fname in image_collector and image_collector[fname] != blob:
                        ext   = Path(fname).suffix
                        fname = f"{Path(fname).stem}_{rId}{ext}"
                    rId_to_info[rId] = (fname, alt_text)
                    image_collector[fname] = blob
                except Exception:
                    pass
    except Exception:
        pass
    return rId_to_info




def _extract_runs(el, hyperlink_map, image_collector=None, rId_to_fname=None):
    """Recursively extract inline HTML from an element's runs and hyperlinks.
    Walks into w:ins, w:del, w:hyperlink, and bare w:r elements so that
    tracked-changes markup and nested structures are all captured.

    *image_collector* and *rId_to_fname* are passed in by para_to_inline_html
    when image extraction is active; both may be None for table cells where
    we don't yet have a per-paragraph rId map.
    """
    parts = []
    for child in el:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            parts.append(run_to_html(child))
        elif tag == "drawing":
            # Emit <img> tags for any images referenced in this drawing element.
            if rId_to_fname is not None:
                NS_BLIP = "http://schemas.openxmlformats.org/drawingml/2006/main"
                NS_REL  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                BLIP_TAG = f"{{{NS_BLIP}}}blip"
                for blip in child.iter(BLIP_TAG):
                    rId = blip.get(f"{{{NS_REL}}}embed")
                    info = rId_to_fname.get(rId) if rId else None
                    if info:
                        fname, alt_text = info
                        safe_src = escape_html(fname)
                        safe_alt = escape_html(alt_text) if alt_text else safe_src
                        parts.append(
                            f'<img src="images/{safe_src}" alt="{safe_alt}" '
                            f'style="max-width:100%;">'
                        )
        elif tag == "hyperlink":
            rid = child.get(qn("r:id"))
            url = hyperlink_map.get(rid, "#")
            inner = _extract_runs(child, hyperlink_map, image_collector, rId_to_fname)
            if inner:
                parts.append(f'<a href="{escape_html(url)}">{inner}</a>')
        elif tag in ("ins", "del", "smartTag", "sdt", "sdtContent"):
            # Recurse into tracked-change wrappers and smart tags
            parts.append(_extract_runs(child, hyperlink_map, image_collector, rId_to_fname))
    return "".join(parts)


def para_to_inline_html(para, image_collector=None):
    """Convert a paragraph's inline content to HTML, handling runs, hyperlinks,
    tracked changes (w:ins / w:del), line breaks, and embedded images.

    *image_collector*, if provided, is a dict {filename: bytes} that will be
    populated with any images found in this paragraph.
    """
    hyperlink_map = {}
    try:
        for rel in para.part.rels.values():
            if "hyperlink" in rel.reltype:
                hyperlink_map[rel.rId] = rel._target
    except Exception:
        pass

    # Build the rId→(filename, alt_text) map for this paragraph's images
    rId_to_fname = None
    if image_collector is not None:
        rId_to_fname = _collect_images(para, image_collector)

    return _extract_runs(para._p, hyperlink_map, image_collector, rId_to_fname)


def _w_attr(el, local_name):
    """Get a w:-namespaced attribute value safely.
    lxml stores XML attributes in Clark notation ({ns}local). Word attributes
    like w:val, w:type are stored as {ns}val, not as 'w:val'.  Try Clark first,
    then fall back to the bare local name for resilience.
    """
    v = el.get(qn(f"w:{local_name}"))
    if v is None:
        v = el.get(local_name)
    return v


def _rpr_flag(rpr, tag):
    """Return True if a run-property toggle tag is active (present and not val='0')."""
    if rpr is None:
        return False
    el = rpr.find(qn(tag))
    if el is None:
        return False
    val = _w_attr(el, "val")
    return val not in ("0", "false", "off")


def run_to_html(r):
    """Convert a single <w:r> element to inline HTML, including line breaks."""
    rpr = r.find(qn("w:rPr"))

    parts = []
    for child in r:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "t":
            parts.append(escape_html(child.text or ""))
        elif tag == "br":
            # w:br — could be a line break or page break; emit <br> for line breaks
            br_type = _w_attr(child, "type") or ""
            if br_type not in ("page", "column"):
                parts.append("<br>")

    text = "".join(parts)
    if not text:
        return ""

    # Apply inline formatting outermost → innermost so nesting is correct:
    # <strong><em><u><del>text</del></u></em></strong>
    if rpr is not None:
        if _rpr_flag(rpr, "w:strike"): text = f"<del>{text}</del>"
        if _rpr_flag(rpr, "w:u"):      text = f"<u>{text}</u>"
        if _rpr_flag(rpr, "w:i"):      text = f"<em>{text}</em>"
        if _rpr_flag(rpr, "w:b"):      text = f"<strong>{text}</strong>"
    return text


# ═══════════════════════════════════════════════════════════════
#  LIST RENDERER
# ═══════════════════════════════════════════════════════════════

def render_list(list_paras, force_tag=None, image_collector=None):
    lines = []
    stack = []
    for para in list_paras:
        level = list_indent_level(para)
        tag   = force_tag if force_tag else ("ol" if is_ordered_para(para) else "ul")
        inline = para_to_inline_html(para, image_collector)
        while len(stack) <= level:
            lines.append("  " * len(stack) + f"<{tag}>")
            stack.append((len(stack), tag))
        while len(stack) > level + 1:
            _, ct = stack.pop()
            lines.append("  " * len(stack) + f"</{ct}>")
        lines.append("  " * (level + 1) + f"<li>{inline}</li>")
    while stack:
        _, ct = stack.pop()
        lines.append("  " * len(stack) + f"</{ct}>")
    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════
#  ACCORDION
# ═══════════════════════════════════════════════════════════════

def render_accordion(cell_paragraphs, acc_heading):
    lines = ['<div class="accordion">']
    in_card = False
    for para in cell_paragraphs:
        if style_name(para) == acc_heading:
            if in_card:
                lines += ['      </div>', '    </div>', '  </div>']
            title = para_to_inline_html(para)
            lines += [
                '  <div class="card">',
                '    <div class="card-header">',
                f'      <h2 class="card-title">{title}</h2>',
                '    </div>',
                '    <div class="collapse">',
                '      <div class="card-body">',
            ]
            in_card = True
        else:
            inline = para_to_inline_html(para)
            if inline.strip():
                prefix = "        " if in_card else ""
                lines.append(f"{prefix}<p>{inline}</p>")
    if in_card:
        lines += ['      </div>', '    </div>', '  </div>']
    lines.append('</div>')
    return "\n".join(lines)


def is_accordion_table(table, acc_heading):
    try:
        cells = list({id(c): c for row in table.rows for c in row.cells}.values())
        if len(cells) != 1:
            return False
        return any(style_name(p) == acc_heading for p in cells[0].paragraphs)
    except Exception:
        return False


# ═══════════════════════════════════════════════════════════════
#  TABLE RENDERER
# ═══════════════════════════════════════════════════════════════

def render_table(table):
    lines = ["<table>"]
    for i, row in enumerate(table.rows):
        lines.append("  <tr>")
        seen = set()
        for cell in row.cells:
            cid = id(cell)
            if cid in seen:
                continue
            seen.add(cid)
            tag = "th" if i == 0 else "td"
            content = " ".join(para_to_inline_html(p) for p in cell.paragraphs).strip()
            lines.append(f"    <{tag}>{content}</{tag}>")
        lines.append("  </tr>")
    lines.append("</table>")
    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════
#  CORE CONVERTER
# ═══════════════════════════════════════════════════════════════

def _accept_tracked_changes(doc):
    """Strip tracked-change markup in memory so all content is treated as accepted.

    - w:ins  (inserted text) : unwrap — keep the runs inside
    - w:del  (deleted text)  : remove entirely — the text was deleted
    - w:rPrChange / w:pPrChange : remove — these are old formatting snapshots

    The original file on disk is never modified.
    """
    body = doc.element.body
    ns   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _tag(local):
        return f"{{{ns}}}{local}"

    # Process repeatedly until no more changes remain (handles nesting)
    for _ in range(10):
        changed = False

        for el in body.iter():
            local = el.tag.split("}")[-1] if "}" in el.tag else el.tag

            # w:ins — keep children, drop the wrapper
            if local == "ins":
                parent = el.getparent()
                if parent is not None:
                    idx = list(parent).index(el)
                    for i, child in enumerate(list(el)):
                        el.remove(child)
                        parent.insert(idx + i, child)
                    parent.remove(el)
                    changed = True
                    break  # restart iteration after mutating tree

            # w:del — drop entirely (deleted content should not appear)
            elif local == "del":
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
                    changed = True
                    break

            # Old formatting snapshots — not needed
            elif local in ("rPrChange", "pPrChange", "sectPrChange",
                           "tblPrChange", "trPrChange", "tcPrChange"):
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
                    changed = True
                    break

        if not changed:
            break


# Matches a bare http/https URL and nothing else
_URL_RE = re.compile(r"^https?://\S+$", re.IGNORECASE)


def _extract_comments(doc):
    """Return {comment_id_str: text} from the document comments part.
    Returns {} when no comments part exists.
    """
    comment_map = {}
    try:
        cp = doc.part.comments_part
        if cp is None:
            return {}
        for cel in cp._element.iter(qn("w:comment")):
            cid = cel.get(qn("w:id"))
            if cid is None:
                continue
            texts = [t.text or "" for t in cel.iter(qn("w:t"))]
            comment_map[cid] = "".join(texts).strip()
    except Exception:
        pass
    return comment_map


def _collect_image_link_annotations(doc):
    """Find images annotated with a URL-only Word comment.

    A "URL-only comment" is a Word comment whose entire text is a single
    http/https URL.  Any image in the same paragraph as that comment range
    is recorded along with the current module heading.

    Returns [{"module": "Module 1: ...", "image": "image1.png",
               "url": "https://..."}, ...]
    """
    comment_map = _extract_comments(doc)
    url_comments = {cid: txt for cid, txt in comment_map.items()
                    if _URL_RE.match(txt)}
    if not url_comments:
        return []

    NS_BLIP = "http://schemas.openxmlformats.org/drawingml/2006/main"
    NS_REL  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    MOD_RE  = re.compile(
        r"Module\s+([1-9][0-9]?)\s*:(.+)", re.IGNORECASE | re.DOTALL)

    results = []
    current_module = "Document"
    open_ids = set()

    for el in doc.element.body:
        local = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if local != "p":
            continue
        # Track module from heading text
        txt = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        m = MOD_RE.match(txt)
        if m:
            current_module = "Module " + m.group(1) + ": " + m.group(2).strip()
        # Open URL-comment ranges
        for crs in el.findall(qn("w:commentRangeStart")):
            cid = crs.get(qn("w:id"))
            if cid in url_comments:
                open_ids.add(cid)
        # Record image + URL if any open ranges
        if open_ids:
            for drawing in el.iter(qn("w:drawing")):
                for blip in drawing.iter(f"{{{NS_BLIP}}}blip"):
                    rId = blip.get(f"{{{NS_REL}}}embed")
                    if rId:
                        try:
                            fname = doc.part.rels[rId].target_part.filename
                        except Exception:
                            fname = rId
                        for cid in sorted(open_ids):
                            results.append({
                                "module": current_module,
                                "image":  fname,
                                "url":    url_comments[cid],
                            })
        # Close comment ranges
        for cre in el.findall(qn("w:commentRangeEnd")):
            open_ids.discard(cre.get(qn("w:id")))

    return results


def _write_image_links_file(link_annotations, out_path):
    """Write image link annotations to a plain-text file.
    Returns True if anything was written.
    """
    if not link_annotations:
        return False
    lines = [
        "Image Link Annotations",
        "=" * 60,
        "Generated by Word -> Brightspace Converter",
        "",
    ]
    cur_mod = None
    for entry in link_annotations:
        mod = entry["module"]
        if mod != cur_mod:
            if cur_mod is not None:
                lines.append("")
            lines.append(f"[ {mod} ]")
            cur_mod = mod
        lines.append(f"  {entry['image']}")
        lines.append(f"    {entry['url']}")
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return True


def convert_docx(docx_path, settings):
    """Convert *docx_path* to HTML.

    Returns a tuple ``(body_html, image_map, log, link_annotations)`` where:
    - *image_map* is ``{filename: bytes}`` of every extracted image
    - *log* is a list of ``{"level": "info"|"warn"|"error", "msg": str}`` dicts
    - *link_annotations* is a list of {module, image, url} dicts
      from URL-only comments attached to images
    """
    hmap     = settings.get("heading_map", DEFAULT_HEADING_MAP)
    ul_out   = settings.get("ul_transform", "ul")
    ol_out   = settings.get("ol_transform", "ol")
    bq_out   = settings.get("bq_transform", "blockquote")
    bq_hr    = settings.get("bq_hr", False)
    acc_head = settings.get("accordion_heading", ACCORDION_HEADING)

    log = []   # list of {"level": "info"|"warn"|"error", "msg": str}

    def _log(level, msg):
        log.append({"level": level, "msg": msg})

    # Copy to a temp file first — sidesteps OneDrive file locks and
    # "file in use by Word" errors that cause python-docx's "Package not found".
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        shutil.copy2(docx_path, tmp_path)
        doc = Document(tmp_path)
    except Exception:
        doc = Document(docx_path)
    _accept_tracked_changes(doc)

    image_collector = {}   # {filename: bytes} — populated as paragraphs are processed
    html_parts = []

    # Track counts for the log summary
    n_headings     = 0
    n_paras        = 0
    n_lists        = 0
    n_blockquotes  = 0
    n_tables       = 0
    n_accordions   = 0
    n_skipped      = 0
    unknown_styles = {}   # style name → count

    # Build ordered block list: ("para", para_obj) | ("table", table_obj)
    para_iter  = iter(doc.paragraphs)
    table_iter = iter(doc.tables)
    next_para  = next(para_iter, None)
    next_table = next(table_iter, None)

    blocks = []
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p" and next_para is not None:
            blocks.append(("para", next_para))
            next_para = next(para_iter, None)
        elif tag == "tbl" and next_table is not None:
            blocks.append(("table", next_table))
            next_table = next(table_iter, None)

    i = 0
    while i < len(blocks):
        kind, obj = blocks[i]

        # ── Table ─────────────────────────────────────────────
        if kind == "table":
            if is_accordion_table(obj, acc_head):
                cells = list({id(c): c for row in obj.rows
                               for c in row.cells}.values())
                html_parts.append(render_accordion(cells[0].paragraphs, acc_head))
                n_accordions += 1
            else:
                html_parts.append(render_table(obj))
                n_tables += 1
            i += 1
            continue

        para = obj
        sn   = style_name(para)

        # ── Heading ───────────────────────────────────────────
        if sn in hmap:
            out_tag = hmap[sn]
            if out_tag != "(skip)":
                inline = para_to_inline_html(para, image_collector)
                if out_tag == "p":
                    html_parts.append(f"<p>{inline}</p>")
                else:
                    html_parts.append(f"<{out_tag}>{inline}</{out_tag}>")
                n_headings += 1
            else:
                n_skipped += 1
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────
        if sn in BLOCKQUOTE_STYLES:
            if bq_out == "(skip)":
                n_skipped += 1
                i += 1
                continue
            open_tag, close_tag = _bq_tags(bq_out)
            collected = []
            while (i < len(blocks) and blocks[i][0] == "para"
                   and style_name(blocks[i][1]) in BLOCKQUOTE_STYLES):
                collected.append(para_to_inline_html(blocks[i][1], image_collector))
                i += 1
            html_parts.append(open_tag)
            if bq_hr and bq_out == "blockquote":
                html_parts.append("  <hr>")
            for line in collected:
                html_parts.append(f"  <p>{line}</p>")
            if bq_hr and bq_out == "blockquote":
                html_parts.append("  <hr>")
            html_parts.append(close_tag)
            n_blockquotes += len(collected)
            continue

        # ── List ──────────────────────────────────────────────
        if is_list_para(para):
            list_paras = []
            while (i < len(blocks) and blocks[i][0] == "para"
                   and is_list_para(blocks[i][1])):
                list_paras.append(blocks[i][1])
                i += 1
            all_ord = all(is_ordered_para(p) for p in list_paras)
            all_bul = all(not is_ordered_para(p) for p in list_paras)
            force   = ol_out if all_ord else (ul_out if all_bul else None)
            html_parts.append(render_list(list_paras, force_tag=force,
                                          image_collector=image_collector))
            n_lists += len(list_paras)
            continue

        # ── Empty ─────────────────────────────────────────────
        inline = para_to_inline_html(para, image_collector)
        if not inline.strip():
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────
        # Track unrecognised non-empty styles for the log
        if sn and sn not in ("Normal", "Body Text", "Default Paragraph Font",
                              "No Spacing", ""):
            unknown_styles[sn] = unknown_styles.get(sn, 0) + 1

        html_parts.append(f"<p>{inline}</p>")
        n_paras += 1
        i += 1

    # ── Build log ─────────────────────────────────────────────

    # Summary line
    parts_summary = []
    if n_headings:    parts_summary.append(f"{n_headings} heading{'s' if n_headings != 1 else ''}")
    if n_paras:       parts_summary.append(f"{n_paras} paragraph{'s' if n_paras != 1 else ''}")
    if n_lists:       parts_summary.append(f"{n_lists} list item{'s' if n_lists != 1 else ''}")
    if n_blockquotes: parts_summary.append(f"{n_blockquotes} blockquote line{'s' if n_blockquotes != 1 else ''}")
    if n_tables:      parts_summary.append(f"{n_tables} table{'s' if n_tables != 1 else ''}")
    if n_accordions:  parts_summary.append(f"{n_accordions} accordion{'s' if n_accordions != 1 else ''}")
    _log("info", "Converted: " + (", ".join(parts_summary) if parts_summary else "no content found"))

    # Images
    if image_collector:
        _log("info", f"{len(image_collector)} image{'s' if len(image_collector) != 1 else ''} extracted: "
             + ", ".join(image_collector.keys()))
    else:
        _log("info", "No images found in document")

    # Skipped elements
    if n_skipped:
        _log("warn", f"{n_skipped} element{'s' if n_skipped != 1 else ''} skipped "
             f"(heading or blockquote set to '(skip)' in Settings)")

    # Unknown styles
    for sn, count in sorted(unknown_styles.items(), key=lambda x: -x[1]):
        _log("warn", f"Unrecognised style '{sn}' ({count}×) — converted as plain <p>")

    link_annotations = _collect_image_link_annotations(doc)
    return "\n".join(html_parts), image_collector, log, link_annotations


def _bq_tags(bq_out):
    if bq_out == "blockquote":
        return "<blockquote>", "</blockquote>"
    if bq_out == "p":
        return "<p>", "</p>"
    tag_name = bq_out.split()[0]
    return f"<{bq_out}>", f"</{tag_name}>"


def wrap_html(body_html, title="", css=""):
    css_block = f"  <style>\n{css}\n  </style>" if css else ""
    return (
        "<!DOCTYPE html>\n"
        '<html lang="en">\n'
        "<head>\n"
        '  <meta charset="UTF-8">\n'
        '  <meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
        f"  <title>{escape_html(title)}</title>\n"
        f"{css_block}\n"
        "</head>\n"
        "<body>\n"
        f"{body_html}\n"
        "</body>\n"
        "</html>"
    )


# ═══════════════════════════════════════════════════════════════
#  MODULE SPLITTER
# ═══════════════════════════════════════════════════════════════

# Matches any heading tag whose text begins with "Module N:" (N = 1–99).
# Case-insensitive, tolerates extra whitespace.
MODULE_RE = re.compile(
    r'<(h[1-6])>(\s*Module\s+([1-9][0-9]?)\s*:.*?)</\1>',
    re.IGNORECASE | re.DOTALL,
)


def split_modules(body_html):
    """Split *body_html* into per-module chunks.

    Returns a list of dicts:
        [{"number": 1, "title": "Module 1: Introduction", "html": "..."}, ...]

    If no Module headings are found, returns an empty list (caller treats
    the document as a single un-split file).
    """
    matches = list(MODULE_RE.finditer(body_html))
    if not matches:
        return []

    modules = []
    for i, m in enumerate(matches):
        start = m.start()
        end   = matches[i + 1].start() if i + 1 < len(matches) else len(body_html)
        tag   = m.group(1)          # e.g. "h2"
        title = re.sub(r"<[^>]+>", "", m.group(2)).strip()  # strip any inline tags
        num   = int(m.group(3))
        chunk = body_html[start:end].strip()
        modules.append({"number": num, "title": title, "html": chunk})
    return modules




PAD    = 12
FONT   = ("Segoe UI", 10)
FONT_B = ("Segoe UI", 10, "bold")
FONT_S = ("Segoe UI", 9)
MONO   = ("Cascadia Code", 9) if True else ("Courier New", 9)

# Dark palette
BG        = "#1a1d27"   # main background
BG2       = "#22263a"   # panel / sidebar background
BG3       = "#2a2f45"   # card / input background
BORDER    = "#383d56"   # subtle borders
ACCENT    = "#6c8ef5"   # primary blue-violet
ACCENT2   = "#a78bfa"   # secondary violet
FG        = "#e2e8f0"   # primary text
FG2       = "#94a3b8"   # secondary / muted text
FG3       = "#64748b"   # very muted
SUCCESS   = "#4ade80"   # green
WARN      = "#fbbf24"   # amber
ERR       = "#f87171"   # red
DROP_BG   = "#252942"   # drop zone background
DROP_HL   = "#6c8ef5"   # drop zone highlight border

# Preview colours (HTML source pane)
PRE_BG    = "#0f1117"
PRE_FG    = "#e2e8f0"

# Default readable CSS injected into "Rendered" preview when no custom CSS loaded
DEFAULT_PREVIEW_CSS = """
/* ── Google Font (Lato — D2L's body font) ──────────────────────────────── */
@import url('https://fonts.googleapis.com/css2?family=Lato:ital,wght@0,400;0,700;1,400&display=swap');

/* ── Page shell — mimics Brightspace's grey outer background ────────────── */
html {
  background: #f5f5f5;
  min-height: 100%;
}

/* ── Content well — the white card D2L renders module pages inside ────────── */
body {
  font-family: 'Lato', 'Segoe UI', system-ui, -apple-system, sans-serif;
  font-size: 16px;
  line-height: 1.6;
  color: #212121;
  background: #ffffff;
  max-width: 860px;
  margin: 28px auto;
  padding: 32px 40px 48px;
  box-shadow: 0 1px 4px rgba(0,0,0,.14);
  border-radius: 2px;
}

/* ── Headings ────────────────────────────────────────────────────────────── */
h1, h2, h3, h4, h5, h6 {
  color: #212121;
  font-weight: 700;
  line-height: 1.3;
  margin: 1.5em 0 0.4em;
}
h1 { font-size: 1.875em; }
h2 { font-size: 1.5em;
     border-bottom: 1px solid #e0e0e0;
     padding-bottom: 0.25em; }
h3 { font-size: 1.25em; }
h4 { font-size: 1.1em; }
h5 { font-size: 1em; }
h6 { font-size: 0.9em; color: #616161; }

.d2l-page-title {
  font-size: 2em;
  font-weight: 700;
  color: #212121;
  margin: 0 0 1.2em;
  padding-bottom: 0.3em;
  border-bottom: 2px solid #006fbf;
}

/* ── Body copy ───────────────────────────────────────────────────────────── */
p { margin: 0 0 1em; }

/* ── Links ───────────────────────────────────────────────────────────────── */
a       { color: #006fbf; text-decoration: underline; }
a:hover { color: #004a8f; }

/* ── Lists ───────────────────────────────────────────────────────────────── */
ul, ol { margin: 0 0 1em 1.5em; padding: 0; }
li     { margin-bottom: 0.3em; }
ul li  { list-style-type: disc; }
ul ul li { list-style-type: circle; }

/* ── Blockquote ──────────────────────────────────────────────────────────── */
blockquote {
  border-left: 4px solid #006fbf;
  margin: 1em 0;
  padding: 0.5em 1em;
  background: #f0f6fb;
  color: #424242;
  border-radius: 0 3px 3px 0;
}

/* ── Tables ──────────────────────────────────────────────────────────────── */
table {
  border-collapse: collapse;
  width: 100%;
  margin: 1em 0 1.5em;
  font-size: 0.95em;
}
th, td {
  border: 1px solid #bdbdbd;
  padding: 9px 13px;
  text-align: left;
  vertical-align: top;
}
th {
  background: #eeeeee;
  font-weight: 700;
  color: #212121;
}
tr:nth-child(even) td { background: #fafafa; }

/* ── Inline formatting ───────────────────────────────────────────────────── */
strong { font-weight: 700; color: #212121; }
em     { font-style: italic; }
u      { text-decoration: underline; }
del    { text-decoration: line-through; color: #757575; }
code {
  background: #f5f5f5;
  border: 1px solid #e0e0e0;
  padding: 1px 5px;
  border-radius: 3px;
  font-size: 0.875em;
  font-family: 'Cascadia Code', 'Consolas', monospace;
  color: #c62828;
}

/* ── Images ──────────────────────────────────────────────────────────────── */
img { max-width: 100%; height: auto; display: block; margin: 0.5em 0; }

/* ── HR ──────────────────────────────────────────────────────────────────── */
hr { border: none; border-top: 1px solid #e0e0e0; margin: 1.5em 0; }

/* ── D2L Accordion component ─────────────────────────────────────────────── */
.accordion {
  border: 1px solid #d3d9de;
  border-radius: 4px;
  overflow: hidden;
  margin: 1em 0 1.5em;
}
.card { border-bottom: 1px solid #d3d9de; }
.card:last-child { border-bottom: none; }
.card-header { background: #f2f3f5; padding: 12px 18px; }
.card-header:hover { background: #e8eaed; }
.card-title { margin: 0; font-size: 1em; font-weight: 700; color: #006fbf; }
.card-body { padding: 14px 18px; background: #ffffff; }
.card-body p:last-child { margin-bottom: 0; }

/* ── D2L Callout / Note box ──────────────────────────────────────────────── */
div.callout {
  background: #fff8e1;
  border-left: 4px solid #f9a825;
  padding: 10px 16px;
  margin: 1em 0;
  border-radius: 0 3px 3px 0;
  color: #424242;
}
"""


class ConverterApp(tk.Tk):

    def __init__(self):
        super().__init__()
        self.title("Word \u2192 Brightspace Converter")
        self.configure(bg=BG)
        self.resizable(True, True)
        self.minsize(1020, 680)

        self.selected_file   = None
        self.custom_css      = ""
        self.custom_css_path = None
        self.presets         = {}          # name → settings dict
        self.preset_var      = tk.StringVar(value="")

        self.heading_vars  = {}
        self.ul_var        = tk.StringVar(value="ul")
        self.ol_var        = tk.StringVar(value="ol")
        self.bq_var        = tk.StringVar(value="blockquote")
        self.bq_hr_var     = tk.BooleanVar(value=False)
        self.acc_head_var  = tk.StringVar(value=ACCORDION_HEADING)
        self.full_html_var = tk.BooleanVar(value=False)
        self.save_next_var = tk.BooleanVar(value=True)
        self.preview_mode  = tk.StringVar(value="source")

        # Module-split state
        self.split_modules_var  = tk.BooleanVar(value=True)
        self._modules           = []   # list of {"number", "title", "html"}
        self._module_idx        = 0    # which module is shown in preview

        # Batch state
        self._mode              = tk.StringVar(value="single")
        self._batch_dir         = Path.cwd()
        self._batch_files       = []   # list of {"path": Path, "var": BooleanVar,
                                       #           "status_var": StringVar,
                                       #           "row": Frame}
        self._batch_running     = False

        self._apply_dark_theme()
        self._build_ui()
        self._load_config()
        self._load_presets_file()

    # ── Dark theme ────────────────────────────────────────────

    def _apply_dark_theme(self):
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure(".", background=BG2, foreground=FG,
                        fieldbackground=BG3, troughcolor=BG3,
                        bordercolor=BORDER, darkcolor=BG2, lightcolor=BG2,
                        selectbackground=ACCENT, selectforeground=FG, font=FONT)
        style.configure("TNotebook", background=BG, borderwidth=0)
        style.configure("TNotebook.Tab", background=BG3, foreground=FG2,
                        padding=[14, 7], borderwidth=0, font=FONT)
        style.map("TNotebook.Tab",
                  background=[("selected", BG2), ("active", BORDER)],
                  foreground=[("selected", FG), ("active", FG)])
        style.configure("TCombobox", fieldbackground=BG3, background=BG3,
                        foreground=FG, arrowcolor=FG2,
                        selectbackground=ACCENT, selectforeground=FG,
                        bordercolor=BORDER, insertcolor=FG)
        style.map("TCombobox", fieldbackground=[("readonly", BG3)],
                  foreground=[("readonly", FG)],
                  bordercolor=[("focus", ACCENT)])
        for orient in ("Vertical", "Horizontal"):
            style.configure(f"{orient}.TScrollbar", background=BG3,
                            troughcolor=BG2, bordercolor=BG2,
                            arrowcolor=FG3, relief="flat", width=10)
        self.option_add("*TCombobox*Listbox.background", BG3)
        self.option_add("*TCombobox*Listbox.foreground", FG)
        self.option_add("*TCombobox*Listbox.selectBackground", ACCENT)
        self.option_add("*TCombobox*Listbox.selectForeground", FG)

    # ── Top-level layout ──────────────────────────────────────

    def _build_ui(self):
        # Header bar
        hdr = tk.Frame(self, bg=BG2, pady=0)
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=ACCENT, width=4).pack(side="left", fill="y")
        tk.Label(hdr, text="  Word \u2192 Brightspace HTML Converter",
                 bg=BG2, fg=FG, font=("Segoe UI", 13, "bold"),
                 pady=11).pack(side="left")
        tk.Frame(hdr, bg=BORDER, height=1).pack(side="bottom", fill="x")

        pane = tk.PanedWindow(self, orient="horizontal", bg=BG,
                              sashwidth=6, sashrelief="flat",
                              sashpad=0)
        pane.pack(fill="both", expand=True)

        left = tk.Frame(pane, bg=BG2, width=380)
        left.pack_propagate(False)
        pane.add(left, minsize=320)

        self.nb = ttk.Notebook(left)
        self.nb.pack(fill="both", expand=True, padx=0, pady=0)

        tc = tk.Frame(self.nb, bg=BG2, padx=PAD, pady=PAD)
        ts = tk.Frame(self.nb, bg=BG2)
        self.nb.add(tc, text="  Convert  ")
        self.nb.add(ts, text="  Settings  ")

        self._build_convert_tab(tc)
        self._build_settings_tab(ts)

        right = tk.Frame(pane, bg=BG)
        pane.add(right, minsize=380)
        self._build_preview_panel(right)

    # ── Convert tab ───────────────────────────────────────────

    def _build_convert_tab(self, p):
        # ── Mode pill toggle ──────────────────────────────────
        pill = tk.Frame(p, bg=BG3, bd=0, highlightthickness=1,
                        highlightbackground=BORDER)
        pill.pack(fill="x", pady=(0, 14))

        def _pill_btn(text, value, side):
            rb = tk.Radiobutton(
                pill, text=text, variable=self._mode, value=value,
                bg=BG3, fg=FG2, font=FONT_B,
                activebackground=ACCENT, activeforeground=FG,
                selectcolor=ACCENT, indicatoron=False,
                relief="flat", bd=0, padx=0, pady=7,
                cursor="hand2",
                command=self._on_mode_change)
            rb.pack(side=side, fill="x", expand=True)
            return rb

        _pill_btn("  Single File  ", "single", "left")
        tk.Frame(pill, bg=BORDER, width=1).pack(side="left", fill="y")
        _pill_btn("  Batch  ", "batch", "left")

        # ── Two swappable inner frames ─────────────────────────
        self._single_frame = tk.Frame(p, bg=BG2)
        self._batch_frame  = tk.Frame(p, bg=BG2)

        self._build_single_frame(self._single_frame)
        self._build_batch_frame(self._batch_frame)

        # Show single by default
        self._single_frame.pack(fill="both", expand=True)

    def _on_mode_change(self):
        if self._mode.get() == "single":
            self._batch_frame.pack_forget()
            self._single_frame.pack(fill="both", expand=True)
        else:
            self._single_frame.pack_forget()
            self._batch_frame.pack(fill="both", expand=True)
            self._batch_scan_dir()

    # ── Single-file inner frame ────────────────────────────────

    def _build_single_frame(self, p):
        # Section label
        tk.Label(p, text="INPUT FILE", font=("Segoe UI", 8, "bold"),
                 bg=BG2, fg=FG3, pady=0).pack(anchor="w", pady=(0, 4))

        # Drop zone
        self.drop_zone = tk.Frame(p, bg=DROP_BG, cursor="hand2",
                                  highlightthickness=2,
                                  highlightbackground=BORDER,
                                  highlightcolor=DROP_HL)
        self.drop_zone.pack(fill="x", pady=(0, 16))

        self.file_icon = tk.Label(self.drop_zone, text="\U0001f4c2",
                                  bg=DROP_BG, font=("Segoe UI", 22), pady=6)
        self.file_icon.pack()
        self.file_label = tk.Label(
            self.drop_zone,
            text="Click or drag & drop a .docx file",
            bg=DROP_BG, fg=FG2, font=FONT, pady=4)
        self.file_label.pack()
        self.file_sub = tk.Label(self.drop_zone, text="",
                                 bg=DROP_BG, fg=FG3, font=FONT_S, pady=2)
        self.file_sub.pack()

        # Bind click and drag-and-drop
        for w in (self.drop_zone, self.file_icon, self.file_label, self.file_sub):
            w.bind("<Button-1>", lambda e: self._pick_file())
            w.bind("<Enter>",    lambda e: self.drop_zone.config(
                highlightbackground=DROP_HL))
            w.bind("<Leave>",    lambda e: self.drop_zone.config(
                highlightbackground=ACCENT if self.selected_file else BORDER))

        self._setup_drag_drop()

        # Output options
        tk.Label(p, text="OUTPUT OPTIONS", font=("Segoe UI", 8, "bold"),
                 bg=BG2, fg=FG3).pack(anchor="w", pady=(0, 6))

        def _dark_cb(text, var, tip_text):
            f = tk.Frame(p, bg=BG2)
            f.pack(fill="x", pady=2)
            cb = tk.Checkbutton(f, text=text, variable=var,
                                bg=BG2, fg=FG, activebackground=BG2,
                                activeforeground=FG, selectcolor=BG3,
                                font=FONT, bd=0, highlightthickness=0,
                                cursor="hand2")
            cb.pack(anchor="w")
            tip(cb, tip_text)
            return cb

        _dark_cb("Wrap in full HTML document", self.full_html_var,
                 "Adds <!DOCTYPE html>, <html>, <head>, <body> tags.\n"
                 "Useful for browser preview. Leave unchecked for a bare\n"
                 "snippet to paste into Brightspace's HTML editor.")
        _dark_cb("Save output alongside source .docx", self.save_next_var,
                 "Saves MyDoc.html next to MyDoc.docx.\n"
                 "Uncheck to choose a custom save location each time.")
        _dark_cb("Split into one file per Module", self.split_modules_var,
                 "When headings like 'Module 1:', 'Module 2:' … are detected,\n"
                 "each module is saved as its own HTML file.\n"
                 "(MyDoc_module_01.html, MyDoc_module_02.html, …)\n\n"
                 "The preview panel will show one module at a time\n"
                 "and let you navigate between them with arrow buttons.")

        btn_convert = tk.Button(p, text="Convert  \u2192", font=FONT_B,
                                bg=ACCENT, fg=FG, relief="flat", bd=0,
                                padx=16, pady=10, cursor="hand2",
                                activebackground=ACCENT2, activeforeground=FG,
                                command=self._run_convert)
        btn_convert.pack(fill="x", pady=(16, 4))
        tip(btn_convert, "Convert the selected .docx and save the HTML output.")

        self.status_var = tk.StringVar(value="")
        self.status_lbl = tk.Label(p, textvariable=self.status_var,
                                   bg=BG2, fg=SUCCESS, font=FONT,
                                   wraplength=320, justify="left")
        self.status_lbl.pack(anchor="w", pady=(4, 0))

        # Module detection info (shown when modules are found)
        self.module_info_var = tk.StringVar(value="")
        self.module_info_lbl = tk.Label(p, textvariable=self.module_info_var,
                                        bg=BG2, fg=ACCENT2, font=FONT_S,
                                        wraplength=320, justify="left")
        self.module_info_lbl.pack(anchor="w", pady=(2, 0))

        btn_preview = tk.Button(p, text="Refresh Preview", font=FONT,
                                bg=BG3, fg=FG, relief="flat", bd=0,
                                padx=10, pady=7, cursor="hand2",
                                activebackground=BORDER, activeforeground=FG,
                                command=self._refresh_preview)
        btn_preview.pack(fill="x", pady=(14, 0))
        tip(btn_preview, "Re-run the conversion and update the preview panel.")

        # ── Conversion Log ────────────────────────────────────
        self._log_expanded = tk.BooleanVar(value=False)

        log_header = tk.Frame(p, bg=BG2, cursor="hand2")
        log_header.pack(fill="x", pady=(14, 0))
        self._log_arrow = tk.Label(log_header, text="▶", font=FONT_S,
                                   bg=BG2, fg=FG3)
        self._log_arrow.pack(side="left")
        self._log_header_lbl = tk.Label(
            log_header, text="CONVERSION LOG", font=("Segoe UI", 8, "bold"),
            bg=BG2, fg=FG3)
        self._log_header_lbl.pack(side="left", padx=(4, 0))
        self._log_badge = tk.Label(log_header, text="", font=FONT_S,
                                   bg=BG2, fg=FG3)
        self._log_badge.pack(side="left", padx=(6, 0))

        for w in (log_header, self._log_arrow, self._log_header_lbl, self._log_badge):
            w.bind("<Button-1>", lambda e: self._toggle_log())

        self._log_body = tk.Frame(p, bg=BG3)
        # Not packed until expanded

        self._log_text = tk.Text(
            self._log_body, font=("Cascadia Code", 8), wrap="word",
            bg=BG3, fg=FG2, relief="flat", bd=0,
            padx=8, pady=6, height=8, state="disabled",
            selectbackground=ACCENT, selectforeground=FG,
            cursor="arrow")
        log_vsb = ttk.Scrollbar(self._log_body, orient="vertical",
                                command=self._log_text.yview)
        self._log_text.configure(yscrollcommand=log_vsb.set)
        log_vsb.pack(side="right", fill="y")
        self._log_text.pack(fill="both", expand=True)

        # Configure colour tags for the three severity levels
        self._log_text.tag_configure("info",  foreground=FG2)
        self._log_text.tag_configure("warn",  foreground=WARN)
        self._log_text.tag_configure("error", foreground=ERR)
        self._log_text.tag_configure("dim",   foreground=FG3)

    # ── Batch inner frame ──────────────────────────────────────

    def _build_batch_frame(self, p):
        # Directory bar
        tk.Label(p, text="FOLDER", font=("Segoe UI", 8, "bold"),
                 bg=BG2, fg=FG3).pack(anchor="w", pady=(0, 4))

        dir_row = tk.Frame(p, bg=BG2)
        dir_row.pack(fill="x", pady=(0, 8))

        self._batch_dir_var = tk.StringVar(value=str(self._batch_dir))
        dir_entry = tk.Entry(dir_row, textvariable=self._batch_dir_var,
                             font=FONT_S, bg=BG3, fg=FG2,
                             insertbackground=FG, relief="flat", bd=4,
                             state="readonly")
        dir_entry.pack(side="left", fill="x", expand=True, padx=(0, 6))

        tk.Button(dir_row, text="Browse…", font=FONT_S, bg=BG3, fg=FG,
                  relief="flat", bd=0, padx=10, pady=5, cursor="hand2",
                  activebackground=BORDER, activeforeground=FG,
                  command=self._batch_browse).pack(side="left")

        # Select-all / none controls
        sel_row = tk.Frame(p, bg=BG2)
        sel_row.pack(fill="x", pady=(0, 4))
        self._batch_count_lbl = tk.Label(sel_row, text="", font=FONT_S,
                                         bg=BG2, fg=FG3)
        self._batch_count_lbl.pack(side="left")

        def _sel_lnk(text, cmd):
            lbl = tk.Label(sel_row, text=text, font=FONT_S, bg=BG2,
                           fg=ACCENT, cursor="hand2")
            lbl.pack(side="right", padx=(6, 0))
            lbl.bind("<Button-1>", lambda e: cmd())
            return lbl

        _sel_lnk("None", self._batch_select_none)
        _sel_lnk("All",  self._batch_select_all)
        tk.Label(sel_row, text="Select:", font=FONT_S, bg=BG2,
                 fg=FG3).pack(side="right", padx=(0, 4))

        # Scrollable file list
        list_outer = tk.Frame(p, bg=BG3, highlightthickness=1,
                              highlightbackground=BORDER)
        list_outer.pack(fill="both", expand=True, pady=(0, 8))

        self._batch_canvas = tk.Canvas(list_outer, bg=BG3,
                                       highlightthickness=0)
        list_vsb = ttk.Scrollbar(list_outer, orient="vertical",
                                 command=self._batch_canvas.yview)
        self._batch_canvas.configure(yscrollcommand=list_vsb.set)
        list_vsb.pack(side="right", fill="y")
        self._batch_canvas.pack(side="left", fill="both", expand=True)

        self._batch_list_frame = tk.Frame(self._batch_canvas, bg=BG3)
        self._batch_list_win = self._batch_canvas.create_window(
            (0, 0), window=self._batch_list_frame, anchor="nw")

        def _on_list_resize(e=None):
            self._batch_canvas.configure(
                scrollregion=self._batch_canvas.bbox("all"))
            self._batch_canvas.itemconfig(
                self._batch_list_win,
                width=self._batch_canvas.winfo_width())

        self._batch_list_frame.bind("<Configure>", lambda e: _on_list_resize())
        self._batch_canvas.bind("<Configure>",     lambda e: _on_list_resize())
        self._batch_canvas.bind_all(
            "<MouseWheel>",
            lambda e: self._batch_canvas.yview_scroll(
                int(-1 * (e.delta / 120)), "units"))

        # Progress bar (hidden until a run starts)
        self._batch_progress_frame = tk.Frame(p, bg=BG2)
        self._batch_progress_var = tk.DoubleVar(value=0)
        self._batch_progress_lbl_var = tk.StringVar(value="")
        tk.Label(self._batch_progress_frame,
                 textvariable=self._batch_progress_lbl_var,
                 font=FONT_S, bg=BG2, fg=FG2).pack(anchor="w")
        self._batch_bar = ttk.Progressbar(
            self._batch_progress_frame, variable=self._batch_progress_var,
            maximum=100, length=300)
        self._batch_bar.pack(fill="x", pady=(2, 0))

        # Convert All button
        self._batch_btn = tk.Button(
            p, text="Convert All  \u2192", font=FONT_B,
            bg=ACCENT, fg=FG, relief="flat", bd=0,
            padx=16, pady=10, cursor="hand2",
            activebackground=ACCENT2, activeforeground=FG,
            command=self._run_batch)
        self._batch_btn.pack(fill="x", pady=(4, 0))
        tip(self._batch_btn,
            "Convert all checked .docx files using the current settings.\n"
            "Click a completed filename in the list to preview it.")

    # ── Batch methods ─────────────────────────────────────────

    def _batch_browse(self):
        d = filedialog.askdirectory(title="Select folder containing .docx files",
                                    initialdir=str(self._batch_dir))
        if d:
            self._batch_dir = Path(d)
            self._batch_dir_var.set(str(self._batch_dir))
            self._batch_scan_dir()

    def _batch_scan_dir(self):
        """Scan _batch_dir for .docx files and populate the file list."""
        # Clear existing rows
        for w in self._batch_list_frame.winfo_children():
            w.destroy()
        self._batch_files = []

        files = sorted(self._batch_dir.glob("*.docx"),
                       key=lambda p: p.name.lower())

        if not files:
            tk.Label(self._batch_list_frame,
                     text="No .docx files found in this folder.",
                     font=FONT_S, bg=BG3, fg=FG3,
                     padx=8, pady=10).pack(anchor="w")
            self._batch_count_lbl.config(text="0 files")
            return

        for i, path in enumerate(files):
            row_bg = BG3 if i % 2 == 0 else "#262b40"
            row = tk.Frame(self._batch_list_frame, bg=row_bg)
            row.pack(fill="x")

            var = tk.BooleanVar(value=True)
            status_var = tk.StringVar(value="")

            cb = tk.Checkbutton(row, variable=var, bg=row_bg,
                                activebackground=row_bg, selectcolor=BG2,
                                bd=0, highlightthickness=0, cursor="hand2")
            cb.pack(side="left", padx=(6, 2))

            name_lbl = tk.Label(row, text=path.name, font=FONT_S,
                                bg=row_bg, fg=FG, anchor="w",
                                cursor="hand2")
            name_lbl.pack(side="left", fill="x", expand=True, pady=5)

            status_lbl = tk.Label(row, textvariable=status_var,
                                  font=FONT_S, bg=row_bg, fg=FG3,
                                  width=6, anchor="e")
            status_lbl.pack(side="right", padx=(0, 8))

            entry = {"path": path, "var": var, "status_var": status_var,
                     "row": row, "name_lbl": name_lbl, "row_bg": row_bg}
            self._batch_files.append(entry)

            # Clicking the filename previews (after conversion) or loads file
            name_lbl.bind("<Button-1>",
                          lambda e, p=path: self._batch_preview_file(p))

        self._batch_count_lbl.config(
            text=f"{len(files)} file{'s' if len(files) != 1 else ''} found")

    def _batch_select_all(self):
        for f in self._batch_files:
            f["var"].set(True)

    def _batch_select_none(self):
        for f in self._batch_files:
            f["var"].set(False)

    def _batch_preview_file(self, path):
        """Load a file into the preview panel from the batch list."""
        self._load_file(str(path))

    def _run_batch(self):
        if self._batch_running:
            return
        checked = [f for f in self._batch_files if f["var"].get()]
        if not checked:
            messagebox.showinfo("Nothing selected",
                                "Check at least one file to convert.")
            return

        # Reset all statuses
        for f in self._batch_files:
            f["status_var"].set("")
            f["name_lbl"].config(fg=FG)

        self._batch_running = True
        self._batch_btn.config(state="disabled", text="Converting…")
        self._batch_progress_frame.pack(fill="x", pady=(0, 6),
                                        before=self._batch_btn)
        self._batch_progress_var.set(0)
        total = len(checked)

        def _convert_one(idx):
            if idx >= total:
                # All done
                self._batch_running = False
                self._batch_btn.config(state="normal",
                                       text="Convert All  \u2192")
                self._batch_progress_lbl_var.set(
                    f"Done — {total} file{'s' if total != 1 else ''} converted")
                return

            entry = checked[idx]
            entry["status_var"].set("⏳")
            entry["name_lbl"].config(fg=FG)
            self._batch_progress_lbl_var.set(
                f"Converting {idx + 1} / {total}: {entry['path'].name}")
            self._batch_progress_var.set((idx / total) * 100)
            self.update_idletasks()

            try:
                settings               = self._collect_settings()
                body_html, img_map, lg, _bl = convert_docx(
                    str(entry["path"]), settings)
                css      = self.custom_css or DEFAULT_PREVIEW_CSS
                src_path = entry["path"]

                def _write_images(out_dir):
                    if not img_map:
                        return
                    img_dir = out_dir / "images"
                    img_dir.mkdir(exist_ok=True)
                    for fname, blob in img_map.items():
                        (img_dir / fname).write_bytes(blob)

                if self.split_modules_var.get():
                    modules = split_modules(body_html)
                else:
                    modules = []

                if modules:
                    for mod in modules:
                        content = mod["html"]
                        if self.full_html_var.get():
                            content = wrap_html(content,
                                                title=mod["title"], css=css)
                        fname = src_path.parent / \
                            f"{src_path.stem}_module_{mod['number']:02d}.html"
                        fname.write_text(content, encoding="utf-8")
                    _write_images(src_path.parent)
                    entry["status_var"].set(f"✅ ×{len(modules)}")
                else:
                    output = (wrap_html(body_html, title=src_path.stem, css=css)
                              if self.full_html_var.get() else body_html)
                    out_path = src_path.with_suffix(".html")
                    out_path.write_text(output, encoding="utf-8")
                    _write_images(src_path.parent)
                    entry["status_var"].set("✅")

                entry["name_lbl"].config(fg=SUCCESS)

            except Exception as exc:
                entry["status_var"].set("❌")
                entry["name_lbl"].config(fg=ERR)
                entry["status_var"].set(f"❌")
                # Store error for tooltip
                tip(entry["name_lbl"], f"Error: {exc}")

            self._batch_progress_var.set(((idx + 1) / total) * 100)
            # Schedule next file via after() so UI stays responsive
            self.after(10, lambda: _convert_one(idx + 1))

        self.after(10, lambda: _convert_one(0))

    # ── Drag-and-drop setup ───────────────────────────────────
    def _setup_drag_drop(self):
        """Enable file drag-and-drop onto the drop zone using tkinterdnd2 if
        available, with a graceful fallback (click-only) if not installed."""
        try:
            # tkinterdnd2 must be installed: pip install tkinterdnd2
            from tkinterdnd2 import DND_FILES

            # Register the drop target on all drop-zone children
            drop_widgets = (self.drop_zone, self.file_icon,
                            self.file_label, self.file_sub)
            for w in drop_widgets:
                w.drop_target_register(DND_FILES)
                w.dnd_bind("<<Drop>>", self._on_drop)

            # Also register the entire window so that dragging over ANY part
            # of the app while holding a file triggers the highlight — the OS
            # only fires <<DragEnter>> on registered targets, so we register
            # the root window as well and forward events to the drop zone.
            self.drop_target_register(DND_FILES)
            self.dnd_bind("<<DragEnter>>", self._on_drag_enter)
            self.dnd_bind("<<DragLeave>>", self._on_drag_leave)
            self.dnd_bind("<<Drop>>",      self._on_drop)

            # Per-widget enter/leave so the highlight works precisely inside
            # the drop zone even without the root-level binding
            for w in drop_widgets:
                w.dnd_bind("<<DragEnter>>", self._on_drag_enter)
                w.dnd_bind("<<DragLeave>>", self._on_drag_leave)

        except Exception:
            # tkinterdnd2 not available — click still works
            pass

    def _on_drag_enter(self, event=None):
        """Highlight the drop zone when a file is dragged over the app.

        Cancel any pending leave debounce so crossing child widget boundaries
        (which fire spurious Leave/Enter pairs) does not cause flickering.
        """
        if hasattr(self, "_drag_leave_id") and self._drag_leave_id:
            self.after_cancel(self._drag_leave_id)
            self._drag_leave_id = None
        self.drop_zone.config(
            highlightbackground=DROP_HL,
            bg="#2e3350")
        for w in (self.file_icon, self.file_label, self.file_sub):
            w.config(bg="#2e3350")

    def _on_drag_leave(self, event=None):
        """Remove highlight when the drag cursor leaves the app.

        Debounced by 50 ms so that crossing internal child-widget boundaries
        (which fire a Leave then immediately an Enter) does not cause a flash.
        """
        if hasattr(self, "_drag_leave_id") and self._drag_leave_id:
            self.after_cancel(self._drag_leave_id)
        self._drag_leave_id = self.after(50, self._do_drag_leave)

    def _do_drag_leave(self):
        """Actually reset the drop zone — called after the debounce delay."""
        self._drag_leave_id = None
        self.drop_zone.config(
            highlightbackground=BORDER if not self.selected_file else ACCENT,
            bg=DROP_BG)
        for w in (self.file_icon, self.file_label, self.file_sub):
            w.config(bg=DROP_BG)

    def _on_drop(self, event):
        """Handle a file dropped onto the drop zone."""
        raw = event.data.strip()
        # tkinterdnd2 wraps paths with spaces in braces: {C:/my file.docx}
        if raw.startswith("{") and raw.endswith("}"):
            path = raw[1:-1]
        else:
            path = raw.split()[0]   # take first file if multiple dropped
        # Cancel any pending debounce and reset the drop zone immediately
        if hasattr(self, "_drag_leave_id") and self._drag_leave_id:
            self.after_cancel(self._drag_leave_id)
            self._drag_leave_id = None
        self._do_drag_leave()
        self._load_file(path)

    def _pick_file(self):
        path = filedialog.askopenfilename(
            title="Select Word Document",
            filetypes=[("Word Documents", "*.docx"), ("All files", "*.*")])
        if path:
            self._load_file(path)

    def _load_file(self, path):
        if not path.lower().endswith(".docx"):
            messagebox.showwarning("Wrong file type",
                                   "Please select a .docx Word document.")
            return
        self.selected_file = path
        if hasattr(self, "_d2l_title_var"):
            self._d2l_title_var.set(Path(path).stem)
        name = Path(path).name
        self.file_icon.config(text="\U0001f4c4")
        self.file_label.config(text=name, fg=FG)
        self.file_sub.config(text=str(Path(path).parent), fg=FG3)
        self.drop_zone.config(highlightbackground=ACCENT)
        self.status_var.set("")
        self._refresh_preview()

    # ── Settings tab ──────────────────────────────────────────

    def _build_settings_tab(self, parent):
        canvas = tk.Canvas(parent, bg=BG2, highlightthickness=0)
        vsb    = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        inner  = tk.Frame(canvas, bg=BG2, padx=PAD, pady=PAD)
        win_id = canvas.create_window((0, 0), window=inner, anchor="nw")

        def _resize(e=None):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(win_id, width=canvas.winfo_width())

        inner.bind("<Configure>", lambda e: _resize())
        canvas.bind("<Configure>",  lambda e: _resize())
        canvas.bind_all("<MouseWheel>",
                        lambda e: canvas.yview_scroll(
                            int(-1 * (e.delta / 120)), "units"))

        # ── Presets ───────────────────────────────────────────
        self._sep(inner, "Presets",
                  "Save the current settings as a named preset and reload them later.\n\n"
                  "Presets store: heading map, list transforms, blockquote style,\n"
                  "accordion heading, and the CSS file path.")

        pf = tk.Frame(inner, bg=BG2)
        pf.pack(fill="x", pady=(0, 8))

        # Dropdown
        self.preset_combo = ttk.Combobox(
            pf, textvariable=self.preset_var,
            values=[], state="readonly", width=22)
        self.preset_combo.grid(row=0, column=0, sticky="w", padx=(0, 6))
        self.preset_combo.bind("<<ComboboxSelected>>", lambda e: self._load_preset())
        tip(self.preset_combo, "Choose a saved preset to load its settings.")

        def _icon_btn(parent, text, cmd, tip_text, col):
            b = tk.Button(parent, text=text, font=FONT, bg=BG3, fg=FG,
                          relief="flat", bd=0, padx=9, pady=4, cursor="hand2",
                          activebackground=BORDER, activeforeground=FG,
                          command=cmd)
            b.grid(row=0, column=col, padx=(0, 4))
            tip(b, tip_text)
            return b

        _icon_btn(pf, "Save…",   self._save_preset_dialog,   "Save current settings as a new preset or overwrite an existing one.", 1)
        _icon_btn(pf, "Rename…", self._rename_preset_dialog,  "Rename the selected preset.",                                          2)
        _icon_btn(pf, "Delete",  self._delete_preset,          "Delete the selected preset.",                                          3)

        # ── Headings ──────────────────────────────────────────
        self._sep(inner, "Heading Transform",
                  "Map each Word heading level to an HTML heading tag.\n\n"
                  "Default: Heading 1 \u2192 h2, Heading 2 \u2192 h3, etc.\n"
                  "This offset exists because Brightspace page titles\n"
                  "already occupy the h1 slot — using h2 as your top-level\n"
                  "heading avoids duplicate h1s in the page outline.\n\n"
                  "Options:\n"
                  "  h1\u2013h6  \u2192  HTML heading tags (h1 = largest)\n"
                  "  p       \u2192  plain body paragraph, no heading style\n"
                  "  (skip)  \u2192  omit this heading level from output entirely")

        hf = tk.Frame(inner, bg=BG2)
        hf.pack(fill="x", pady=(0, 8))
        tk.Label(hf, text="Word Style", font=("Segoe UI", 9, "bold"),
                 fg=FG3, bg=BG2).grid(row=0, column=0, sticky="w", padx=(0, 8))
        tk.Label(hf, text="Output Tag", font=("Segoe UI", 9, "bold"),
                 fg=FG3, bg=BG2).grid(row=0, column=2, sticky="w")

        word_styles = [f"Heading {n}" for n in range(1, 7)]
        defaults    = list(DEFAULT_HEADING_MAP.values())
        for idx, (ws, dfl) in enumerate(zip(word_styles, defaults)):
            tk.Label(hf, text=ws, font=FONT, bg=BG2, fg=FG).grid(
                row=idx+1, column=0, sticky="w", pady=3, padx=(0, 6))
            tk.Label(hf, text="\u2192", font=FONT, bg=BG2,
                     fg=FG3).grid(row=idx+1, column=1, padx=8)
            var = tk.StringVar(value=dfl)
            self.heading_vars[ws] = var
            cb = ttk.Combobox(hf, textvariable=var, values=HEADING_OPTS,
                              state="readonly", width=10)
            cb.grid(row=idx+1, column=2, sticky="w", pady=3)
            cb.bind("<<ComboboxSelected>>", lambda e: self._refresh_preview())
            tip(cb,
                f"Output HTML tag for Word's '{ws}'.\n\n"
                "h1  \u2192  largest heading (rarely needed — Brightspace owns h1)\n"
                "h2  \u2192  section heading (recommended for top-level)\n"
                "h3  \u2192  sub-section heading\n"
                "h4  \u2192  sub-sub-section heading\n"
                "h5/h6  \u2192  minor headings\n"
                "p  \u2192  plain paragraph text, no heading styling\n"
                "(skip)  \u2192  remove this heading from the output entirely")

        # ── Lists ─────────────────────────────────────────────
        self._sep(inner, "List Transform",
                  "Choose the HTML list element used for bullet and numbered lists.\n\n"
                  "ul = unordered list  \u2192  renders as bullet points ( \u2022 )\n"
                  "ol = ordered list    \u2192  renders as numbered items ( 1. 2. 3. )\n\n"
                  "You can override the list type here — for example,\n"
                  "forcing all bullet lists to render as numbered lists.")

        lf = tk.Frame(inner, bg=BG2)
        lf.pack(fill="x", pady=(0, 8))
        list_rows = [
            ("Bullet list (UL) \u2192", self.ul_var,
             "Controls how Word bullet lists are output.\n\n"
             "ul  \u2192  <ul><li> — unordered list, renders as bullet points\n"
             "ol  \u2192  <ol><li> — ordered list, renders as 1. 2. 3.\n\n"
             "Leave as 'ul' unless you want to force bullets into a numbered list."),
            ("Numbered list (OL) \u2192", self.ol_var,
             "Controls how Word numbered lists are output.\n\n"
             "ol  \u2192  <ol><li> — ordered list, renders as 1. 2. 3.\n"
             "ul  \u2192  <ul><li> — unordered list, renders as bullet points\n\n"
             "Leave as 'ol' unless you want to force numbers into a bullet list."),
        ]
        for row_i, (lbl, var, tip_text) in enumerate(list_rows):
            tk.Label(lf, text=lbl, font=FONT, bg=BG2, fg=FG).grid(
                row=row_i, column=0, sticky="w", pady=3, padx=(0, 8))
            cb = ttk.Combobox(lf, textvariable=var, values=LIST_OPTS,
                              state="readonly", width=10)
            cb.grid(row=row_i, column=1, sticky="w")
            cb.bind("<<ComboboxSelected>>", lambda e: self._refresh_preview())
            tip(cb, tip_text)

        # ── Blockquote ────────────────────────────────────────
        self._sep(inner, "Blockquote Transform",
                  "Word paragraphs using the 'Quote', 'Block Text', or\n"
                  "'Intense Quote' styles are treated as blockquotes.\n\n"
                  "Choose how they appear in the HTML output.")

        bqf = tk.Frame(inner, bg=BG2)
        bqf.pack(fill="x", pady=(0, 8))
        tk.Label(bqf, text="Quote style \u2192", font=FONT, bg=BG2,
                 fg=FG).grid(row=0, column=0, sticky="w", padx=(0, 8), pady=2)
        cb_bq = ttk.Combobox(bqf, textvariable=self.bq_var, values=BQ_OPTS,
                              state="readonly", width=22)
        cb_bq.grid(row=0, column=1, sticky="w")
        cb_bq.bind("<<ComboboxSelected>>", lambda e: self._refresh_preview())
        tip(cb_bq,
            "blockquote\n"
            "  Semantic HTML blockquote — renders with a left border\n"
            "  and indented styling. Best for pull quotes and citations.\n\n"
            "p\n"
            "  Plain paragraph — no special blockquote styling.\n\n"
            'div class="callout"\n'
            "  Brightspace callout box — styled highlight panel.\n"
            "  Requires your Brightspace theme CSS to render correctly.\n\n"
            "(skip)\n"
            "  Remove quoted paragraphs from the output entirely.")

        cb_hr = tk.Checkbutton(bqf, text="Add <hr> dividers inside blockquote",
                               variable=self.bq_hr_var, bg=BG2, fg=FG,
                               activebackground=BG2, activeforeground=FG,
                               selectcolor=BG3, font=FONT, bd=0,
                               highlightthickness=0,
                               command=self._refresh_preview)
        cb_hr.grid(row=1, column=0, columnspan=2, sticky="w", pady=(4, 0))
        tip(cb_hr,
            "Wraps each blockquote with <hr> lines above and below,\n"
            "creating a visual separator. Only applies when output\n"
            "is set to 'blockquote'.")

        # ── Accordion ─────────────────────────────────────────
        self._sep(inner, "Accordion Card Title Style",
                  "Single-cell Word tables are automatically converted\n"
                  "to D2L accordion components.\n\n"
                  "Each paragraph using the selected heading style inside\n"
                  "the table cell becomes a collapsible card title.\n"
                  "All body paragraphs that follow it (until the next\n"
                  "matching heading) become that card's content.")

        af = tk.Frame(inner, bg=BG2)
        af.pack(fill="x", pady=(0, 8))
        tk.Label(af, text="Trigger heading \u2192", font=FONT, bg=BG2,
                 fg=FG).grid(row=0, column=0, sticky="w", padx=(0, 8))
        cb_acc = ttk.Combobox(af, textvariable=self.acc_head_var,
                              values=[f"Heading {n}" for n in range(1, 7)],
                              state="readonly", width=12)
        cb_acc.grid(row=0, column=1, sticky="w")
        cb_acc.bind("<<ComboboxSelected>>", lambda e: self._refresh_preview())
        tip(cb_acc,
            "The Word heading style that marks the start of each\n"
            "accordion card inside a single-cell table.\n\n"
            "Default: Heading 4\n\n"
            "Change this if your document uses a different heading\n"
            "level inside tables for accordion titles.")

        # ── CSS Upload ────────────────────────────────────────
        self._sep(inner, "Preview CSS (optional)",
                  "Load a CSS file to style the preview panel.\n\n"
                  "Loading your Brightspace theme stylesheet gives an\n"
                  "accurate preview of how the page will look to students.\n\n"
                  "CSS is only embedded in saved output when\n"
                  "'Wrap in full HTML document' is enabled.")

        cf = tk.Frame(inner, bg=BG2)
        cf.pack(fill="x", pady=(0, 8))
        self.css_label = tk.Label(
            cf, text="No CSS loaded \u2014 default preview styles",
            font=FONT_S, fg=FG3, bg=BG2, wraplength=300, justify="left")
        self.css_label.pack(anchor="w")
        br = tk.Frame(cf, bg=BG2)
        br.pack(anchor="w", pady=(6, 0))

        def _dark_btn(parent, text, cmd, tip_text):
            b = tk.Button(parent, text=text, font=FONT, bg=BG3, fg=FG,
                          relief="flat", bd=0, padx=10, pady=5, cursor="hand2",
                          activebackground=BORDER, activeforeground=FG,
                          command=cmd)
            b.pack(side="left", padx=(0, 6))
            tip(b, tip_text)
            return b

        _dark_btn(br, "Upload CSS\u2026", self._pick_css,
                  "Load a .css file to style the preview panel\n"
                  "and full-document HTML output.")
        _dark_btn(br, "Clear", self._clear_css,
                  "Remove the loaded CSS and revert to\n"
                  "the built-in default preview styles.")

    # ── Preset management ────────────────────────────────────

    def _preset_snapshot(self):
        """Return a dict of all current settings (for saving as a preset)."""
        return {
            "headings":    {ws: v.get() for ws, v in self.heading_vars.items()},
            "ul":          self.ul_var.get(),
            "ol":          self.ol_var.get(),
            "bq":          self.bq_var.get(),
            "bq_hr":       self.bq_hr_var.get(),
            "acc_heading": self.acc_head_var.get(),
            "css_path":    self.custom_css_path or "",
        }

    def _apply_preset_dict(self, d):
        """Apply a preset dict to the current UI controls."""
        for ws, var in self.heading_vars.items():
            if ws in d.get("headings", {}):
                var.set(d["headings"][ws])
        self.ul_var.set(d.get("ul", "ul"))
        self.ol_var.set(d.get("ol", "ol"))
        self.bq_var.set(d.get("bq", "blockquote"))
        self.bq_hr_var.set(d.get("bq_hr", False))
        self.acc_head_var.set(d.get("acc_heading", ACCORDION_HEADING))
        css_path = d.get("css_path", "")
        if css_path and Path(css_path).exists():
            self.custom_css = Path(css_path).read_text(encoding="utf-8", errors="replace")
            self.custom_css_path = css_path
            self.css_label.config(text=f"\u2705  {Path(css_path).name}", fg=SUCCESS)
        elif not css_path:
            self.custom_css = ""
            self.custom_css_path = None
            self.css_label.config(
                text="No CSS loaded \u2014 default preview styles", fg=FG3)
        self._refresh_preview()

    def _refresh_preset_combo(self, select=None):
        """Rebuild the preset dropdown values; optionally select a name."""
        names = sorted(self.presets.keys())
        self.preset_combo.configure(values=names)
        if select and select in names:
            self.preset_var.set(select)
        elif names:
            self.preset_var.set(names[0])
        else:
            self.preset_var.set("")

    def _load_preset(self):
        name = self.preset_var.get()
        if name and name in self.presets:
            self._apply_preset_dict(self.presets[name])
            self._save_config()   # persist the selected preset name

    def _save_preset_dialog(self):
        """Open a small dialog to name and save a preset."""
        dlg = tk.Toplevel(self)
        dlg.title("Save Preset")
        dlg.configure(bg=BG2)
        dlg.resizable(False, False)
        dlg.grab_set()

        tk.Label(dlg, text="Preset name:", font=FONT, bg=BG2, fg=FG,
                 padx=14, pady=(14, 0)).pack(anchor="w")

        existing = sorted(self.presets.keys())
        name_var = tk.StringVar(value=self.preset_var.get() or "")
        combo = ttk.Combobox(dlg, textvariable=name_var, values=existing,
                             font=FONT, width=28)
        combo.pack(padx=14, pady=6)
        combo.focus_set()

        msg = tk.Label(dlg, text="", font=FONT_S, bg=BG2, fg=WARN,
                       padx=14, wraplength=240)
        msg.pack(anchor="w")

        def _do_save():
            name = name_var.get().strip()
            if not name:
                msg.config(text="Please enter a name.")
                return
            overwriting = name in self.presets
            self.presets[name] = self._preset_snapshot()
            self._save_presets_file()
            self._refresh_preset_combo(select=name)
            dlg.destroy()
            verb = "updated" if overwriting else "saved"
            self.status_var.set(f"\u2705  Preset \"{name}\" {verb}.")

        bf = tk.Frame(dlg, bg=BG2)
        bf.pack(pady=(4, 14), padx=14, anchor="e")
        tk.Button(bf, text="Save", font=FONT, bg=ACCENT, fg="#fff",
                  relief="flat", bd=0, padx=12, pady=5, cursor="hand2",
                  activebackground=ACCENT2, activeforeground="#fff",
                  command=_do_save).pack(side="left", padx=(0, 6))
        tk.Button(bf, text="Cancel", font=FONT, bg=BG3, fg=FG,
                  relief="flat", bd=0, padx=12, pady=5, cursor="hand2",
                  activebackground=BORDER, activeforeground=FG,
                  command=dlg.destroy).pack(side="left")
        dlg.bind("<Return>", lambda e: _do_save())
        dlg.bind("<Escape>", lambda e: dlg.destroy())

    def _rename_preset_dialog(self):
        name = self.preset_var.get()
        if not name:
            messagebox.showinfo("No preset selected",
                                "Select a preset from the dropdown first.")
            return
        dlg = tk.Toplevel(self)
        dlg.title("Rename Preset")
        dlg.configure(bg=BG2)
        dlg.resizable(False, False)
        dlg.grab_set()

        tk.Label(dlg, text=f"Rename \"{name}\" to:", font=FONT, bg=BG2, fg=FG,
                 padx=14, pady=(14, 0)).pack(anchor="w")
        new_var = tk.StringVar(value=name)
        entry = tk.Entry(dlg, textvariable=new_var, font=FONT,
                         bg=BG3, fg=FG, insertbackground=FG,
                         relief="flat", bd=4, width=28)
        entry.pack(padx=14, pady=6)
        entry.focus_set()
        entry.select_range(0, "end")

        msg = tk.Label(dlg, text="", font=FONT_S, bg=BG2, fg=WARN, padx=14)
        msg.pack(anchor="w")

        def _do_rename():
            new_name = new_var.get().strip()
            if not new_name:
                msg.config(text="Name cannot be empty.")
                return
            if new_name == name:
                dlg.destroy()
                return
            if new_name in self.presets:
                msg.config(text="A preset with that name already exists.")
                return
            self.presets[new_name] = self.presets.pop(name)
            self._save_presets_file()
            self._refresh_preset_combo(select=new_name)
            dlg.destroy()
            self.status_var.set(f'\u2705  Renamed to "{new_name}".')

        bf = tk.Frame(dlg, bg=BG2)
        bf.pack(pady=(4, 14), padx=14, anchor="e")
        tk.Button(bf, text="Rename", font=FONT, bg=ACCENT, fg="#fff",
                  relief="flat", bd=0, padx=12, pady=5, cursor="hand2",
                  activebackground=ACCENT2, activeforeground="#fff",
                  command=_do_rename).pack(side="left", padx=(0, 6))
        tk.Button(bf, text="Cancel", font=FONT, bg=BG3, fg=FG,
                  relief="flat", bd=0, padx=12, pady=5, cursor="hand2",
                  activebackground=BORDER, activeforeground=FG,
                  command=dlg.destroy).pack(side="left")
        dlg.bind("<Return>", lambda e: _do_rename())
        dlg.bind("<Escape>", lambda e: dlg.destroy())

    def _delete_preset(self):
        name = self.preset_var.get()
        if not name:
            messagebox.showinfo("No preset selected",
                                "Select a preset from the dropdown first.")
            return
        if not messagebox.askyesno("Delete preset",
                                   f'Delete preset "{name}"?',
                                   icon="warning"):
            return
        self.presets.pop(name, None)
        self._save_presets_file()
        self._refresh_preset_combo()
        self.status_var.set(f'\u274c  Preset "{name}" deleted.')

    def _save_presets_file(self):
        try:
            PRESETS_FILE.write_text(
                json.dumps(self.presets, indent=2), encoding="utf-8")
        except Exception:
            pass

    def _load_presets_file(self):
        try:
            self.presets = json.loads(
                PRESETS_FILE.read_text(encoding="utf-8"))
        except Exception:
            self.presets = {}

        # First launch: seed a sensible default so the dropdown is never empty
        if not self.presets:
            self.presets["Profile 1"] = self._preset_snapshot()
            self._save_presets_file()

        # Restore the last-used preset; fall back to the first available name
        last = getattr(self, "_last_preset_name", "")
        if last and last in self.presets:
            self._refresh_preset_combo(select=last)
        else:
            self._refresh_preset_combo()

    # ── Conversion log ────────────────────────────────────────

    def _toggle_log(self):
        if self._log_expanded.get():
            self._log_body.pack_forget()
            self._log_expanded.set(False)
            self._log_arrow.config(text="▶")
        else:
            self._log_body.pack(fill="x", pady=(2, 0))
            self._log_expanded.set(True)
            self._log_arrow.config(text="▼")
            # Scroll to the bottom so most recent entries are visible
            self._log_text.see("end")

    def _update_log(self, log_entries):
        """Populate the log panel from a fresh list of log dicts.

        Called after every conversion. Clears the previous run's entries,
        then auto-expands the panel if there are any warnings or errors.
        """
        t = self._log_text
        t.config(state="normal")
        t.delete("1.0", "end")

        n_warn  = sum(1 for e in log_entries if e["level"] == "warn")
        n_error = sum(1 for e in log_entries if e["level"] == "error")

        ICONS = {"info": "·", "warn": "⚠", "error": "✕"}
        for entry in log_entries:
            lvl  = entry["level"]
            icon = ICONS.get(lvl, "·")
            t.insert("end", f" {icon} ", lvl)
            t.insert("end", entry["msg"] + "\n", lvl)

        t.config(state="disabled")

        # Update the badge summary
        if n_error:
            badge = f"  {n_error} error{'s' if n_error != 1 else ''}"
            self._log_badge.config(text=badge, fg=ERR)
        elif n_warn:
            badge = f"  {n_warn} warning{'s' if n_warn != 1 else ''}"
            self._log_badge.config(text=badge, fg=WARN)
        else:
            self._log_badge.config(text="  OK", fg=SUCCESS)

        # Auto-expand on warnings/errors; auto-collapse on clean runs
        if n_warn or n_error:
            if not self._log_expanded.get():
                self._toggle_log()
        else:
            if self._log_expanded.get():
                self._toggle_log()

    def _log_append(self, level, msg):
        """Append a single entry to the log panel (used by _run_convert for
        save-time info like file paths and sizes)."""
        t = self._log_text
        t.config(state="normal")
        ICONS = {"info": "·", "warn": "⚠", "error": "✕"}
        icon = ICONS.get(level, "·")
        t.insert("end", f" {icon} ", level)
        t.insert("end", msg + "\n", level)
        t.config(state="disabled")
        t.see("end")

    def _sep(self, parent, label, tooltip_text=None):
        lbl = tk.Label(parent, text=label, font=("Segoe UI", 10, "bold"),
                       bg=BG2, fg=ACCENT2)
        lbl.pack(anchor="w", pady=(14, 3))
        if tooltip_text:
            tip(lbl, tooltip_text)
        tk.Frame(parent, bg=BORDER, height=1).pack(fill="x", pady=(0, 8))

    # ── Preview panel ─────────────────────────────────────────

    def _build_preview_panel(self, parent):
        # Toolbar
        bar = tk.Frame(parent, bg=BG2)
        bar.pack(fill="x")
        tk.Frame(bar, bg=BORDER, height=1).pack(side="bottom", fill="x")

        tk.Label(bar, text="  Preview", font=FONT_B, bg=BG2,
                 fg=FG, pady=8).pack(side="left")

        def _mode_btn(text, value):
            rb = tk.Radiobutton(bar, text=f"  {text}  ",
                                variable=self.preview_mode, value=value,
                                bg=BG2, fg=FG2, font=FONT,
                                activebackground=BG2, activeforeground=FG,
                                selectcolor=BG3, indicatoron=False,
                                relief="flat", bd=0, padx=8, pady=5,
                                cursor="hand2",
                                command=self._refresh_preview)
            rb.pack(side="right", padx=4, pady=4)
            return rb

        _mode_btn("Rendered", "rendered")
        _mode_btn("HTML Source", "source")

        # Module navigation bar (hidden until modules are detected)
        self.mod_nav_frame = tk.Frame(bar, bg=BG2)
        # Will be packed when modules are found

        self._mod_prev_btn = tk.Button(
            self.mod_nav_frame, text="◀", font=FONT, bg=BG3, fg=FG,
            relief="flat", bd=0, padx=8, pady=4, cursor="hand2",
            activebackground=BORDER, activeforeground=FG,
            command=self._module_prev)
        self._mod_prev_btn.pack(side="left")

        self._mod_label_var = tk.StringVar(value="")
        tk.Label(self.mod_nav_frame, textvariable=self._mod_label_var,
                 bg=BG2, fg=FG, font=FONT, padx=8).pack(side="left")

        self._mod_next_btn = tk.Button(
            self.mod_nav_frame, text="▶", font=FONT, bg=BG3, fg=FG,
            relief="flat", bd=0, padx=8, pady=4, cursor="hand2",
            activebackground=BORDER, activeforeground=FG,
            command=self._module_next)
        self._mod_next_btn.pack(side="left")

        # Search bar (hidden by default, toggled with Ctrl+F)
        self.search_frame = tk.Frame(parent, bg=BG2)
        # Don't pack yet — shown on demand
        tk.Frame(self.search_frame, bg=BORDER, height=1).pack(side="bottom", fill="x")
        self._search_var      = tk.StringVar()
        self._search_matches  = []
        self._search_idx      = 0

        sf_inner = tk.Frame(self.search_frame, bg=BG2, padx=8, pady=5)
        sf_inner.pack(fill="x")
        tk.Label(sf_inner, text="Find:", font=FONT, bg=BG2, fg=FG2).pack(side="left")
        self._search_entry = tk.Entry(
            sf_inner, textvariable=self._search_var,
            font=FONT, bg=BG3, fg=FG, insertbackground=FG,
            relief="flat", bd=4, width=28)
        self._search_entry.pack(side="left", padx=(6, 4))
        self._search_entry.bind("<Return>",       lambda e: self._search_next())
        self._search_entry.bind("<Shift-Return>", lambda e: self._search_prev())
        self._search_entry.bind("<Escape>",       lambda e: self._hide_search())
        self._search_var.trace_add("write", lambda *_: self._search_run())

        self._search_count_lbl = tk.Label(sf_inner, text="", font=FONT_S,
                                          bg=BG2, fg=FG3, width=12, anchor="w")
        self._search_count_lbl.pack(side="left", padx=(2, 8))

        tk.Button(sf_inner, text="▲", font=FONT_S, bg=BG3, fg=FG,
                  relief="flat", bd=0, padx=6, pady=2, cursor="hand2",
                  activebackground=BORDER, activeforeground=FG,
                  command=self._search_prev).pack(side="left", padx=(0, 2))
        tk.Button(sf_inner, text="▼", font=FONT_S, bg=BG3, fg=FG,
                  relief="flat", bd=0, padx=6, pady=2, cursor="hand2",
                  activebackground=BORDER, activeforeground=FG,
                  command=self._search_next).pack(side="left", padx=(0, 8))
        tk.Button(sf_inner, text="✕", font=FONT_S, bg=BG2, fg=FG3,
                  relief="flat", bd=0, padx=6, pady=2, cursor="hand2",
                  activebackground=BG2, activeforeground=FG,
                  command=self._hide_search).pack(side="left")

        self._search_visible = False

        # Preview area — Text widget for source, Frame+scrollable for rendered
        self.preview_outer = tk.Frame(parent, bg=PRE_BG)
        self.preview_outer.pack(fill="both", expand=True)

        # Source pane (Text widget)
        self.source_frame = tk.Frame(self.preview_outer, bg=PRE_BG)
        self.preview_text = tk.Text(
            self.source_frame, font=("Cascadia Code", 9), wrap="none",
            bg=PRE_BG, fg=PRE_FG,
            insertbackground=FG, relief="flat", bd=0,
            padx=12, pady=12,
            selectbackground=ACCENT, selectforeground=FG)
        vsb = ttk.Scrollbar(self.source_frame, orient="vertical",
                            command=self.preview_text.yview)
        hsb = ttk.Scrollbar(self.source_frame, orient="horizontal",
                            command=self.preview_text.xview)
        self.preview_text.configure(yscrollcommand=vsb.set,
                                    xscrollcommand=hsb.set)
        vsb.pack(side="right", fill="y")
        hsb.pack(side="bottom", fill="x")
        self.preview_text.pack(fill="both", expand=True)

        # ── Brightspace chrome: dark top navbar ───────────────────────────
        # Mimics the breadcrumb strip that appears above every D2L content page.
        D2L_NAV    = "#1a2632"
        D2L_NAV_FG = "#ffffff"
        D2L_SEP    = "#3a4a58"
        D2L_SHELL  = "#f5f5f5"

        bs_chrome = tk.Frame(self.preview_outer, bg=D2L_NAV)
        bs_chrome.pack(fill="x")

        bc_row = tk.Frame(bs_chrome, bg=D2L_NAV, padx=14, pady=8)
        bc_row.pack(fill="x")

        def _crumb(text, fg="#78909c", bold=False):
            font = ("Segoe UI", 9, "bold") if bold else ("Segoe UI", 9)
            tk.Label(bc_row, text=text, font=font,
                     bg=D2L_NAV, fg=fg, pady=0).pack(side="left")

        _crumb("🏠  Course Home")
        _crumb("  ›  ", fg="#546e7a")
        _crumb("Content")
        _crumb("  ›  ", fg="#546e7a")

        self._d2l_title_var = tk.StringVar(value="(no file loaded)")
        tk.Label(bc_row, textvariable=self._d2l_title_var,
                 font=("Segoe UI", 9, "bold"),
                 bg=D2L_NAV, fg=D2L_NAV_FG).pack(side="left")

        tk.Frame(self.preview_outer, bg=D2L_SEP, height=1).pack(fill="x")

        # ── Rendered pane: grey shell + white content well ─────────────────
        self.rendered_frame = tk.Frame(self.preview_outer, bg=D2L_SHELL)

        card_frame = tk.Frame(
            self.rendered_frame,
            bg="#ffffff",
            highlightbackground="#d0d0d0",
            highlightthickness=1,
        )
        card_frame.pack(fill="both", expand=True, padx=28, pady=24)

        self.rendered_text = tk.Text(
            card_frame, font=("Lato", 11), wrap="word",
            bg="#ffffff", fg="#212121",
            relief="flat", bd=0, padx=32, pady=28,
            spacing1=0, spacing2=0, spacing3=0,
            selectbackground="#b3d4f5", selectforeground="#212121",
            cursor="arrow", state="disabled")
        rvsb = ttk.Scrollbar(card_frame, orient="vertical",
                             command=self.rendered_text.yview)
        self.rendered_text.configure(yscrollcommand=rvsb.set)
        rvsb.pack(side="right", fill="y")
        self.rendered_text.pack(fill="both", expand=True)
        self._configure_rendered_tags()

        # Bottom bar
        bot = tk.Frame(parent, bg=BG2)
        bot.pack(fill="x")
        tk.Frame(bot, bg=BORDER, height=1).pack(side="top", fill="x")

        btn_browser = tk.Button(bot,
                                text="Open in Browser  \U0001f310",
                                font=FONT, bg=BG2, fg=FG2, relief="flat", bd=0,
                                padx=10, pady=6, cursor="hand2",
                                activebackground=BG3, activeforeground=FG,
                                command=self._open_in_browser)
        btn_browser.pack(side="right", padx=8, pady=4)
        tip(btn_browser, "Open the converted HTML in your default browser.")

        self._copy_btn = tk.Button(bot,
                                   text="Copy HTML  \U0001f4cb",
                                   font=FONT, bg=BG2, fg=FG2, relief="flat", bd=0,
                                   padx=10, pady=6, cursor="hand2",
                                   activebackground=BG3, activeforeground=FG,
                                   command=self._copy_html)
        self._copy_btn.pack(side="left", padx=8, pady=4)
        tip(self._copy_btn,
            "Copy the HTML snippet to the clipboard (Ctrl+Shift+C).\n"
            "Pastes cleanly into Brightspace's HTML editor.\n"
            "When modules are detected, copies the currently shown module.")

        self._show_pane("source")
        self._set_preview("Select a .docx file to see a preview here.")

        # Bind Ctrl+F globally to toggle search
        self.bind_all("<Control-f>",       lambda e: self._toggle_search())
        self.bind_all("<Control-F>",       lambda e: self._toggle_search())
        self.bind_all("<Control-Shift-c>", lambda e: self._copy_html())
        self.bind_all("<Control-Shift-C>", lambda e: self._copy_html())

    def _configure_rendered_tags(self):
        """Set up Text widget tags styled to match a Brightspace content page.

        spacing1 = pixels above the first line of a tagged run
        spacing3 = pixels below the last line of a tagged run
        Because newlines are always inserted untagged, these values apply
        only to content lines — no stacking or double-spacing.
        """
        t = self.rendered_text

        # ── Headings ──────────────────────────────────────────────────────
        t.tag_configure("h1", font=("Lato", 24, "bold"),
                        foreground="#212121", spacing1=8, spacing3=3)
        t.tag_configure("h2", font=("Lato", 19, "bold"),
                        foreground="#212121", spacing1=7, spacing3=3)
        t.tag_configure("h3", font=("Lato", 16, "bold"),
                        foreground="#212121", spacing1=5, spacing3=2)
        t.tag_configure("h4", font=("Lato", 14, "bold"),
                        foreground="#212121", spacing1=4, spacing3=2)
        t.tag_configure("h5", font=("Lato", 12, "bold"),
                        foreground="#424242", spacing1=3, spacing3=1)
        t.tag_configure("h6", font=("Lato", 11, "bold"),
                        foreground="#616161", spacing1=3, spacing3=1)

        # ── Body copy ─────────────────────────────────────────────────────
        t.tag_configure("p",  font=("Lato", 11),
                        foreground="#212121", spacing1=0, spacing3=4)

        # ── Lists ─────────────────────────────────────────────────────────
        t.tag_configure("li", font=("Lato", 11),
                        foreground="#212121", spacing1=2, spacing3=2,
                        lmargin1=28, lmargin2=44)

        # ── Blockquote ────────────────────────────────────────────────────
        t.tag_configure("blockquote", font=("Lato", 11, "italic"),
                        foreground="#424242", lmargin1=28, lmargin2=28,
                        spacing1=2, spacing3=2)

        # ── Inline formatting ─────────────────────────────────────────────
        t.tag_configure("bold",          font=("Lato", 11, "bold"))
        t.tag_configure("italic",        font=("Lato", 11, "italic"))
        t.tag_configure("underline",     font=("Lato", 11, "underline"))
        t.tag_configure("strikethrough", font=("Lato", 11),
                        overstrike=True, foreground="#757575")

        # ── Links ─────────────────────────────────────────────────────────
        t.tag_configure("link", foreground="#006fbf",
                        font=("Lato", 11, "underline"))

        # ── Misc ──────────────────────────────────────────────────────────
        t.tag_configure("hr",  foreground="#e0e0e0", spacing1=4, spacing3=4)
        t.tag_configure("img", foreground="#006fbf",
                        font=("Lato", 10, "italic"),
                        spacing1=4, spacing3=4, lmargin1=4)

        # ── Accordion card title ───────────────────────────────────────────
        t.tag_configure("card_title", font=("Lato", 11, "bold"),
                        foreground="#006fbf", spacing1=4, spacing3=1,
                        lmargin1=12, lmargin2=12)
        t.tag_configure("card_body",  font=("Lato", 11),
                        foreground="#212121", spacing1=1, spacing3=1,
                        lmargin1=20, lmargin2=20)

    def _toggle_search(self):
        if self._search_visible:
            self._hide_search()
        else:
            self._show_search()

    def _show_search(self):
        if self._search_visible:
            return
        # Insert search bar between toolbar and preview_outer
        self.search_frame.pack(fill="x", before=self.preview_outer)
        self._search_visible = True
        self._search_entry.focus_set()
        self._search_entry.select_range(0, "end")
        self._search_run()

    def _hide_search(self):
        if not self._search_visible:
            return
        self.search_frame.pack_forget()
        self._search_visible = False
        self._search_clear_highlights()
        self._search_matches = []
        self._search_count_lbl.config(text="")
        # Return focus to preview
        if self.preview_mode.get() == "source":
            self.preview_text.focus_set()
        else:
            self.rendered_text.focus_set()

    def _active_text_widget(self):
        """Return whichever Text widget is currently visible."""
        if self.preview_mode.get() == "source":
            return self.preview_text
        return self.rendered_text

    def _search_clear_highlights(self):
        for tw in (self.preview_text, self.rendered_text):
            tw.tag_remove("search_match",   "1.0", "end")
            tw.tag_remove("search_current", "1.0", "end")

    def _search_run(self):
        """Run the search on the active widget and highlight all matches."""
        self._search_clear_highlights()
        self._search_matches = []
        self._search_idx = 0
        query = self._search_var.get()
        if not query:
            self._search_count_lbl.config(text="")
            return

        tw = self._active_text_widget()
        tw.tag_configure("search_match",   background="#fbbf24", foreground="#1e293b")
        tw.tag_configure("search_current", background="#f97316", foreground="#fff")

        start = "1.0"
        while True:
            pos = tw.search(query, start, stopindex="end", nocase=True)
            if not pos:
                break
            end = f"{pos}+{len(query)}c"
            tw.tag_add("search_match", pos, end)
            self._search_matches.append((pos, end))
            start = end

        count = len(self._search_matches)
        if count == 0:
            self._search_count_lbl.config(text="No results", fg=ERR)
        else:
            self._search_idx = 0
            self._search_highlight_current()

    def _search_highlight_current(self):
        if not self._search_matches:
            return
        tw = self._active_text_widget()
        # Clear previous current highlight
        tw.tag_remove("search_current", "1.0", "end")
        pos, end = self._search_matches[self._search_idx]
        tw.tag_add("search_current", pos, end)
        tw.see(pos)
        total = len(self._search_matches)
        self._search_count_lbl.config(
            text=f"{self._search_idx + 1} / {total}", fg=FG2)

    def _search_next(self):
        if not self._search_matches:
            return
        self._search_idx = (self._search_idx + 1) % len(self._search_matches)
        self._search_highlight_current()

    def _search_prev(self):
        if not self._search_matches:
            return
        self._search_idx = (self._search_idx - 1) % len(self._search_matches)
        self._search_highlight_current()

    def _render_html_to_text(self, body_html):
        """Parse body_html and insert styled text into the rendered_text widget.

        Key spacing rule: newline characters are always inserted without any
        tag so that spacing1/spacing3 only applies to lines that carry actual
        content — preventing double/triple spacing from stacked tag spacings.
        """
        import html as html_lib
        import html.parser

        t = self.rendered_text
        t.config(state="normal")
        t.delete("1.0", "end")

        class _Parser(html.parser.HTMLParser):
            def __init__(self, widget):
                super().__init__()
                self.w          = widget
                self.tag_stack  = []
                self.list_stack = []   # stack of ("ul"|"ol", counter)
                self.skip       = False
                self._pending_nl = 0   # untagged newlines queued to insert

            def _flush_nl(self):
                """Emit any queued plain newlines before content."""
                if self._pending_nl:
                    self.w.insert("end", "\n" * self._pending_nl)
                    self._pending_nl = 0

            def _nl(self, n=1):
                """Queue n newlines (merged so tags never receive \n)."""
                self._pending_nl = max(self._pending_nl, n)

            def _insert(self, text, *tags):
                if self.skip or not text:
                    return
                self._flush_nl()
                self.w.insert("end", text, tags)

            def handle_starttag(self, tag, attrs):
                self.tag_stack.append(tag)
                attrs = dict(attrs)
                if tag in ("h1","h2","h3","h4","h5","h6"):
                    # Extra blank line before headings (except at very start)
                    idx = self.w.index("end-1c")
                    if idx != "1.0":
                        self._nl(2)
                elif tag == "p":
                    pass   # gap handled on </p>
                elif tag == "ul":
                    self.list_stack.append(("ul", [0]))
                    self._nl(1)
                elif tag == "ol":
                    self.list_stack.append(("ol", [0]))
                    self._nl(1)
                elif tag == "li":
                    self._flush_nl()
                    if self.list_stack:
                        kind, ctr = self.list_stack[-1]
                        ctr[0] += 1
                        bullet = f"  {ctr[0]}. " if kind == "ol" else "  \u2022 "
                        self.w.insert("end", bullet, ("li",))
                elif tag == "blockquote":
                    self._nl(1)
                elif tag == "hr":
                    self._nl(1)
                    self._flush_nl()
                    self.w.insert("end", "\u2500" * 40, ("hr",))
                    self._nl(1)
                elif tag == "img":
                    src   = attrs.get("src", "")
                    alt   = attrs.get("alt", "")
                    label = alt or (Path(src).name if src else "image")
                    self._nl(1)
                    self._flush_nl()
                    self.w.insert("end", f"  \U0001f5bc\ufe0f  [{label}]", ("img",))
                    self._nl(1)
                elif tag in ("table","thead","tbody","tr","th","td"):
                    if tag == "tr":
                        self._nl(1)
                    elif tag == "td":
                        self._flush_nl()
                        self.w.insert("end", "  | ", ("p",))
                    elif tag == "th":
                        self._flush_nl()
                        self.w.insert("end", "  | ", ("p","bold"))
                elif tag in ("script","style","head"):
                    self.skip = True

            def handle_endtag(self, tag):
                if tag in ("script","style","head"):
                    self.skip = False
                if self.tag_stack and self.tag_stack[-1] == tag:
                    self.tag_stack.pop()
                if tag in ("h1","h2","h3","h4","h5","h6"):
                    self._nl(1)
                elif tag == "p":
                    self._nl(2)
                elif tag == "li":
                    self._nl(1)
                elif tag in ("ul","ol"):
                    if self.list_stack:
                        self.list_stack.pop()
                    self._nl(1)
                elif tag == "blockquote":
                    self._nl(1)

            def handle_data(self, data):
                if self.skip:
                    return
                text = html_lib.unescape(data)
                # Strip leading/trailing newlines from raw HTML whitespace
                text = text.strip("\n\r")
                if not text:
                    return
                cur = self.tag_stack
                if any(h in cur for h in ("h1","h2","h3","h4","h5","h6")):
                    htag = next(h for h in cur
                                if h in ("h1","h2","h3","h4","h5","h6"))
                    self._insert(text, htag)
                elif "strong" in cur or "b" in cur:
                    self._insert(text, "bold")
                elif "em" in cur or "i" in cur:
                    self._insert(text, "italic")
                elif "a" in cur:
                    self._insert(text, "link")
                elif "blockquote" in cur:
                    self._insert(text, "blockquote")
                elif "li" in cur:
                    self._insert(text, "li")
                elif "td" in cur:
                    self._insert(text, "p")
                elif "th" in cur:
                    self._insert(text, "p", "bold")
                else:
                    self._insert(text, "p")

        _Parser(t).feed(body_html)
        # Flush any trailing queued newlines, then lock
        t.config(state="disabled")

    def _show_pane(self, mode):
        """Switch between source and rendered panes."""
        if mode == "source":
            self.rendered_frame.pack_forget()
            self.source_frame.pack(fill="both", expand=True)
        else:
            self.source_frame.pack_forget()
            self.rendered_frame.pack(fill="both", expand=True)

    # ── Module navigation ─────────────────────────────────────

    def _update_module_nav(self):
        """Show/hide the module nav bar and update its label."""
        if self._modules and self.split_modules_var.get():
            total = len(self._modules)
            mod   = self._modules[self._module_idx]
            self._mod_label_var.set(f"Module {self._module_idx + 1} of {total}")
            tip(self._mod_prev_btn, f"Previous module")
            tip(self._mod_next_btn, f"Next module")
            self._mod_prev_btn.config(
                state="normal" if self._module_idx > 0 else "disabled",
                fg=FG if self._module_idx > 0 else FG3)
            self._mod_next_btn.config(
                state="normal" if self._module_idx < total - 1 else "disabled",
                fg=FG if self._module_idx < total - 1 else FG3)
            # Pack nav bar left of the "Preview" label if not already shown
            if not self.mod_nav_frame.winfo_ismapped():
                self.mod_nav_frame.pack(side="left", padx=(8, 0))
            self.module_info_var.set(
                f"📦  {total} module{'s' if total != 1 else ''} detected "
                f"— will save {total} separate file{'s' if total != 1 else ''}")
        else:
            self.mod_nav_frame.pack_forget()
            self._mod_label_var.set("")
            self.module_info_var.set("")

    def _module_prev(self):
        if self._module_idx > 0:
            self._module_idx -= 1
            self._show_module()

    def _module_next(self):
        if self._module_idx < len(self._modules) - 1:
            self._module_idx += 1
            self._show_module()

    def _show_module(self):
        """Render the currently-selected module in the preview."""
        self._update_module_nav()
        if not self._modules:
            return
        html = self._modules[self._module_idx]["html"]
        mode = self.preview_mode.get()
        self._show_pane(mode)
        if mode == "source":
            self._set_preview(html)
        else:
            self._render_html_to_text(html)
        if self._search_visible:
            self._search_run()

    # ── Actions ───────────────────────────────────────────────

    def _pick_css(self):
        path = filedialog.askopenfilename(
            title="Select Brightspace CSS File",
            filetypes=[("CSS files", "*.css"), ("All files", "*.*")])
        if path:
            try:
                self.custom_css = Path(path).read_text(
                    encoding="utf-8", errors="replace")
                self.custom_css_path = path
                self.css_label.config(
                    text=f"\u2705  {Path(path).name}", fg=SUCCESS)
                self._refresh_preview()
            except Exception as e:
                messagebox.showerror("CSS Error", str(e))

    def _clear_css(self):
        self.custom_css = ""
        self.custom_css_path = None
        self.css_label.config(
            text="No CSS loaded \u2014 default preview styles", fg=FG3)
        self._refresh_preview()

    def _collect_settings(self):
        hmap = {ws: v.get() for ws, v in self.heading_vars.items()
                if v.get() != "(skip)"}
        return {
            "heading_map":       hmap,
            "ul_transform":      self.ul_var.get(),
            "ol_transform":      self.ol_var.get(),
            "bq_transform":      self.bq_var.get(),
            "bq_hr":             self.bq_hr_var.get(),
            "accordion_heading": self.acc_head_var.get(),
        }

    def _refresh_preview(self, *_):
        if not self.selected_file:
            return
        try:
            settings                   = self._collect_settings()
            body_html, _imgs, log, _links = convert_docx(self.selected_file, settings)
            self._update_log(log)

            # Detect modules and update navigation state
            if self.split_modules_var.get():
                self._modules = split_modules(body_html)
            else:
                self._modules = []

            # Clamp the index in case the document changed
            if self._modules:
                self._module_idx = max(0, min(self._module_idx,
                                              len(self._modules) - 1))
                self._update_module_nav()
                self._show_module()
                return   # _show_module handles the rest
            else:
                self._update_module_nav()

            mode = self.preview_mode.get()
            self._show_pane(mode)
            if mode == "source":
                self._set_preview(body_html)
            else:
                self._render_html_to_text(body_html)
        except Exception as exc:
            self._show_pane("source")
            self._set_preview(f"Preview error:\n{exc}")
        # Re-run any active search so highlights are up to date
        if self._search_visible:
            self._search_run()

    def _set_preview(self, text):
        self.preview_text.config(state="normal")
        self.preview_text.delete("1.0", "end")
        self.preview_text.insert("1.0", text)
        self.preview_text.config(state="disabled")

    def _open_in_browser(self):
        if not self.selected_file:
            messagebox.showinfo("No file", "Please select a .docx file first.")
            return
        try:
            settings                  = self._collect_settings()
            body_html, img_map, _log, _links = convert_docx(self.selected_file, settings)
            css  = self.custom_css or DEFAULT_PREVIEW_CSS
            full = wrap_html(body_html,
                             title=Path(self.selected_file).stem,
                             css=css)
            tmp = tempfile.NamedTemporaryFile(
                mode="w", suffix=".html", delete=False, encoding="utf-8")
            tmp.write(full)
            tmp.close()
            # Write images alongside the temp HTML so the browser can load them
            if img_map:
                tmp_img_dir = Path(tmp.name).parent / "images"
                tmp_img_dir.mkdir(exist_ok=True)
                for fname, blob in img_map.items():
                    (tmp_img_dir / fname).write_bytes(blob)
            webbrowser.open(f"file://{tmp.name}")
        except Exception as exc:
            messagebox.showerror("Error", str(exc))

    def _copy_html(self):
        """Copy the current HTML snippet to the clipboard.

        When modules are detected copies only the currently-previewed module.
        Always copies a bare body snippet — no <!DOCTYPE> wrapper.
        """
        if not self.selected_file:
            self.status_var.set("⚠  No file loaded — nothing to copy.")
            return
        try:
            if self._modules and self.split_modules_var.get():
                html  = self._modules[self._module_idx]["html"]
                label = f"Module {self._module_idx + 1}"
            else:
                settings    = self._collect_settings()
                html, _, _l, _links = convert_docx(self.selected_file, settings)
                label       = "HTML"

            self.clipboard_clear()
            self.clipboard_append(html)

            # Flash the button to confirm
            self._copy_btn.config(text="✓ Copied!", fg=SUCCESS)
            self.after(1500, lambda: self._copy_btn.config(
                text="Copy HTML  \U0001f4cb", fg=FG2))

            self.status_var.set(f"📋  {label} copied to clipboard.")
        except Exception as exc:
            messagebox.showerror("Copy error", str(exc))

    def _run_convert(self):
        if not self.selected_file:
            messagebox.showwarning("No file", "Please select a .docx file first.")
            return
        try:
            settings                    = self._collect_settings()
            body_html, img_map, log, link_annotations = convert_docx(self.selected_file, settings)
            css      = self.custom_css or DEFAULT_PREVIEW_CSS
            src_path = Path(self.selected_file)

            # Log any save-time entries (file sizes etc.) after we know the paths
            self._update_log(log)   # show conversion log immediately

            def _write_images(output_dir):
                """Write img_map blobs into output_dir/images/. Returns count."""
                if not img_map:
                    return 0
                img_dir = output_dir / "images"
                img_dir.mkdir(exist_ok=True)
                for fname, blob in img_map.items():
                    (img_dir / fname).write_bytes(blob)
                return len(img_map)

            def _img_note(n):
                if n == 0:
                    return ""
                return f"  ({n} image{'s' if n != 1 else ''} → images/)"

            # ── Module-split output ───────────────────────────
            if self.split_modules_var.get():
                modules = split_modules(body_html)
            else:
                modules = []

            if modules:
                if self.save_next_var.get():
                    out_dir = src_path.parent
                else:
                    out_dir_str = filedialog.askdirectory(
                        title="Choose folder to save module files")
                    if not out_dir_str:
                        self.status_var.set("Save cancelled.")
                        return
                    out_dir = Path(out_dir_str)

                stem  = src_path.stem
                saved = []
                for mod in modules:
                    num     = mod["number"]
                    content = mod["html"]
                    if self.full_html_var.get():
                        content = wrap_html(content, title=mod["title"], css=css)
                    fname = out_dir / f"{stem}_module_{num:02d}.html"
                    fname.write_text(content, encoding="utf-8")
                    saved.append(fname.name)
                    self._log_append("info",
                        f"Saved {fname.name} ({fname.stat().st_size // 1024} KB)")

                n_imgs = _write_images(out_dir)
                if n_imgs:
                    self._log_append("info",
                        f"Images written to {out_dir / 'images'}")
                if link_annotations:
                    lp = out_dir / f"{stem}_image_links.txt"
                    if _write_image_links_file(link_annotations, lp):
                        self._log_append("info",
                            f"Image links saved to {lp.name}")
                self._save_config()
                short = ", ".join(saved[:3]) + (" …" if len(saved) > 3 else "")
                self.status_var.set(
                    f"✅  Saved {len(saved)} module file"
                    f"{'s' if len(saved) != 1 else ''}: {short}"
                    f"{_img_note(n_imgs)}")
                self._refresh_preview()
                return

            # ── Single-file output ────────────────────────────
            if self.full_html_var.get():
                output = wrap_html(body_html, title=src_path.stem, css=css)
            else:
                output = body_html

            if self.save_next_var.get():
                out_path = src_path.with_suffix(".html")
            else:
                out_path = filedialog.asksaveasfilename(
                    title="Save HTML as",
                    defaultextension=".html",
                    filetypes=[("HTML files", "*.html"), ("All files", "*.*")],
                    initialfile=src_path.stem + ".html")
                if not out_path:
                    self.status_var.set("Save cancelled.")
                    return
                out_path = Path(out_path)

            out_path.write_text(output, encoding="utf-8")
            n_imgs = _write_images(out_path.parent)
            self._log_append("info",
                f"Saved {out_path.name} ({out_path.stat().st_size // 1024} KB)")
            if n_imgs:
                self._log_append("info",
                    f"Images written to {out_path.parent / 'images'}")
            if link_annotations:
                lp = out_path.parent / f"{src_path.stem}_image_links.txt"
                if _write_image_links_file(link_annotations, lp):
                    self._log_append("info",
                        f"Image links saved to {lp.name}")
            self._save_config()
            self.status_var.set(
                f"✅  Saved: {out_path.name}{_img_note(n_imgs)}")
            self._refresh_preview()
        except Exception as exc:
            messagebox.showerror("Conversion error", str(exc))
            self.status_var.set(f"❌  Error: {exc}")

    # ── Config persistence ────────────────────────────────────

    def _save_config(self):
        cfg = {
            "headings":       {ws: v.get() for ws, v in self.heading_vars.items()},
            "ul":             self.ul_var.get(),
            "ol":             self.ol_var.get(),
            "bq":             self.bq_var.get(),
            "bq_hr":          self.bq_hr_var.get(),
            "acc_heading":    self.acc_head_var.get(),
            "full_html":      self.full_html_var.get(),
            "save_next":      self.save_next_var.get(),
            "split_modules":  self.split_modules_var.get(),
            "css_path":       self.custom_css_path or "",
            "last_preset":    self.preset_var.get(),
        }
        try:
            CONFIG_FILE.write_text(json.dumps(cfg, indent=2), encoding="utf-8")
        except Exception:
            pass

    def _load_config(self):
        try:
            cfg = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
            for ws, var in self.heading_vars.items():
                if ws in cfg.get("headings", {}):
                    var.set(cfg["headings"][ws])
            self.ul_var.set(cfg.get("ul", "ul"))
            self.ol_var.set(cfg.get("ol", "ol"))
            self.bq_var.set(cfg.get("bq", "blockquote"))
            self.bq_hr_var.set(cfg.get("bq_hr", False))
            self.acc_head_var.set(cfg.get("acc_heading", ACCORDION_HEADING))
            self.full_html_var.set(cfg.get("full_html", False))
            self.save_next_var.set(cfg.get("save_next", True))
            self.split_modules_var.set(cfg.get("split_modules", True))
            css_path = cfg.get("css_path", "")
            if css_path and Path(css_path).exists():
                self.custom_css = Path(css_path).read_text(
                    encoding="utf-8", errors="replace")
                self.custom_css_path = css_path
                self.css_label.config(
                    text=f"\u2705  {Path(css_path).name}", fg=SUCCESS)
            # Restore the last-used preset selection (applied after presets load)
            self._last_preset_name = cfg.get("last_preset", "")
        except Exception:
            self._last_preset_name = ""

# ═══════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    try:
        # tkinterdnd2 may have just been installed by the bootstrap — re-import
        import tkinterdnd2
        ConverterApp.__bases__ = (tkinterdnd2.TkinterDnD.Tk,)
    except ImportError:
        pass  # drag-and-drop gracefully unavailable; click still works

    app = ConverterApp()
    app.mainloop()